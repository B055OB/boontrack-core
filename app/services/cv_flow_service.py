import os
import re
import json
import asyncio
import tempfile
from typing import Dict, Any, Optional

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from aiogram import types
from aiogram.types import InlineKeyboardMarkup, InlineKeyboardButton, InputFile

from app.core.bot import bot, send_chunked_message
from app.core.database import (
    save_dropoff,
    save_cv_version,
    track_event,
    check_user_paid,
    count_referrals
)
from app.services.ai_gateway import ai_gateway
from app.services.brain_engine import BrainEngine
from app.repositories.session_repository import SessionRepository
from app.engines.cv_review_engine import cv_review_engine
from app.services.cv_review_service import cv_review_service
from app.services.cloudflare_service import get_user_slug

TOTAL_STEPS = 9
REQUIRED_REFERRALS = 5

_session_repo = SessionRepository()
_brain_engine = BrainEngine(session_repo=_session_repo, ai_gateway=ai_gateway)

CLOSING_WORDS = [
    "Mantap! Terus melangkah ya,",
    "Keren banget progress-mu,",
    "Langkah bagus, selangkah lebih dekat!",
    "Bagus sekali, profilmu makin matang,",
    "Sip, data sudah tersimpan dengan aman,"
]

def get_progress_bar(step: int) -> str:
    filled = "🟩" * step
    empty = "⬜" * (TOTAL_STEPS - step)
    percent = int((step / TOTAL_STEPS) * 100)
    return f"{filled}{empty} <b>{percent}%</b>"

def clean_val(val: Any) -> str:
    if not val:
        return ""
    v = str(val).strip().lower()
    if v in ["-", "skip", "lewati", "tidak ada", "ga ada", "ngga ada", "belum ada", "lupa", "kosong"]:
        return ""
    return str(val).strip()

def get_question_text(step: int, target_lang: str = "ID", status_kerja: str = "Berpengalaman") -> str:
    is_full_en = target_lang == "EN"
    is_fresh = "fresh" in str(status_kerja).lower()
    
    if is_full_en:
        questions = {
            1: "👤 <b>What is your Full Name?</b>\n<i>(Recommended official name for the top of your CV)</i>",
            2: "📧 <b>Your Active Email?</b>\n<i>(Email regularly checked for recruiter replies)</i>",
            3: (
                "💼 <b>Your Work / Organization / Freelance Experience?</b>\n\n"
                "💬 <i>Describe it naturally, e.g.:\n"
                "'Cashier at Toko Makmur 2021-2023, then Sales Admin at PT ABC 2023-2024'\n"
                "I will format it into professional statements!</i>\n\n"
                "<i>(Type '-' or 'Fresh Grad' if none)</i>"
            ),
            4: (
                "🏆 <b>Any projects, organizations, competitions, or key accomplishments?</b>\n<i>(Be honest, I will refine the wording for you!)</i>"
                if is_fresh else
                "🏆 <b>What were your key responsibilities or accomplishments in that role?</b>\n<i>(Be honest, I will refine the wording for you!)</i>"
            ),
            5: "🎓 <b>Your Latest Education?</b>\n<i>(e.g., Bachelor of Management - Universitas Terbuka, 2023)</i>",
            6: "🛠️ <b>Your Top Skills / Expertise?</b>\n<i>(e.g., Ms. Excel, Customer Service, Canvassing, Python)</i>",
            7: "📱 <b>WhatsApp / Phone Number (Optional)</b>\n<i>Recruiters need contact info to reach you if you pass selection. You can enter your number here, or click 'Skip' to add it manually in Word later.</i>",
            8: "📍 <b>Current City of Residence?</b>\n<i>Optional if you prefer not to display your location on your CV right now.</i>",
            9: "🔗 <b>LinkedIn / Portfolio / GitHub Link?</b>\n<i>Optional. If you don't have one yet, feel free to skip!</i>"
        }
    else:
        questions = {
            1: "👤 <b>Siapa nama lengkapmu?</b>\n<i>(Nama resmi yang ingin dicantumkan di paling atas CV)</i>",
            2: "📧 <b>Email aktif yang bisa dihubungi recruiter?</b>\n<i>(contoh: nama@gmail.com)</i>",
            3: (
                "💼 <b>Pengalaman Kerja / Organisasi / Freelance terakhirmu?</b>\n\n"
                "💬 <i>Ceritakan santai saja seperti ke teman, contoh:\n"
                "'Kasir di Toko Makmur 2021-2023, lalu Admin Sales di PT ABC 2023-2024'\n"
                "Nanti saya bantu susun menjadi kalimat profesional!</i>\n\n"
                "<i>(Ketik '-' atau 'Fresh Grad' jika belum ada pengalaman)</i>"
            ),
            4: (
                "🏆 <b>Ada project, organisasi, lomba, magang, atau pencapaian yang pernah kamu lakukan?</b>\n"
                "<i>Ceritakan santai saja. Misalnya: 'Pernah bikin website untuk target kuliah' atau 'Aktif panitia kampus'.\n"
                "Kalau belum ada juga tidak masalah, tinggal tekan tombol Lewati 😊</i>"
                if is_fresh else
                "🏆 <b>Apa saja tugas utama atau pencapaianmu di pekerjaan tersebut?</b>\n"
                "<i>(Tulis apa adanya, tidak perlu dibuat-buat. Nanti saya rapikan!)</i>"
            ),
            5: "🎓 <b>Pendidikan terakhirmu?</b>\n<i>(contoh: S1 Manajemen - Universitas Terbuka, 2023)</i>",
            6: "🛠️ <b>Skill / Keahlian utama kamu?</b>\n<i>(contoh: Ms. Excel, Customer Service, Canvassing, Python)</i>",
            7: "📱 <b>Nomor WhatsApp / HP (Opsional)</b>\n<i>Rekruter butuh nomor kontak untuk menghubungi kamu jika lolos seleksi. Kamu bisa masukkan nomor HP di sini, atau tekan tombol [ ⏩ Lewati ] jika ingin mengisinya sendiri nanti di Word.</i>",
            8: "📍 <b>Kota Domisili saat ini?</b>\n<i>Tidak wajib diisi kalau kamu belum ingin mencantumkan lokasi di CV.</i>",
            9: "🔗 <b>Link LinkedIn / Portfolio / GitHub?</b>\n<i>Kalau belum punya, tidak masalah. Kamu bisa menyusul menambahkannya nanti.</i>"
        }
    return questions.get(step, "")

def format_telegram_review_response(data: dict, target_position: str) -> dict:
    scores = data.get("scores", {})
    confidence = data.get("confidence", {})
    
    msg = f"📊 <b>CV REVIEW DIAGNOSIS</b>\n"
    msg += f"🎯 Target: <b>{target_position}</b>\n\n"
    msg += f"📄 CV Quality        : <b>{scores.get('cv_quality', 0)}/100</b>\n"
    msg += f"🎯 Job Match         : <b>{scores.get('job_match', 0)}/100</b>\n"
    msg += f"💪 Evidence Strength : <b>{scores.get('evidence_strength', 0)}/100</b>\n"
    msg += f"───────────────\n"
    msg += f"📈 <b>Overall Score   : {data.get('overall_score', 0)}/100</b>\n\n"
  
    if data.get("strengths"):
        msg += "<b>💪 Kekuatan Utama:</b>\n"
        for s in data["strengths"]:
            msg += f"• {s}\n"
        msg += "\n"
        
    if data.get("weaknesses"):
        msg += "<b>⚠️ Celah Perbaikan:</b>\n"
        for w in data["weaknesses"]:
            msg += f"• {w}\n"
        msg += "\n"
        
    if data.get("action_plan"):
        msg += "<b>🎯 Prioritas Action Plan:</b>\n"
        for act in data["action_plan"][:3]:
            icon = "🔴" if act.get("priority") == "HIGH" else ("🟡" if act.get("priority") == "MEDIUM" else "🟢")
            msg += f"{icon} <b>{act.get('section')}</b>: {act.get('recommendation')}\n"
        msg += "\n"

    msg += f"🔍 <i>Confidence: {confidence.get('level', 'MEDIUM')} ({confidence.get('reason', '')})</i>\n"
    
    response = {
        "text": msg,
        "parse_mode": "HTML"
    }

    if data.get("is_locked"):
        msg += f"\n🔒 <i>{data.get('upgrade_cta')}</i>"
        response["text"] = msg
        response["reply_markup"] = {
            "inline_keyboard": [
                [{"text": "🚀 Buat Career Page Saya (Rp10.000)", "callback_data": "cp_build_now"}],
                [{"text": "🏠 Kembali ke Menu Utama", "callback_data": "home_back_main"}]
            ]
        }
        
    return response

async def handle_cv_review_process(user_id: int, target_position: str, cv_text: str, is_paid: bool = False) -> dict:
    det_result = cv_review_engine.evaluate_cv(cv_text, target_position)
    prompt = cv_review_engine.build_llm_prompt(det_result, cv_text, target_position, is_paid)
    
    try:
        llm_raw_response = await ai_gateway.generate(
            user_message=prompt,
            context={"user_id": user_id, "feature": "cv_review"}
        )
        if llm_raw_response:
            llm_json = json.loads(llm_raw_response)
            det_result.update(llm_json)
    except Exception as e:
        print(f"[CV Review Engine] LLM Error / Timeout: {e}", flush=True)

    final_output = cv_review_engine.apply_access_control(det_result, is_paid)
    
    await cv_review_service.save_review(
        user_id=user_id,
        target_position=target_position,
        overall_score=final_output.get("overall_score", 0),
        quality_score=det_result["scores"]["cv_quality"],
        job_match_score=det_result["scores"]["job_match"],
        evidence_score=det_result["scores"]["evidence_strength"],
        review_json=final_output,
        confidence_level=det_result["confidence"]["level"]
    )
    
    return format_telegram_review_response(final_output, target_position)

async def ai_career_chat_response(user_query: str, user_context: Optional[dict] = None) -> str:
    user_context = user_context or {}
    try:
        response = await _brain_engine.handle_message(
            user_message=user_query,
            context=user_context
        )
        if response:
            if isinstance(response, dict):
                return response.get("text", "")
            return str(response)
    except Exception as e:
        print(f"[BRAIN ENGINE ERROR]: {type(e).__name__}: {e}", flush=True)

    try:
        response = await ai_gateway.generate(
            user_message=user_query,
            context=user_context
        )
        if response:
            return response
    except Exception as e:
        print(f"[AI GATEWAY DIRECT ERROR]: {type(e).__name__}: {e}", flush=True)

    return "Maaf, staf kami yang menjawab untuk kebutuhan karir sedang tidak di tempat. Mungkin bisa coba lagi nanti ya 🙏"

def create_cv_docx(user_id: int, data: dict) -> str:
    doc = Document()
    target_lang = data.get("target_lang", "ID")
    status_kerja = data.get("status_kerja", "Berpengalaman")
    is_en = target_lang in ["EN", "HYBRID"]
    
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    name = clean_val(data.get("1", "NAMA LENGKAP"))
    email = clean_val(data.get("2", ""))
    phone = clean_val(data.get("7", ""))
    domicile = clean_val(data.get("8", ""))
    linkedin = clean_val(data.get("9", ""))

    p_name = doc.add_paragraph()
    p_name.paragraph_format.space_after = Pt(2)
    r_name = p_name.add_run(name.upper())
    r_name.font.name = 'Calibri'
    r_name.font.size = Pt(16)
    r_name.font.bold = True
    r_name.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    contact_parts = [p for p in [email, phone, domicile, linkedin] if p]
    if contact_parts:
        p_contact = doc.add_paragraph()
        p_contact.paragraph_format.space_after = Pt(12)
        r_contact = p_contact.add_run(" | ".join(contact_parts))
        r_contact.font.name = 'Calibri'
        r_contact.font.size = Pt(10)
        r_contact.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    def add_section_header(title):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(10)
        h.paragraph_format.space_after = Pt(4)
        run = h.add_run(title.upper())
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
        
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '1F4E78')
        pBdr.append(bottom)
        h._p.get_or_add_pPr().append(pBdr)

    add_section_header("PROFESSIONAL SUMMARY" if is_en else "RINGKASAN PROFESIONAL")
    position_text = clean_val(data.get("target_position", "Profesional"))
    summary_text = f"Profesional yang berdedikasi dan berorientasi pada hasil dengan fokus pada bidang {position_text}. Memiliki kemampuan komunikasi yang baik serta siap memberikan kontribusi positif."
    p_sum = doc.add_paragraph(summary_text)
    p_sum.paragraph_format.space_after = Pt(8)
    for r in p_sum.runs:
        r.font.name = 'Calibri'
        r.font.size = Pt(10.5)

    exp = clean_val(data.get("3", ""))
    ach_raw = clean_val(data.get("4", ""))

    if exp:
        section_title = "ORGANIZATION & PROJECTS" if (is_en and "fresh" in str(status_kerja).lower()) else ("PROFESSIONAL EXPERIENCE" if is_en else "PENGALAMAN KERJA / ORGANISASI")
        add_section_header(section_title)
        raw_jobs = [j.strip() for j in re.split(r'[\n|]', exp) if j.strip()]
        
        for job_title in raw_jobs:
            if not job_title:
                continue
            
            p_job = doc.add_paragraph()
            p_job.paragraph_format.space_before = Pt(6)
            p_job.paragraph_format.space_after = Pt(2)
            r_job = p_job.add_run(job_title)
            r_job.font.name = 'Calibri'
            r_job.font.size = Pt(10.5)
            r_job.font.bold = True

            if ach_raw:
                for bullet in ach_raw.split("\n"):
                    b_text = bullet.strip().lstrip("-*• ")
                    if b_text:
                        p_b = doc.add_paragraph(style='List Bullet')
                        p_b.paragraph_format.space_after = Pt(2)
                        r_b = p_b.add_run(b_text)
                        r_b.font.name = 'Calibri'
                        r_b.font.size = Pt(10)

    edu = clean_val(data.get("5", ""))
    if edu:
        add_section_header("EDUCATION" if is_en else "PENDIDIKAN")
        p_edu = doc.add_paragraph(edu)
        p_edu.paragraph_format.space_after = Pt(8)
        for r in p_edu.runs:
            r.font.name = 'Calibri'
            r.font.size = Pt(10.5)

    skill = clean_val(data.get("6", ""))
    if skill:
        add_section_header("SKILLS & EXPERTISE" if is_en else "KEAHLIAN")
        for line in skill.split("\n"):
            line_str = line.strip()
            if line_str:
                p_skill = doc.add_paragraph(line_str)
                p_skill.paragraph_format.space_after = Pt(3)
                for r in p_skill.runs:
                    r.font.name = 'Calibri'
                    r.font.size = Pt(10)

    temp_dir = tempfile.gettempdir()
    file_path = os.path.join(temp_dir, f"CV_{user_id}.docx")
    doc.save(file_path)
    return file_path

def get_career_home_keyboard() -> types.InlineKeyboardMarkup:
    kbd = types.InlineKeyboardMarkup(row_width=1)
    kbd.add(
        types.InlineKeyboardButton("📝 Buat / Edit CV Baru", callback_data="home_create_cv"),
        types.InlineKeyboardButton("🔍 Review CV Saya", callback_data="trigger_cv_review"),
        types.InlineKeyboardButton("🌐 Buat Career Page Profesional (Rp10.000)", callback_data="don_10000"),
        types.InlineKeyboardButton("📚 Ebook & Program Digital", callback_data="home_digital_products"),
        types.InlineKeyboardButton("🎁 Cek Referral Saya", callback_data="home_check_ref"),
        types.InlineKeyboardButton("💼 Tanya Seputar Dunia Kerja", callback_data="home_career_qa")
    )
    return kbd

def get_donation_options_keyboard() -> InlineKeyboardMarkup:
    kbd = InlineKeyboardMarkup(row_width=1)
    kbd.add(
        InlineKeyboardButton("🌐 Buat Career Page Profesional (Rp10.000)", callback_data="don_10000"),
        InlineKeyboardButton("📣 Gratis via Invite 5 Teman (Referral)", callback_data="home_check_ref"),
        InlineKeyboardButton("⏩ Nanti Dulu / Cukup CV Word", callback_data="home_back_main")
    )
    return kbd

async def process_and_send_cv(message: types.Message, user_id: int, user_data: dict, user_state: Optional[dict] = None):
    if user_state is not None and user_id in user_state:
        user_state[user_id]["step"] = 0
    await save_dropoff(user_id, TOTAL_STEPS, user_data)
    
    processing_msg = await message.reply(
        "⏳ <b>Merapikan & menyusun CV kamu agar mudah dibaca sistem rekrutmen...</b>\n"
        "Mohon tunggu sekitar 15-20 detik ya!",
        parse_mode="HTML"
    )

    try:
        file_path = await asyncio.to_thread(create_cv_docx, user_id, user_data)
        position = clean_val(user_data.get("target_position", "General"))
        await save_cv_version(user_id, position, user_data)
        asyncio.create_task(track_event(user_id, "resume_generated", meta={"position": position}))

        document = InputFile(file_path)
        user_name = user_data.get("nama_panggilan", message.from_user.first_name or "Teman")

        await bot.send_document(
            chat_id=user_id,
            document=document,
            caption=f"🎉 <b>CV kamu sudah selesai, {user_name}!</b>\n\n"
                    "File dalam format Word (.docx) sudah saya kirim di atas. Bisa kamu edit kapan saja!",
            parse_mode="HTML"
        )
        
        try:
            await bot.delete_message(chat_id=user_id, message_id=processing_msg.message_id)
        except Exception:
            pass

        # CV Review Engine
        cv_text_summary = f"{user_data.get('3', '')} {user_data.get('4', '')} {user_data.get('6', '')}"
        is_paid = await check_user_paid(user_id)
        review_response = await handle_cv_review_process(user_id, position, cv_text_summary, is_paid)

        if isinstance(review_response, dict):
            await send_chunked_message(
                chat_id=user_id,
                text=review_response.get("text", ""),
                reply_markup=review_response.get("reply_markup"),
                parse_mode=review_response.get("parse_mode", "HTML")
            )
        else:
            await send_chunked_message(user_id, review_response, parse_mode="HTML")

        value_text = (
            "💡 <b>Tips Penting Sebelum Melamar:</b>\n\n"
            "1. <b>Subjek Email Jelas:</b> Gunakan format <code>[Posisi] - [Nama Kamu]</code> (Contoh: <i>Admin Operasional - Rayi Gemilang</i>)\n"
            "2. <b>Body Email Terisi:</b> Jangan biarkan pesan email kosong; sertakan Surat Lamaran/Cover Letter singkat.\n"
            "3. <b>Pencapaian Terukur:</b> Cantumkan angka atau pencapaian konkret saat wawancara nanti.\n\n"
            "CV ini sudah bisa kamu edit kapan saja di Word jika ada bagian yang ingin kamu sesuaikan kembali. 🚀"
        )
        await send_chunked_message(user_id, value_text, parse_mode="HTML")

        slug = get_user_slug(user_data, message.from_user.first_name)

        insight_text = (
            f"📊 <b>Career Insight untuk Posisi {position}:</b>\n\n"
            f"Berdasarkan data profilmu, kekuatan utamamu ada pada keahlian operasional & komunikasi. "
            f"Rekruter di bidang ini akan sangat menyukai portofolio interaktif yang bisa diakses langsung via link bio/LinkedIn.\n\n"
            f"Tampilkan CV, pengalaman, skill & portofolio kamu dalam satu halaman web profesional yang siap dibagikan ke rekruter."
        )

        if is_paid:
            kbd_paid = InlineKeyboardMarkup(row_width=1)
            kbd_paid.add(
                InlineKeyboardButton("🌐 Kelola Career Page Saya", callback_data="cp_manage"),
                InlineKeyboardButton("🔙 Kembali ke Menu Utama", callback_data="home_back_main")
            )
            monetize_text = (
                f"{insight_text}\n\n"
                f"👉 <i>Link Website Live Kamu:</i> https://{slug}.boontrack.com\n"
                f"Kamu bisa memperbarui foto, posisi, atau mengimpor data CV terbaru kapan saja!"
            )
            await send_chunked_message(user_id, monetize_text, reply_markup=kbd_paid, parse_mode="HTML")
        else:
            monetize_text = (
                f"{insight_text}\n\n"
                f"🌐 <b>Buat Career Page Profesional</b>\n"
                f"Contoh Live: <code>rayigemilang.boontrack.com</code>\n"
                f"<i>(Sekali aktivasi seumur hidup — Rp10.000)</i>"
            )
            await send_chunked_message(user_id, monetize_text, reply_markup=get_donation_options_keyboard(), parse_mode="HTML")

        if os.path.exists(file_path):
            os.remove(file_path)

        referrer_id = user_state.get(user_id, {}).get("meta", {}).get("referrer_id")
        if referrer_id:
            ref_count_referrer = await count_referrals(referrer_id)
            if ref_count_referrer >= REQUIRED_REFERRALS:
                reward_text = (
                    f"🎉 <b>SELAMAT! Target {REQUIRED_REFERRALS} Referral Kamu Tercapai!</b>\n\n"
                    f"{REQUIRED_REFERRALS} teman yang kamu ajak telah berhasil menyusun CV.\n"
                    "Kamu berhak klaim <b>Website Portfolio Personal Gratis</b>!\n\n"
                    "Ketik /claim_website untuk klaim websitemu! 🌐"
                )
                try:
                    await bot.send_message(chat_id=int(referrer_id), text=reward_text, parse_mode="HTML")
                except Exception as e:
                    print(f"Error send referral reward: {e}", flush=True)

    except Exception as e:
        print(f"Error Generate CV Flow: {e}", flush=True)
        await message.reply("❌ Terjadi kendala teknis. Silakan tekan /start untuk coba lagi!", parse_mode="HTML")
