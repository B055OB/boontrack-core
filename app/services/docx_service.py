import re
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from app.services.ai_service import ai_generate_summary, ai_rewrite_achievement

def clean_val(val) -> str:
    """Membersihkan teks input dari karakter/kata tempat penampung kosong."""
    if not val:
        return ""
    if isinstance(val, list):
        return ", ".join([str(v) for v in val if v])
    v = str(val).strip().lower()
    if v in ["-", "skip", "tidak ada", "ga ada", "ngga ada", "belum ada", "hangus", "hilang", "lupa", "kosong"]:
        return ""
    return str(val).strip()

def create_cv_docx(user_id: int, data: dict) -> str:
    """Membuat dokumen CV format .docx dan mengembalikan path file temporary-nya."""
    doc = Document()
    
    # Atur Margin Halaman (0.75 inci)
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    name = clean_val(data.get("step_1", "NAMA LENGKAP"))
    email = clean_val(data.get("step_2", ""))
    phone = clean_val(data.get("step_3", ""))
    domicile = clean_val(data.get("step_4", ""))
    linkedin = clean_val(data.get("step_5", ""))

    # Header Nama Lengkap
    p_name = doc.add_paragraph()
    p_name.paragraph_format.space_after = Pt(2)
    r_name = p_name.add_run(name.upper())
    r_name.font.name = 'Calibri'
    r_name.font.size = Pt(16)
    r_name.font.bold = True
    r_name.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    # Info Kontak
    contact_parts = [p for p in [email, phone, domicile, linkedin] if p]
    if contact_parts:
        p_contact = doc.add_paragraph()
        p_contact.paragraph_format.space_after = Pt(12)
        r_contact = p_contact.add_run(" | ".join(contact_parts))
        r_contact.font.name = 'Calibri'
        r_contact.font.size = Pt(10)
        r_contact.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    def add_section_header(title: str):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(10)
        h.paragraph_format.space_after = Pt(4)
        run = h.add_run(title.upper())
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
        
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '1F4E78')
        pBdr.append(bottom)
        h._p.get_or_add_pPr().append(pBdr)

    # Section 1: Professional Summary
    summary_text = clean_val(data.get("step_6", ""))
    if summary_text:
        add_section_header("PROFESSIONAL SUMMARY")
        p_sum = doc.add_paragraph(summary_text)
        p_sum.paragraph_format.space_after = Pt(8)
        for r in p_sum.runs:
            r.font.name = 'Calibri'
            r.font.size = Pt(10.5)

    # Section 2: Professional Experience
    exp = clean_val(data.get("step_7", ""))
    ach_raw = clean_val(data.get("step_8", ""))

    if exp:
        add_section_header("PROFESSIONAL EXPERIENCE")
        raw_jobs = [j.strip() for j in re.split(r'[\n|]', exp) if j.strip()]
        
        for job_title in raw_jobs:
            if not job_title:
                continue
            
            p_job = doc.add_paragraph()
            p_job.paragraph_format.space_before = Pt(6)
            p_job.paragraph_format.space_after = Pt(2)
            r_job = p_job.add_run(job_title)
            r_job.font.name = 'Calibri'
            r_job.font.size = Pt(10.5)
            r_job.font.bold = True

            if ach_raw:
                for bullet in ach_raw.split("\n"):
                    b_text = bullet.strip().lstrip("-*• ")
                    if b_text:
                        p_b = doc.add_paragraph(style='List Bullet')
                        p_b.paragraph_format.space_after = Pt(2)
                        r_b = p_b.add_run(b_text)
                        r_b.font.name = 'Calibri'
                        r_b.font.size = Pt(10)

    # Section 3: Education
    edu = clean_val(data.get("step_9", ""))
    if edu:
        add_section_header("EDUCATION")
        p_edu = doc.add_paragraph(edu)
        p_edu.paragraph_format.space_after = Pt(8)
        for r in p_edu.runs:
            r.font.name = 'Calibri'
            r.font.size = Pt(10.5)

    # Section 4: Skills
    skill = clean_val(data.get("step_10", ""))
    if skill:
        add_section_header("SKILLS")
        for line in skill.split("\n"):
            line_str = line.strip()
            if line_str:
                p_skill = doc.add_paragraph(line_str)
                p_skill.paragraph_format.space_after = Pt(3)
                for r in p_skill.runs:
                    r.font.name = 'Calibri'
                    r.font.size = Pt(10)

    file_path = f"/tmp/CV_{user_id}.docx"
    doc.save(file_path)
    return file_path

async def generate_cv_docx(user_id, data: dict) -> str:
    """Wrapper async terpadu untuk merender dokumen Word CV."""
    clean_id = "".join(filter(str.isdigit, str(user_id))) or "1"
    
    summary_val = data.get("summary") or data.get(10) or data.get("step_6", "")
    if not clean_val(summary_val):
        pos_title = data.get("position") or data.get(5) or "General Professional"
        summary_val = await ai_generate_summary(str(pos_title))

    ach_val = data.get("achievements") or data.get("step_8", "")
    exp_val = data.get("experience") or data.get(6) or data.get("step_7", "")
    if not clean_val(ach_val) and clean_val(exp_val):
        ach_val = await ai_rewrite_achievement(str(exp_val))

    normalized_data = {
        "step_1": data.get("name") or data.get(1) or data.get("step_1", "NAMA LENGKAP"),
        "step_2": data.get("email") or data.get(3) or data.get("step_2", ""),
        "step_3": data.get("phone") or data.get(2) or data.get("step_3", ""),
        "step_4": data.get("city") or data.get(4) or data.get("step_4", ""),
        "step_5": data.get("linkedin") or data.get(9) or data.get("step_5", ""),
        "step_6": summary_val,
        "step_7": exp_val,
        "step_8": ach_val,
        "step_9": data.get("education") or data.get(7) or data.get("step_9", ""),
        "step_10": data.get("skills") or data.get(8) or data.get("step_10", "")
    }
    
    return create_cv_docx(int(clean_id) if clean_id.isdigit() else 1, normalized_data)