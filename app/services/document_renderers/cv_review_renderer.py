"""app/services/document_renderers/cv_review_renderer.py
Template renderer for HR CV Review & ATS Audit Report (Laporan_Review_CV_HR.docx).
"""

import io
from typing import Dict, Any
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.services.document_renderers.common import (
    COLOR_PRIMARY,
    COLOR_SECONDARY,
    COLOR_DARK,
    COLOR_MUTED,
    COLOR_SUCCESS,
    COLOR_WARNING,
    COLOR_DANGER,
    set_document_margins,
    add_section_header,
    add_bullet_item,
    add_compliance_footer
)


def render_cv_review_docx(data: Dict[str, Any]) -> bytes:
    """Merender Laporan Audit CV HR & Evaluasi ATS ke format Word (.docx)."""
    doc = Document()
    set_document_margins(doc, 0.75)

    score = data.get("ats_score") or data.get("overall_score") or 0
    target_role = data.get("target_role") or "Kandidat Profesional"
    summary = data.get("summary") or "Laporan evaluasi komprehensif profil pelamar."

    # 1. Header Laporan
    p_badge = doc.add_paragraph()
    p_badge.paragraph_format.space_before = Pt(0)
    p_badge.paragraph_format.space_after = Pt(2)
    p_badge.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_badge = p_badge.add_run("★ BOONTRACK ATS AUDIT & HR REVIEW REPORT ★")
    r_badge.font.name = 'Calibri'
    r_badge.font.size = Pt(9.5)
    r_badge.font.bold = True
    r_badge.font.color.rgb = COLOR_SECONDARY

    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_after = Pt(2)
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run("LAPORAN EVALUASI & REVIEW CV HR")
    r_title.font.name = 'Calibri'
    r_title.font.size = Pt(16)
    r_title.font.bold = True
    r_title.font.color.rgb = COLOR_PRIMARY

    p_role = doc.add_paragraph()
    p_role.paragraph_format.space_after = Pt(8)
    p_role.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_role = p_role.add_run(f"Target Posisi: {target_role}")
    r_role.font.name = 'Calibri'
    r_role.font.size = Pt(11)
    r_role.font.bold = True
    r_role.font.color.rgb = COLOR_MUTED

    # 2. Kotak Skor ATS Visual
    p_sc = doc.add_paragraph()
    p_sc.paragraph_format.space_before = Pt(4)
    p_sc.paragraph_format.space_after = Pt(4)
    p_sc.alignment = WD_ALIGN_PARAGRAPH.CENTER

    score_color = COLOR_SUCCESS if score >= 80 else (COLOR_WARNING if score >= 60 else COLOR_DANGER)
    r_sc = p_sc.add_run(f"SKOR ATS KESELURUHAN: {score}/100")
    r_sc.font.name = 'Calibri'
    r_sc.font.size = Pt(14)
    r_sc.font.bold = True
    r_sc.font.color.rgb = score_color

    # 3. Executive Summary
    add_section_header(doc, "Ringkasan Eksekutif Hasil Audit")
    p_sum = doc.add_paragraph(str(summary))
    p_sum.paragraph_format.space_before = Pt(2)
    p_sum.paragraph_format.space_after = Pt(6)
    p_sum.paragraph_format.line_spacing = 1.15
    for r in p_sum.runs:
        r.font.name = 'Calibri'
        r.font.size = Pt(10)
        r.font.color.rgb = COLOR_DARK

    # 4. Breakdown Parameter Skor
    breakdown = data.get("breakdown_scores") or {}
    if isinstance(breakdown, dict) and breakdown:
        add_section_header(doc, "Skor Parameter ATS & Kompatibilitas Sistem")
        t = doc.add_table(rows=1, cols=2)
        t.style = 'Table Grid'
        hdr_cells = t.rows[0].cells
        hdr_cells[0].text = "Parameter Evaluasi"
        hdr_cells[1].text = "Skor Indikator"
        for cell in hdr_cells:
            for p in cell.paragraphs:
                p.runs[0].font.bold = True
                p.runs[0].font.size = Pt(9.5)
                p.runs[0].font.color.rgb = COLOR_PRIMARY

        for k, v in breakdown.items():
            row_cells = t.add_row().cells
            row_cells[0].text = k.replace("_", " ").title()
            row_cells[1].text = f"{v}/100"
            for cell in row_cells:
                for p in cell.paragraphs:
                    p.runs[0].font.size = Pt(9.5)

    # 5. Kelebihan & Kekuatan CV (Strengths)
    strengths = data.get("strengths") or []
    if isinstance(strengths, list) and strengths:
        add_section_header(doc, "Kelebihan & Kekuatan Profil (Strengths)", color_rgb=COLOR_SUCCESS)
        for item in strengths:
            add_bullet_item(doc, str(item), bold_prefix="✓ ")

    # 6. Catatan Kritis & Red Flags HR
    red_flags = data.get("red_flags") or []
    if isinstance(red_flags, list) and red_flags:
        add_section_header(doc, "Catatan Kritis & Red Flags HR", color_rgb=COLOR_DANGER)
        for rf in red_flags:
            add_bullet_item(doc, str(rf), bold_prefix="⚠️ ")

    # 7. Missing Keywords (ATS Gaps)
    missing_kw = data.get("missing_keywords") or []
    if isinstance(missing_kw, list) and missing_kw:
        add_section_header(doc, "Kata Kunci Penting yang Belum Ditemukan (Missing Keywords)", color_rgb=COLOR_SECONDARY)
        kw_str = ", ".join([str(k) for k in missing_kw])
        p_kw = doc.add_paragraph(f"Kata Kunci Relevan Disarankan: {kw_str}")
        p_kw.paragraph_format.space_before = Pt(2)
        p_kw.paragraph_format.space_after = Pt(6)
        for r in p_kw.runs:
            r.font.name = 'Calibri'
            r.font.size = Pt(10)
            r.font.color.rgb = COLOR_DARK

    # 8. Actionable Fixes (Rekomendasi Konkret per Bagian)
    fixes = data.get("actionable_fixes") or data.get("findings") or data.get("recommendations") or []
    if isinstance(fixes, list) and fixes:
        add_section_header(doc, "Rekomendasi Perbaikan Konkret (Actionable Fixes)")
        for item in fixes:
            if isinstance(item, dict):
                section = item.get("section", "Umum")
                issue = item.get("issue") or item.get("finding") or ""
                rec = item.get("fix") or item.get("recommendation") or ""
                full_text = f"{issue} -> Langkah Perbaikan: {rec}" if rec else issue
                add_bullet_item(doc, full_text, bold_prefix=f"[{section}] ")
            else:
                add_bullet_item(doc, str(item))

    # 9. Prioritas Perbaikan (Priority Improvements)
    priority_imps = data.get("priority_improvements") or []
    if isinstance(priority_imps, list) and priority_imps:
        add_section_header(doc, "Langkah Perbaikan Prioritas Utama (Quick Wins)", color_rgb=COLOR_PRIMARY)
        for idx, imp in enumerate(priority_imps):
            add_bullet_item(doc, str(imp), bold_prefix=f"{idx + 1}. ")

    add_compliance_footer(doc)

    out_io = io.BytesIO()
    doc.save(out_io)
    return out_io.getvalue()
