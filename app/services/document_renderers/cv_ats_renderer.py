"""app/services/document_renderers/cv_ats_renderer.py
Template renderer for single-column ATS CV Optimization (CV_ATS_Optimasi.docx).
"""

import io
from typing import Dict, Any
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.services.document_renderers.common import (
    COLOR_PRIMARY,
    COLOR_SECONDARY,
    COLOR_DARK,
    COLOR_MUTED,
    set_document_margins,
    add_section_header,
    add_bullet_item,
    add_compliance_footer
)


def render_cv_ats_docx(data: Dict[str, Any]) -> bytes:
    """Merender hasil perombakan CV ATS ke file Word (.docx) standar HR internasional."""
    doc = Document()
    set_document_margins(doc, 0.75)

    full_name = data.get("full_name") or data.get("name") or "KANDIDAT PROFESIONAL"
    email = data.get("email") or ""
    phone = data.get("phone") or ""
    location = data.get("location") or data.get("domicile") or ""
    linkedin = data.get("linkedin") or ""
    portfolio = data.get("portfolio") or data.get("website") or ""
    target_pos = data.get("target_position") or data.get("title") or ""

    # 1. Header Nama Kandidat
    p_name = doc.add_paragraph()
    p_name.paragraph_format.space_before = Pt(0)
    p_name.paragraph_format.space_after = Pt(2)
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_name = p_name.add_run(str(full_name).upper())
    r_name.font.name = 'Calibri'
    r_name.font.size = Pt(16)
    r_name.font.bold = True
    r_name.font.color.rgb = COLOR_PRIMARY

    # Target Posisi jika ada
    if target_pos:
        p_t = doc.add_paragraph()
        p_t.paragraph_format.space_before = Pt(0)
        p_t.paragraph_format.space_after = Pt(2)
        p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_t = p_t.add_run(str(target_pos).title())
        r_t.font.name = 'Calibri'
        r_t.font.size = Pt(11)
        r_t.font.bold = True
        r_t.font.color.rgb = COLOR_SECONDARY

    # Kontak gabungan (Satu baris dipisah pipa)
    contact_parts = [p for p in [email, phone, location, linkedin, portfolio] if p]
    if contact_parts:
        p_contact = doc.add_paragraph()
        p_contact.paragraph_format.space_before = Pt(0)
        p_contact.paragraph_format.space_after = Pt(8)
        p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_c = p_contact.add_run(" | ".join(contact_parts))
        r_c.font.name = 'Calibri'
        r_c.font.size = Pt(9.5)
        r_c.font.color.rgb = COLOR_MUTED

    # 2. Ringkasan Profesional (Summary)
    summary = data.get("summary") or data.get("professional_summary") or ""
    if summary:
        add_section_header(doc, "Professional Summary")
        p_sum = doc.add_paragraph(str(summary))
        p_sum.paragraph_format.space_before = Pt(2)
        p_sum.paragraph_format.space_after = Pt(6)
        p_sum.paragraph_format.line_spacing = 1.15
        for r in p_sum.runs:
            r.font.name = 'Calibri'
            r.font.size = Pt(10)
            r.font.color.rgb = COLOR_DARK

    # 3. Keahlian (Skills)
    skills = data.get("skills") or []
    if skills:
        add_section_header(doc, "Core Competencies & Skills")
        if isinstance(skills, dict):
            for cat, s_list in skills.items():
                cat_name = str(cat).replace("_", " ").title()
                s_str = ", ".join(s_list) if isinstance(s_list, list) else str(s_list)
                add_bullet_item(doc, f": {s_str}", bold_prefix=f"{cat_name}")
        elif isinstance(skills, list):
            s_str = ", ".join([str(s) for s in skills if s])
            p_sk = doc.add_paragraph(s_str)
            p_sk.paragraph_format.space_before = Pt(2)
            p_sk.paragraph_format.space_after = Pt(6)
            p_sk.paragraph_format.line_spacing = 1.15
            for r in p_sk.runs:
                r.font.name = 'Calibri'
                r.font.size = Pt(10)
                r.font.color.rgb = COLOR_DARK

    # 4. Pengalaman Kerja (Professional Experience)
    experiences = data.get("experience") or data.get("work_experience") or []
    if isinstance(experiences, list) and experiences:
        add_section_header(doc, "Professional Experience")
        for exp in experiences:
            if not isinstance(exp, dict):
                continue
            role = exp.get("role") or exp.get("position") or exp.get("title") or "Posisi"
            company = exp.get("company") or exp.get("organization") or ""
            period = exp.get("period") or exp.get("dates") or ""
            loc = exp.get("location") or ""
            
            p_role = doc.add_paragraph()
            p_role.paragraph_format.space_before = Pt(6)
            p_role.paragraph_format.space_after = Pt(1)
            
            r_role = p_role.add_run(role)
            r_role.font.name = 'Calibri'
            r_role.font.size = Pt(10.5)
            r_role.font.bold = True
            r_role.font.color.rgb = COLOR_PRIMARY
            
            if company:
                r_comp = p_role.add_run(f" — {company}")
                r_comp.font.name = 'Calibri'
                r_comp.font.size = Pt(10)
                r_comp.font.bold = True
                r_comp.font.color.rgb = COLOR_DARK

            meta_parts = [m for m in [period, loc] if m]
            if meta_parts:
                p_meta = doc.add_paragraph(" | ".join(meta_parts))
                p_meta.paragraph_format.space_before = Pt(0)
                p_meta.paragraph_format.space_after = Pt(3)
                for r in p_meta.runs:
                    r.font.name = 'Calibri'
                    r.font.size = Pt(9)
                    r.font.italic = True
                    r.font.color.rgb = COLOR_MUTED

            bullets = exp.get("bullets") or exp.get("achievements") or exp.get("responsibilities") or []
            if isinstance(bullets, list):
                for b in bullets:
                    add_bullet_item(doc, str(b))

    # 5. Pendidikan (Education)
    education = data.get("education") or []
    if isinstance(education, list) and education:
        add_section_header(doc, "Education")
        for edu in education:
            if isinstance(edu, dict):
                degree = edu.get("degree") or ""
                inst = edu.get("institution") or edu.get("university") or edu.get("school") or ""
                year = edu.get("year") or edu.get("period") or ""
                gpa = edu.get("gpa") or ""
                title_str = f"{degree} — {inst}" if inst else degree
                if year:
                    title_str += f" ({year})"
                if gpa:
                    title_str += f" | IPK/GPA: {gpa}"
                add_bullet_item(doc, title_str)
            else:
                add_bullet_item(doc, str(edu))

    # 6. Sertifikasi & Lisensi
    certs = data.get("certifications") or data.get("licenses") or []
    if isinstance(certs, list) and certs:
        add_section_header(doc, "Certifications & Training")
        for c in certs:
            add_bullet_item(doc, str(c))

    add_compliance_footer(doc)

    out_io = io.BytesIO()
    doc.save(out_io)
    return out_io.getvalue()
