"""app/services/document_renderers/polish_rephrase_renderer.py
Template renderer for Academic Polish & Paraphrase (Naskah_Hasil_Parafrase.docx).
Includes robust Markdown-to-Docx Native Table parser with fixed column widths & Artifact Sanitizer.
"""

import io
import re
import logging
from typing import Dict, Any, List, Union
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.services.document_renderers.common import (
    COLOR_PRIMARY,
    COLOR_SECONDARY,
    COLOR_DARK,
    COLOR_MUTED,
    set_document_margins,
    add_section_header,
    add_bullet_item,
    add_compliance_footer
)

logger = logging.getLogger("DOCX_RENDERER.POLISH")


def _sanitize_academic_text(text: str) -> str:
    """Membersihkan artefak chunking/prompt dan merapikan tabel markdown multi-baris."""
    if not text:
        return ""

    # 1. Hapus marker chunk 'Bagian X' / '**Bagian X**'
    cleaned = re.sub(r'(?i)^\s*(\*\*|#+)?\s*Bagian\s+\d+\s*(\*\*)?\s*$', '', text, flags=re.MULTILINE)

    # 2. Satukan baris tabel Markdown yang terpotong enter/newline antar sel
    cleaned = re.sub(r'(\n\s*\|\s*)|(\s*\|\s*\n)', ' | ', cleaned)
    cleaned = re.sub(r'(\|\s*)\n+(\s*\|)', r'\1\n\2', cleaned)

    # 3. Rapikan multiple empty newlines
    cleaned = re.sub(r'\n{3,}', '\n\n', cleaned)
    return cleaned.strip()


def _is_markdown_table_row(line: str) -> bool:
    """Mengecek apakah satu baris string benar-benar baris tabel Markdown (memiliki pipa ganda)."""
    l = line.strip()
    return bool((l.startswith('|') or l.endswith('|')) and l.count('|') >= 2)


def _render_markdown_table_to_docx(doc: Document, table_lines: List[str]):
    """Mengonversi baris Markdown Table menjadi native table Microsoft Word bergaris dengan proporsi kolom stabil."""
    raw_rows = []
    for line in table_lines:
        # Abaikan baris separator kolom seperti |---|---| atau :---:
        if re.match(r'^\s*\|?(\s*:?-+:?\s*\|)+\s*$', line.strip()):
            continue
        cells = [c.strip() for c in line.strip().split('|')]
        # Bersihkan sel kosong di awal/akhir baris pipa
        if line.strip().startswith('|') and cells:
            cells.pop(0)
        if line.strip().endswith('|') and cells:
            cells.pop()
        if cells and any(c for c in cells):
            raw_rows.append(cells)

    if not raw_rows:
        return

    num_cols = max(len(r) for r in raw_rows)
    table = doc.add_table(rows=len(raw_rows), cols=num_cols)
    table.style = 'Table Grid'
    table.autofit = False  # Matikan autofit agar lebar kolom stabil sesuai kustomisasi

    # Atur lebar kolom manual agar proporsional (Kolom No, Kriteria, Jumlah)
    col_widths = [Inches(0.6), Inches(4.5), Inches(1.2)]
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(col_widths):
                cell.width = col_widths[idx]

    for row_idx, row_data in enumerate(raw_rows):
        is_header = (row_idx == 0)
        for col_idx in range(num_cols):
            cell_text = row_data[col_idx] if col_idx < len(row_data) else ""
            cell = table.cell(row_idx, col_idx)
            cell.text = cell_text

            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)

            for r in p.runs:
                r.font.name = 'Calibri'
                r.font.size = Pt(9.5)
                if is_header:
                    r.font.bold = True
                    r.font.color.rgb = COLOR_PRIMARY
                else:
                    r.font.color.rgb = COLOR_DARK

    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(0)
    p_space.paragraph_format.space_after = Pt(6)


def _normalize_data(data: Union[Dict[str, Any], str, None]) -> Dict[str, Any]:
    """Normalisasi struktur input renderer naskah."""
    if isinstance(data, str):
        sanitized = _sanitize_academic_text(data)
        return {
            "title": "Naskah Hasil Parafrase Akademis",
            "sections": [{"heading": "", "content": sanitized}],
            "full_text": sanitized,
            "full_paraphrased_text": sanitized,
        }
    if not isinstance(data, dict):
        return {}

    sections = data.get("sections") or []
    normalized_sections = []
    for sec in sections:
        if isinstance(sec, dict):
            sec["content"] = _sanitize_academic_text(sec.get("content", ""))
            sec["heading"] = _sanitize_academic_text(sec.get("heading", ""))
            normalized_sections.append(sec)
        elif isinstance(sec, str):
            normalized_sections.append({"heading": "", "content": _sanitize_academic_text(sec)})

    data = dict(data)
    data["sections"] = normalized_sections
    if data.get("full_text"):
        data["full_text"] = _sanitize_academic_text(data["full_text"])
    if data.get("full_paraphrased_text"):
        data["full_paraphrased_text"] = _sanitize_academic_text(data["full_paraphrased_text"])
    return data


def render_polish_rephrase_docx(data: Union[Dict[str, Any], str]) -> bytes:
    """Merender hasil parafrase akademik ke file Word (.docx) berstandar karya ilmiah formal."""
    data = _normalize_data(data)

    doc = Document()
    set_document_margins(doc, 1.0)  # Margin 1 inch formal

    title = data.get("title") or "Naskah Hasil Parafrase Akademis"
    tone = data.get("tone") or "Akademik Formal (EYD V)"
    orig_words = data.get("original_word_count") or 0
    final_words = data.get("paraphrased_word_count") or 0
    full_text = data.get("full_text") or data.get("full_paraphrased_text") or ""
    sections = data.get("sections") or []

    # 1. Header Judul
    p_badge = doc.add_paragraph()
    p_badge.paragraph_format.space_before = Pt(0)
    p_badge.paragraph_format.space_after = Pt(2)
    p_badge.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_badge = p_badge.add_run("★ BOONTRACK ACADEMIC EDITING & REPHRASE ★")
    r_badge.font.name = 'Calibri'
    r_badge.font.size = Pt(9.5)
    r_badge.font.bold = True
    r_badge.font.color.rgb = COLOR_SECONDARY

    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_after = Pt(4)
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run(str(title).upper())
    r_title.font.name = 'Calibri'
    r_title.font.size = Pt(15)
    r_title.font.bold = True
    r_title.font.color.rgb = COLOR_PRIMARY

    # Metadata Penyempurnaan Naskah
    p_meta = doc.add_paragraph()
    p_meta.paragraph_format.space_after = Pt(12)
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_text = f"Standar: {tone}"
    if orig_words > 0 and final_words > 0:
        meta_text += f" | Panjang Naskah: {final_words} kata (Input: {orig_words} kata)"
    elif final_words > 0:
        meta_text += f" | Panjang Naskah: {final_words} kata"
    r_meta = p_meta.add_run(meta_text)
    r_meta.font.name = 'Calibri'
    r_meta.font.size = Pt(9.5)
    r_meta.font.color.rgb = COLOR_MUTED

    # 2. Catatan Kualitas Naskah
    takeaways = data.get("key_takeaways") or []
    if isinstance(takeaways, list) and takeaways:
        add_section_header(doc, "Catatan Penyempurnaan Naskah", color_rgb=COLOR_SECONDARY)
        for t in takeaways:
            add_bullet_item(doc, str(t), bold_prefix="✓ ")

    # 3. Isi Naskah Akademis
    add_section_header(doc, "Naskah Hasil Penyempurnaan (EYD V)", color_rgb=COLOR_PRIMARY)

    raw_text_stream = ""
    if sections:
        raw_text_stream = "\n\n".join([sec.get("content", "") for sec in sections if isinstance(sec, dict)])
    if not raw_text_stream and full_text:
        raw_text_stream = full_text

    # Stream Parser: pisahkan blok tabel Markdown dengan paragraf teks biasa
    lines = raw_text_stream.split('\n')
    current_table_lines: List[str] = []

    def flush_table(doc_ref, t_lines):
        if t_lines:
            _render_markdown_table_to_docx(doc_ref, t_lines)
            t_lines.clear()

    for line in lines:
        line_str = line.strip()
        if not line_str:
            flush_table(doc, current_table_lines)
            continue

        if _is_markdown_table_row(line_str):
            current_table_lines.append(line_str)
        else:
            flush_table(doc, current_table_lines)

            # Render paragraf teks biasa
            p_body = doc.add_paragraph(line_str)
            p_body.paragraph_format.space_before = Pt(2)
            p_body.paragraph_format.space_after = Pt(6)
            p_body.paragraph_format.line_spacing = 1.25
            for r in p_body.runs:
                r.font.name = 'Calibri'
                r.font.size = Pt(10.5)
                r.font.color.rgb = COLOR_DARK

    flush_table(doc, current_table_lines)
    add_compliance_footer(doc)

    out_io = io.BytesIO()
    doc.save(out_io)
    return out_io.getvalue()