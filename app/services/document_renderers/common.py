"""app/services/document_renderers/common.py
Common styling primitives and typography tokens for python-docx document renderers.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.services.pricing_engine import COMPLIANCE_DISCLAIMER, OFFICIAL_PRODUCT_NAME

# Color Palette Constants
COLOR_PRIMARY = RGBColor(0x1F, 0x4E, 0x78)    # Deep Navy
COLOR_SECONDARY = RGBColor(0x2E, 0x75, 0xB6)  # Accent Blue
COLOR_DARK = RGBColor(0x22, 0x22, 0x22)       # Text Dark
COLOR_MUTED = RGBColor(0x66, 0x66, 0x66)      # Subtext Gray
COLOR_SUCCESS = RGBColor(0x2E, 0x7D, 0x32)    # Green (High Score)
COLOR_WARNING = RGBColor(0xE6, 0x51, 0x00)    # Orange (Medium Score / Warnings)
COLOR_DANGER = RGBColor(0xC6, 0x28, 0x28)     # Red (Critical Red Flags)


def set_document_margins(doc: Document, margin_inches: float = 0.75):
    """Mengatur margin dokumen seragam."""
    for section in doc.sections:
        section.top_margin = Inches(margin_inches)
        section.bottom_margin = Inches(margin_inches)
        section.left_margin = Inches(margin_inches)
        section.right_margin = Inches(margin_inches)


def add_section_header(doc: Document, title: str, color_rgb: RGBColor = COLOR_PRIMARY):
    """Menambahkan heading section bergaris bawah tipis standar HR & korporat."""
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(4)
    run = h.add_run(title.upper())
    run.font.name = 'Calibri'
    run.font.size = Pt(11.5)
    run.font.bold = True
    run.font.color.rgb = color_rgb

    # Garis pembatas bawah elegan
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '2')
    bottom.set(qn('w:color'), '1F4E78')
    pBdr.append(bottom)
    h._p.get_or_add_pPr().append(pBdr)


def add_bullet_item(doc: Document, text: str, bold_prefix: str = ""):
    """Menambahkan item bullet point dengan format rapi."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    
    if bold_prefix:
        r_bold = p.add_run(bold_prefix)
        r_bold.font.name = 'Calibri'
        r_bold.font.size = Pt(10)
        r_bold.font.bold = True
        r_bold.font.color.rgb = COLOR_DARK

    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.color.rgb = COLOR_DARK


def add_compliance_footer(doc: Document):
    """Menambahkan disclaimer kepatuhan resmi di bagian bawah dokumen."""
    p_f = doc.add_paragraph()
    p_f.paragraph_format.space_before = Pt(16)
    p_f.paragraph_format.space_after = Pt(4)
    
    # Divider line
    pBdr = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '4')
    top.set(qn('w:space'), '4')
    top.set(qn('w:color'), 'CCCCCC')
    pBdr.append(top)
    p_f._p.get_or_add_pPr().append(pBdr)

    r_f = p_f.add_run(f"📌 {COMPLIANCE_DISCLAIMER}")
    r_f.font.name = 'Calibri'
    r_f.font.size = Pt(8.5)
    r_f.font.italic = True
    r_f.font.color.rgb = COLOR_MUTED
