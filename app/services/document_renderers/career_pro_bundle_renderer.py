"""app/services/document_renderers/career_pro_bundle_renderer.py
Template renderer for Career Pro Bundle (Paket_Lengkap_Karir_Pro.docx).
Combines:
1. CV Tailored ATS
2. HR Strategic Roadmap & STAR Interview Guide
3. Targeted Formal Cover Letter
"""

import io
from typing import Dict, Any
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.services.document_renderers.common import (
    COLOR_PRIMARY,
    COLOR_SECONDARY,
    COLOR_DARK,
    COLOR_MUTED,
    COLOR_SUCCESS,
    set_document_margins,
    add_section_header,
    add_bullet_item,
    add_compliance_footer
)


def render_career_pro_bundle_docx(data: Dict[str, Any]) -> bytes:
    """Merender Paket Lengkap Karir Pro (3 Pilar) ke file Word (.docx)."""
    doc = Document()
    set_document_margins(doc, 0.75)

    full_name = data.get("full_name") or data.get("name") or "KANDIDAT PROFESIONAL"
    email = data.get("email") or ""
    phone = data.get("phone") or ""
    location = data.get("location") or data.get("domicile") or ""
    linkedin = data.get("linkedin") or ""
    portfolio = data.get("portfolio") or data.get("website") or ""
    target_pos = data.get("target_position") or data.get("title") or "Target Posisi Karir"

    # ==========================================
    # BAGIAN 1: HEADER BUNDLE & CV ATS TAILORED
    # ==========================================
    p_badge = doc.add_paragraph()
    p_badge.paragraph_format.space_before = Pt(0)
    p_badge.paragraph_format.space_after = Pt(2)
    p_badge.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_badge = p_badge.add_run("★ BOONTRACK CAREER PRO BUNDLE ★")
    r_badge.font.name = 'Calibri'
    r_badge.font.size = Pt(9.5)
    r_badge.font.bold = True
    r_badge.font.color.rgb = COLOR_SECONDARY

    p_name = doc.add_paragraph()
    p_name.paragraph_format.space_after = Pt(2)
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_name = p_name.add_run(str(full_name).upper())
    r_name.font.name = 'Calibri'
    r_name.font.size = Pt(16)
    r_name.font.bold = True
    r_name.font.color.rgb = COLOR_PRIMARY

    p_target = doc.add_paragraph()
    p_target.paragraph_format.space_after = Pt(2)
    p_target.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_t = p_target.add_run(str(target_pos).title())
    r_t.font.name = 'Calibri'
    r_t.font.size = Pt(11)
    r_t.font.bold = True
    r_t.font.color.rgb = COLOR_SECONDARY

    contact_items = [x for x in [email, phone, location, linkedin, portfolio] if x]
    if contact_items:
        p_c = doc.add_paragraph()
        p_c.paragraph_format.space_after = Pt(10)
        p_c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_c = p_c.add_run(" | ".join(contact_items))
        r_c.font.name = 'Calibri'
        r_c.font.size = Pt(9.5)
        r_c.font.color.rgb = COLOR_MUTED

    # 1.1 Summary CV
    summary = data.get("summary") or ""
    if summary:
        add_section_header(doc, "Professional Summary (ATS Tailored)")
        p_s = doc.add_paragraph(str(summary))
        p_s.paragraph_format.space_after = Pt(6)
        p_s.paragraph_format.line_spacing = 1.15
        for r in p_s.runs:
            r.font.name = 'Calibri'
            r.font.size = Pt(10)
            r.font.color.rgb = COLOR_DARK

    # 1.2 Skills
    skills = data.get("skills") or []
    if skills:
        add_section_header(doc, "Core Competencies & Skills")
        if isinstance(skills, dict):
            for cat, s_list in skills.items():
                cat_name = str(cat).replace("_", " ").title()
                s_str = ", ".join(s_list) if isinstance(s_list, list) else str(s_list)
                add_bullet_item(doc, f": {s_str}", bold_prefix=f"{cat_name}")
        elif isinstance(skills, list):
            s_str = ", ".join([str(s) for s in skills if s])
            p_sk = doc.add_paragraph(s_str)
            p_sk.paragraph_format.space_after = Pt(6)
            for r in p_sk.runs:
                r.font.name = 'Calibri'
                r.font.size = Pt(10)
                r.font.color.rgb = COLOR_DARK

    # 1.3 Experience
    experiences = data.get("experience") or []
    if isinstance(experiences, list) and experiences:
        add_section_header(doc, "Professional Experience (Metric-Based ATS)")
        for exp in experiences:
            if not isinstance(exp, dict):
                continue
            role = exp.get("role") or exp.get("position") or "Position"
            company = exp.get("company") or ""
            period = exp.get("period") or ""
            loc = exp.get("location") or ""
            
            p_exp = doc.add_paragraph()
            p_exp.paragraph_format.space_before = Pt(6)
            p_exp.paragraph_format.space_after = Pt(1)
            r_role = p_exp.add_run(f"{role} — {company}" if company else role)
            r_role.font.name = 'Calibri'
            r_role.font.size = Pt(10.5)
            r_role.font.bold = True
            r_role.font.color.rgb = COLOR_PRIMARY
            
            if period or loc:
                p_meta = doc.add_paragraph(" | ".join([x for x in [period, loc] if x]))
                p_meta.paragraph_format.space_after = Pt(3)
                for r in p_meta.runs:
                    r.font.name = 'Calibri'
                    r.font.size = Pt(9.5)
                    r.font.italic = True
                    r.font.color.rgb = COLOR_MUTED

            bullets = exp.get("bullets") or []
            if isinstance(bullets, list):
                for b in bullets:
                    add_bullet_item(doc, str(b))

    # 1.4 Education & Certifications
    education = data.get("education") or []
    if isinstance(education, list) and education:
        add_section_header(doc, "Education")
        for edu in education:
            if isinstance(edu, dict):
                degree = edu.get("degree") or ""
                inst = edu.get("institution") or ""
                year = edu.get("year") or ""
                title_str = f"{degree} — {inst}" if inst else degree
                if year:
                    title_str += f" ({year})"
                add_bullet_item(doc, title_str)
            else:
                add_bullet_item(doc, str(edu))

    certs = data.get("certifications") or []
    if isinstance(certs, list) and certs:
        add_section_header(doc, "Certifications")
        for c in certs:
            add_bullet_item(doc, str(c))

    # ==========================================
    # BAGIAN 2: REKOMENDASI HR PROFESIONAL
    # ==========================================
    hr_recs = data.get("hr_recommendations") or {}
    if isinstance(hr_recs, dict) and hr_recs:
        add_section_header(doc, "Rekomendasi HR & Career Strategy", color_rgb=COLOR_SECONDARY)
        
        readiness = hr_recs.get("profile_readiness") or ""
        if readiness:
            p_r = doc.add_paragraph()
            p_r.paragraph_format.space_before = Pt(4)
            p_r.paragraph_format.space_after = Pt(4)
            r_r = p_r.add_run(f"Kesiapan Profil: {readiness}")
            r_r.font.name = 'Calibri'
            r_r.font.size = Pt(10.5)
            r_r.font.bold = True
            r_r.font.color.rgb = COLOR_SUCCESS

        strengths = hr_recs.get("key_strengths") or []
        if isinstance(strengths, list) and strengths:
            p_st_title = doc.add_paragraph("Kekuatan Utama Profil:")
            p_st_title.paragraph_format.space_before = Pt(4)
            p_st_title.paragraph_format.space_after = Pt(2)
            p_st_title.runs[0].font.bold = True
            for s in strengths:
                add_bullet_item(doc, str(s))

        improvements = hr_recs.get("strategic_improvements") or []
        if isinstance(improvements, list) and improvements:
            p_im_title = doc.add_paragraph("Rekomendasi Peningkatan Strategis:")
            p_im_title.paragraph_format.space_before = Pt(4)
            p_im_title.paragraph_format.space_after = Pt(2)
            p_im_title.runs[0].font.bold = True
            for imp in improvements:
                add_bullet_item(doc, str(imp))

        interview_tips = hr_recs.get("interview_tips") or []
        if isinstance(interview_tips, list) and interview_tips:
            p_it_title = doc.add_paragraph("Panduan Wawancara HR (Metode STAR):")
            p_it_title.paragraph_format.space_before = Pt(4)
            p_it_title.paragraph_format.space_after = Pt(2)
            p_it_title.runs[0].font.bold = True
            for tip in interview_tips:
                add_bullet_item(doc, str(tip))

    # ==========================================
    # BAGIAN 3: SURAT LAMARAN (COVER LETTER)
    # ==========================================
    cl = data.get("cover_letter") or {}
    if isinstance(cl, dict) and cl:
        add_section_header(doc, "Surat Lamaran Kerja (Cover Letter)", color_rgb=COLOR_PRIMARY)
        
        recipient = cl.get("recipient") or "Yth. Tim Rekrutmen & Hiring Manager"
        subject = cl.get("subject") or f"Aplikasi Lamaran Pekerjaan - {target_pos}"
        salutation = cl.get("salutation") or "Dengan hormat,"
        opening = cl.get("opening") or ""
        body_paragraphs = cl.get("body_paragraphs") or []
        closing = cl.get("closing") or ""
        sign_off = cl.get("sign_off") or f"Hormat saya,\n{full_name}"

        p_rec = doc.add_paragraph(str(recipient))
        p_rec.paragraph_format.space_before = Pt(6)
        p_rec.paragraph_format.space_after = Pt(2)
        p_rec.runs[0].font.bold = True

        p_sub = doc.add_paragraph(f"Perihal: {subject}")
        p_sub.paragraph_format.space_after = Pt(6)
        p_sub.runs[0].font.bold = True

        p_sal = doc.add_paragraph(str(salutation))
        p_sal.paragraph_format.space_after = Pt(6)

        if opening:
            p_op = doc.add_paragraph(str(opening))
            p_op.paragraph_format.space_after = Pt(6)
            p_op.paragraph_format.line_spacing = 1.15

        if isinstance(body_paragraphs, list):
            for para in body_paragraphs:
                p_bp = doc.add_paragraph(str(para))
                p_bp.paragraph_format.space_after = Pt(6)
                p_bp.paragraph_format.line_spacing = 1.15
        elif isinstance(body_paragraphs, str) and body_paragraphs:
            p_bp = doc.add_paragraph(str(body_paragraphs))
            p_bp.paragraph_format.space_after = Pt(6)
            p_bp.paragraph_format.line_spacing = 1.15

        if closing:
            p_cl = doc.add_paragraph(str(closing))
            p_cl.paragraph_format.space_after = Pt(8)
            p_cl.paragraph_format.line_spacing = 1.15

        p_so = doc.add_paragraph(str(sign_off))
        p_so.paragraph_format.space_after = Pt(6)

    add_compliance_footer(doc)

    out_io = io.BytesIO()
    doc.save(out_io)
    return out_io.getvalue()
