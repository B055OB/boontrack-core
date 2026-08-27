import os
import io
import re
import logging
from typing import Optional
import httpx
from pypdf import PdfReader
from docx import Document

logger = logging.getLogger(__name__)

# Magic byte signatures
MAGIC_BYTES_PDF = b"%PDF"
MAGIC_BYTES_ZIP = b"PK\x03\x04"  # DOCX is a zipped XML container


def _get_active_wa_token(tenant_id: str = "boontrack-career") -> str:
    """Mengambil WhatsApp Access Token secara dinamis dari config/env."""
    try:
        from app.services.whatsapp_service import get_wa_credentials
        token, _, _ = get_wa_credentials(tenant_id)
        if token:
            return token
    except Exception:
        pass

    return (
        os.getenv("CAREER_ACCESS_TOKEN")
        or os.getenv("WHATSAPP_TOKEN")
        or os.getenv("META_WA_TOKEN")
        or os.getenv("WA_TOKEN")
        or os.getenv("META_WA_ACCESS_TOKEN")
        or os.getenv("META_ACCESS_TOKEN")
        or ""
    ).strip()


async def download_whatsapp_media(media_id: str, tenant_id: str = "boontrack-career") -> bytes:
    """Mengambil URL download dari Meta Graph API dan mengunduh filenya secara aman."""
    if not media_id:
        raise ValueError("media_id tidak boleh kosong")

    token = _get_active_wa_token(tenant_id)
    version = os.getenv("META_GRAPH_VERSION", "v20.0")
    meta_url = f"https://graph.facebook.com/{version}/{media_id}"
    headers = {"Authorization": f"Bearer {token}"} if token else {}

    async with httpx.AsyncClient(timeout=45.0, follow_redirects=True) as client:
        logger.info(f"[DocumentParser] Fetching media metadata for media_id={media_id} from {meta_url}")
        res = await client.get(meta_url, headers=headers)
        if res.status_code != 200:
            logger.error(f"[DocumentParser] Failed to get media metadata: HTTP {res.status_code} - {res.text}")
            res.raise_for_status()

        download_url = res.json().get("url")
        if not download_url:
            raise ValueError(f"Meta Graph API response did not contain download 'url': {res.text}")

        # Unduh binary file (beberapa Meta CDN URL memerlukan Bearer auth, beberapa tidak)
        try:
            file_res = await client.get(download_url, headers=headers)
            file_res.raise_for_status()
            content = file_res.content
        except httpx.HTTPStatusError as auth_err:
            if auth_err.response.status_code in (400, 401, 403):
                # Retry tanpa Authorization header jika CDN URL sudah presigned
                logger.warning(f"[DocumentParser] Download with auth returned {auth_err.response.status_code}, retrying without auth header...")
                file_res = await client.get(download_url)
                file_res.raise_for_status()
                content = file_res.content
            else:
                raise

        if not content:
            raise ValueError(f"Downloaded media is empty (0 bytes) for media_id={media_id}")

        logger.info(f"[DocumentParser] Successfully downloaded {len(content)} bytes for media_id={media_id}")
        return content


def extract_text_from_bytes(file_bytes: bytes, filename: str = "") -> str:
    """Mengekstrak teks mentah dari file PDF atau DOCX dengan deteksi magic bytes & layout fallback."""
    if not file_bytes or len(file_bytes) < 4:
        logger.warning(f"[Document Parser] file_bytes kosong atau terlalu kecil ({len(file_bytes) if file_bytes else 0} bytes)")
        return ""

    filename_lower = str(filename or "").lower().strip()
    is_pdf = file_bytes.startswith(MAGIC_BYTES_PDF) or filename_lower.endswith(".pdf")
    is_docx = file_bytes.startswith(MAGIC_BYTES_ZIP) or filename_lower.endswith(".docx")
    extracted_text = ""

    # 1. Ekstraksi Dokumen PDF
    if is_pdf:
        try:
            reader = PdfReader(io.BytesIO(file_bytes), strict=False)
            if reader.is_encrypted:
                try:
                    reader.decrypt("")
                except Exception as enc_err:
                    logger.warning(f"[Document Parser] PDF is encrypted, empty decrypt attempt: {enc_err}")

            page_texts = []
            for page_idx, page in enumerate(reader.pages):
                txt = page.extract_text() or ""
                if not txt.strip():
                    try:
                        txt = page.extract_text(extraction_mode="layout") or ""
                    except Exception:
                        pass
                if txt.strip():
                    page_texts.append(txt.strip())

            if page_texts:
                extracted_text = "\n\n".join(page_texts)
        except Exception as e:
            logger.error(f"[Document Parser Error] Gagal mengekstrak PDF {filename}: {e}", exc_info=True)

    # 2. Ekstraksi Dokumen DOCX
    elif is_docx:
        try:
            doc = Document(io.BytesIO(file_bytes))
            para_texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

            # Tambahkan teks dari tabel dokumen (sering digunakan di naskah/CV)
            table_texts = []
            for table in doc.tables:
                for row in table.rows:
                    row_parts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_parts:
                        table_texts.append(" | ".join(row_parts))

            all_parts = para_texts + table_texts
            if all_parts:
                extracted_text = "\n\n".join(all_parts)
        except Exception as e:
            logger.error(f"[Document Parser Error] Gagal mengekstrak DOCX {filename}: {e}", exc_info=True)

    # 3. Fallback Teks Biasa (UTF-8 / Latin-1) jika bukan binary PDF/DOCX
    if not extracted_text.strip():
        if not file_bytes.startswith(b"\x00") and not file_bytes.startswith(MAGIC_BYTES_PDF) and not file_bytes.startswith(MAGIC_BYTES_ZIP):
            for enc in ("utf-8", "latin-1"):
                try:
                    decoded = file_bytes.decode(enc).strip()
                    if len(decoded) > 20 and not re.search(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", decoded):
                        extracted_text = decoded
                        break
                except Exception:
                    pass

    clean_result = extracted_text.strip()
    if not clean_result:
        logger.warning(
            f"[Document Parser] Teks kosong hasil ekstraksi dari {filename} "
            f"(bytes_len={len(file_bytes)}, is_pdf={is_pdf}, is_docx={is_docx})"
        )
    else:
        logger.info(
            f"[Document Parser] Berhasil mengekstrak {len(clean_result.split())} kata "
            f"({len(clean_result)} karakter) dari {filename}"
        )

    return clean_result


parse_cv_document = extract_text_from_bytes
