import io
import re
import math
import logging
from typing import Dict, Any, List, Optional
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.services.pricing_engine import COMPLIANCE_DISCLAIMER, OFFICIAL_PRODUCT_NAME

logger = logging.getLogger(__name__)

# Color Palette Constants
COLOR_PRIMARY = RGBColor(0x1F, 0x4E, 0x78)    # Deep Navy
COLOR_SECONDARY = RGBColor(0x2E, 0x75, 0xB6)  # Accent Blue
COLOR_DARK = RGBColor(0x22, 0x22, 0x22)       # Text Dark
COLOR_MUTED = RGBColor(0x66, 0x66, 0x66)      # Subtext Gray
COLOR_SUCCESS = RGBColor(0x2E, 0x7D, 0x32)    # Green (High Score)
COLOR_WARNING = RGBColor(0xE6, 0x51, 0x00)    # Orange (Medium Score)


def chunk_document_text(text: str, max_chunk_words: int = 1500) -> List[Dict[str, Any]]:
    """Membagi naskah panjang per bab/paragraf dengan tetap mempertahankan kutipan, teori, & sitasi.
    
    Returns:
        List of chunks with 'chunk_index', 'heading', 'text', 'word_count'.
    """
    if not text:
        return []

    paragraphs = [p.strip() for p in text.split("\n") if p.strip()]
    chunks: List[Dict[str, Any]] = []
    current_paragraphs: List[str] = []
    current_word_count = 0
    chunk_idx = 1
    current_heading = "Pendahuluan / Bagian Awal"

    # Pattern deteksi heading bab / section (e.g. "BAB I", "1.1", "Metodologi", "A. Latar Belakang")
    heading_pattern = re.compile(r"^(bab\s+[ivxlcdm\d]+|(\d+\.){1,3}\d*|[a-z]\.\s+|abstrak|pendahuluan|metode|pembahasan|kesimpulan|daftar pustaka)", re.IGNORECASE)

    for p in paragraphs:
        p_words = len(p.split())
        is_heading = bool(heading_pattern.match(p)) and len(p.split()) < 12

        if is_heading and current_paragraphs:
            # Selesaikan chunk sebelumnya jika sudah cukup panjang
            if current_word_count >= 300:
                chunks.append({
                    "chunk_index": chunk_idx,
                    "heading": current_heading,
                    "text": "\n\n".join(current_paragraphs),
                    "word_count": current_word_count
                })
                chunk_idx += 1
                current_paragraphs = []
                current_word_count = 0
            current_heading = p

        if current_word_count + p_words > max_chunk_words and current_paragraphs:
            chunks.append({
                "chunk_index": chunk_idx,
                "heading": current_heading,
                "text": "\n\n".join(current_paragraphs),
                "word_count": current_word_count
            })
            chunk_idx += 1
            current_paragraphs = [p]
            current_word_count = p_words
        else:
            current_paragraphs.append(p)
            current_word_count += p_words

    if current_paragraphs:
        chunks.append({
            "chunk_index": chunk_idx,
            "heading": current_heading,
            "text": "\n\n".join(current_paragraphs),
            "word_count": current_word_count
        })

    return chunks


def _set_document_margins(doc: Document, margin_inches: float = 0.75):
    """Mengatur margin dokumen seragam."""
    for section in doc.sections:
        section.top_margin = Inches(margin_inches)
        section.bottom_margin = Inches(margin_inches)
        section.left_margin = Inches(margin_inches)
        section.right_margin = Inches(margin_inches)


def _add_section_header(doc: Document, title: str, color_rgb: RGBColor = COLOR_PRIMARY):
    """Menambahkan heading section bergaris bawah tipis khas standar HR/profesional."""
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(4)
    run = h.add_run(title.upper())
    run.font.name = 'Calibri'
    run.font.size = Pt(11.5)
    run.font.bold = True
    run.font.color.rgb = color_rgb

    # Garis pembatas bawah elegan
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '2')
    bottom.set(qn('w:color'), '1F4E78')
    pBdr.append(bottom)
    h._p.get_or_add_pPr().append(pBdr)


def _add_bullet_item(doc: Document, text: str, bold_prefix: str = ""):
    """Menambahkan item bullet point dengan format rapi."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    
    if bold_prefix:
        r_bold = p.add_run(bold_prefix)
        r_bold.font.name = 'Calibri'
        r_bold.font.size = Pt(10)
        r_bold.font.bold = True
        r_bold.font.color.rgb = COLOR_DARK

    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.color.rgb = COLOR_DARK


def _add_compliance_footer(doc: Document):
    """Menambahkan disclaimer kepatuhan resmi di bagian bawah dokumen."""
    p_f = doc.add_paragraph()
    p_f.paragraph_format.space_before = Pt(16)
    p_f.paragraph_format.space_after = Pt(4)
    
    # Divider line
    pBdr = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '4')
    top.set(qn('w:space'), '4')
    top.set(qn('w:color'), 'CCCCCC')
    pBdr.append(top)
    p_f._p.get_or_add_pPr().append(pBdr)

    r_f = p_f.add_run(f"📌 {COMPLIANCE_DISCLAIMER}")
    r_f.font.name = 'Calibri'
    r_f.font.size = Pt(8.5)
    r_f.font.italic = True
    r_f.font.color.rgb = COLOR_MUTED


def render_ats_review_docx(data: Dict[str, Any]) -> bytes:
    """Merender hasil audit ATS Review ke file Word (.docx) terstruktur profesional."""
    doc = Document()
    _set_document_margins(doc, 0.75)

    # 1. Header Dokumen
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(2)
    r_title = p_title.add_run("BOONTRACK ATS AUDIT & CV DIAGNOSTIC REPORT")
    r_title.font.name = 'Calibri'
    r_title.font.size = Pt(16)
    r_title.font.bold = True
    r_title.font.color.rgb = COLOR_PRIMARY

    target_role = data.get("target_role") or data.get("target_position") or "General Professional"
    score = int(data.get("overall_score", 0))

    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(10)
    r_sub = p_sub.add_run(f"Target Posisi: {target_role} | Skor ATS: {score}/100")
    r_sub.font.name = 'Calibri'
    r_sub.font.size = Pt(11)
    r_sub.font.bold = True
    r_sub.font.color.rgb = COLOR_SUCCESS if score >= 75 else COLOR_WARNING

    # 2. Executive Summary
    summary = data.get("summary") or data.get("executive_summary") or ""
    if summary:
        _add_section_header(doc, "Executive Summary")
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(str(summary))
        r.font.name = 'Calibri'
        r.font.size = Pt(10.5)
        r.font.color.rgb = COLOR_DARK

    # 3. Breakdown Scores
    breakdown = data.get("breakdown_scores") or data.get("section_scores") or {}
    if isinstance(breakdown, dict) and breakdown:
        _add_section_header(doc, "Skor Parameter ATS")
        for metric, val in breakdown.items():
            metric_clean = str(metric).replace("_", " ").title()
            _add_bullet_item(doc, f": {val}/100", bold_prefix=f"{metric_clean}")

    # 4. Strengths
    strengths = data.get("strengths") or []
    if isinstance(strengths, list) and strengths:
        _add_section_header(doc, "Kelebihan & Kekuatan CV")
        for item in strengths:
            _add_bullet_item(doc, str(item))

    # 5. Weaknesses & Actionable Recommendations
    findings = data.get("findings") or data.get("recommendations") or data.get("improvements") or []
    if isinstance(findings, list) and findings:
        _add_section_header(doc, "Area Peningkatan & Rekomendasi Perbaikan")
        for item in findings:
            if isinstance(item, dict):
                section = item.get("section", "Umum")
                issue = item.get("issue") or item.get("finding") or ""
                rec = item.get("recommendation") or item.get("fix") or ""
                full_text = f"{issue} -> Rekomendasi: {rec}" if rec else issue
                _add_bullet_item(doc, full_text, bold_prefix=f"[{section}] ")
            else:
                _add_bullet_item(doc, str(item))

    _add_compliance_footer(doc)

    out_io = io.BytesIO()
    doc.save(out_io)
    return out_io.getvalue()


def render_cv_rewrite_docx(data: Dict[str, Any]) -> bytes:
    """Merender hasil perombakan CV Rewrite ke file Word (.docx) standar HR internasional."""
    doc = Document()
    _set_document_margins(doc, 0.75)

    full_name = data.get("full_name") or data.get("name") or "KANDIDAT PROFESIONAL"
    email = data.get("email") or ""
    phone = data.get("phone") or ""
    location = data.get("location") or data.get("domicile") or ""
    linkedin = data.get("linkedin") or ""
    portfolio = data.get("portfolio") or data.get("website") or ""

    # 1. Header Nama
    p_name = doc.add_paragraph()
    p_name.paragraph_format.space_before = Pt(0)
    p_name.paragraph_format.space_after = Pt(2)
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_name = p_name.add_run(str(full_name).upper())
    r_name.font.name = 'Calibri'
    r_name.font.size = Pt(16)
    r_name.font.bold = True
    r_name.font.color.rgb = COLOR_PRIMARY

    # Target Subtitle
    target_pos = data.get("target_position") or data.get("title") or ""
    if target_pos:
        p_target = doc.add_paragraph()
        p_target.paragraph_format.space_after = Pt(2)
        p_target.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_t = p_target.add_run(str(target_pos).title())
        r_t.font.name = 'Calibri'
        r_t.font.size = Pt(11)
        r_t.font.bold = True
        r_t.font.color.rgb = COLOR_SECONDARY

    # Contact Info
    contact_items = [x for x in [email, phone, location, linkedin, portfolio] if x]
    if contact_items:
        p_c = doc.add_paragraph()
        p_c.paragraph_format.space_after = Pt(12)
        p_c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_c = p_c.add_run(" | ".join(contact_items))
        r_c.font.name = 'Calibri'
        r_c.font.size = Pt(9.5)
        r_c.font.color.rgb = COLOR_MUTED

    # 2. Professional Summary
    summary = data.get("summary") or data.get("professional_summary") or ""
    if summary:
        _add_section_header(doc, "Professional Summary")
        p_s = doc.add_paragraph()
        p_s.paragraph_format.space_after = Pt(6)
        p_s.paragraph_format.line_spacing = 1.15
        r_s = p_s.add_run(str(summary))
        r_s.font.name = 'Calibri'
        r_s.font.size = Pt(10)
        r_s.font.color.rgb = COLOR_DARK

    # 3. Skills
    skills = data.get("skills") or []
    if skills:
        _add_section_header(doc, "Core Competencies & Skills")
        if isinstance(skills, dict):
            for cat, s_list in skills.items():
                cat_name = str(cat).replace("_", " ").title()
                s_str = ", ".join(s_list) if isinstance(s_list, list) else str(s_list)
                _add_bullet_item(doc, f": {s_str}", bold_prefix=f"{cat_name}")
        elif isinstance(skills, list):
            s_str = ", ".join([str(s) for s in skills if s])
            p_sk = doc.add_paragraph()
            p_sk.paragraph_format.space_after = Pt(6)
            r_sk = p_sk.add_run(s_str)
            r_sk.font.name = 'Calibri'
            r_sk.font.size = Pt(10)
            r_sk.font.color.rgb = COLOR_DARK

    # 4. Professional Experience
    experiences = data.get("experience") or data.get("experiences") or []
    if isinstance(experiences, list) and experiences:
        _add_section_header(doc, "Professional Experience")
        for exp in experiences:
            if not isinstance(exp, dict):
                continue
            role = exp.get("role") or exp.get("job_title") or exp.get("position") or "Position"
            company = exp.get("company") or ""
            period = exp.get("period") or exp.get("dates") or ""
            loc = exp.get("location") or ""
            
            p_exp = doc.add_paragraph()
            p_exp.paragraph_format.space_before = Pt(6)
            p_exp.paragraph_format.space_after = Pt(1)
            
            r_role = p_exp.add_run(f"{role} — {company}" if company else role)
            r_role.font.name = 'Calibri'
            r_role.font.size = Pt(10.5)
            r_role.font.bold = True
            r_role.font.color.rgb = COLOR_PRIMARY
            
            if period or loc:
                p_meta = doc.add_paragraph()
                p_meta.paragraph_format.space_after = Pt(3)
                meta_str = " | ".join([x for x in [period, loc] if x])
                r_meta = p_meta.add_run(meta_str)
                r_meta.font.name = 'Calibri'
                r_meta.font.size = Pt(9.5)
                r_meta.font.italic = True
                r_meta.font.color.rgb = COLOR_MUTED

            bullets = exp.get("bullets") or exp.get("achievements") or exp.get("responsibilities") or []
            if isinstance(bullets, list):
                for b in bullets:
                    _add_bullet_item(doc, str(b))

    # 5. Education
    education = data.get("education") or []
    if isinstance(education, list) and education:
        _add_section_header(doc, "Education")
        for edu in education:
            if isinstance(edu, dict):
                degree = edu.get("degree") or edu.get("major") or ""
                institution = edu.get("institution") or edu.get("school") or edu.get("university") or ""
                year = edu.get("year") or edu.get("graduation_year") or ""
                title = f"{degree} — {institution}" if institution else degree
                p_edu = doc.add_paragraph()
                p_edu.paragraph_format.space_before = Pt(4)
                p_edu.paragraph_format.space_after = Pt(1)
                r_e = p_edu.add_run(title)
                r_e.font.name = 'Calibri'
                r_e.font.size = Pt(10)
                r_e.font.bold = True
                if year:
                    r_y = p_edu.add_run(f" ({year})")
                    r_y.font.name = 'Calibri'
                    r_y.font.size = Pt(9.5)
                    r_y.font.italic = True
            else:
                _add_bullet_item(doc, str(edu))

    # 6. Certifications
    certs = data.get("certifications") or data.get("licenses") or []
    if isinstance(certs, list) and certs:
        _add_section_header(doc, "Certifications & Training")
        for c in certs:
            _add_bullet_item(doc, str(c))

    _add_compliance_footer(doc)

    out_io = io.BytesIO()
    doc.save(out_io)
    return out_io.getvalue()


def render_paraphrase_docx(data: Dict[str, Any]) -> bytes:
    """Merender hasil polish & rephrase dokumen ke file Word (.docx)."""
    doc = Document()
    _set_document_margins(doc, 0.75)

    title = data.get("title") or f"{OFFICIAL_PRODUCT_NAME.upper()}"
    tone = data.get("tone") or "Profesional & Formal"
    orig_words = data.get("original_word_count", 0)
    para_words = data.get("paraphrased_word_count", 0)

    # 1. Header Dokumen
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(2)
    r_title = p_title.add_run(str(title).upper())
    r_title.font.name = 'Calibri'
    r_title.font.size = Pt(16)
    r_title.font.bold = True
    r_title.font.color.rgb = COLOR_PRIMARY

    # Subtitle Metadata
    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(10)
    r_sub = p_sub.add_run(f"Tone: {tone} | Kata Asli: {orig_words} -> Hasil: {para_words}")
    r_sub.font.name = 'Calibri'
    r_sub.font.size = Pt(10)
    r_sub.font.color.rgb = COLOR_MUTED

    # 2. Executive Key Takeaways / Summary
    key_points = data.get("key_takeaways") or data.get("summary") or []
    if key_points:
        _add_section_header(doc, "Ringkasan Poin Utama")
        if isinstance(key_points, list):
            for pt in key_points:
                _add_bullet_item(doc, str(pt))
        elif isinstance(key_points, str):
            p_s = doc.add_paragraph(key_points)
            p_s.paragraph_format.space_after = Pt(6)
            for r in p_s.runs:
                r.font.name = 'Calibri'
                r.font.size = Pt(10)

    # 3. Main Paraphrased Content
    sections = data.get("sections") or []
    if isinstance(sections, list) and sections:
        _add_section_header(doc, "Naskah Hasil Polish & Rephrase")
        for sec in sections:
            if isinstance(sec, dict):
                heading = sec.get("heading") or sec.get("title") or ""
                content = sec.get("content") or sec.get("text") or ""
                if heading:
                    p_h = doc.add_paragraph()
                    p_h.paragraph_format.space_before = Pt(6)
                    p_h.paragraph_format.space_after = Pt(2)
                    r_h = p_h.add_run(str(heading))
                    r_h.font.name = 'Calibri'
                    r_h.font.size = Pt(11)
                    r_h.font.bold = True
                    r_h.font.color.rgb = COLOR_SECONDARY
                if content:
                    p_c = doc.add_paragraph(str(content))
                    p_c.paragraph_format.space_after = Pt(6)
                    p_c.paragraph_format.line_spacing = 1.15
                    for r in p_c.runs:
                        r.font.name = 'Calibri'
                        r.font.size = Pt(10.5)
                        r.font.color.rgb = COLOR_DARK
            elif isinstance(sec, str):
                p_c = doc.add_paragraph(sec)
                p_c.paragraph_format.space_after = Pt(6)
                p_c.paragraph_format.line_spacing = 1.15
                for r in p_c.runs:
                    r.font.name = 'Calibri'
                    r.font.size = Pt(10.5)
                    r.font.color.rgb = COLOR_DARK
    else:
        full_text = data.get("full_text") or data.get("paraphrased_text") or ""
        if full_text:
            _add_section_header(doc, "Naskah Hasil Polish & Rephrase")
            p_f = doc.add_paragraph(str(full_text))
            p_f.paragraph_format.space_after = Pt(8)
            p_f.paragraph_format.line_spacing = 1.15
            for r in p_f.runs:
                r.font.name = 'Calibri'
                r.font.size = Pt(10.5)
                r.font.color.rgb = COLOR_DARK

    _add_compliance_footer(doc)

    out_io = io.BytesIO()
    doc.save(out_io)
    return out_io.getvalue()


def build_document_result(task_type: str, structured_data: Dict[str, Any]) -> bytes:
    """Dispatcher utama pembuat file Word (.docx) berdasarkan task_type."""
    normalized = str(task_type).upper().strip()
    if normalized in ["ATS_DIAGNOSTIC", "ATS_REVIEW"]:
        return render_ats_review_docx(structured_data)
    elif normalized in ["CV_POLISH_REWRITE", "CV_REWRITE", "CAREER_PRO_BUNDLE"]:
        return render_cv_rewrite_docx(structured_data)
    elif normalized in ["POLISH_REPHRASE", "PARAPHRASE"]:
        return render_paraphrase_docx(structured_data)
    else:
        return render_cv_rewrite_docx(structured_data)
