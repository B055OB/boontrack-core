import os
import re
import json
import random
import tempfile
import uuid
import time
import asyncio
import logging
from datetime import datetime, timedelta

from aiogram import types, Dispatcher, Bot
from aiogram.types import InlineKeyboardMarkup, InlineKeyboardButton, InputFile

from app.core.bot import bot, send_chunked_message
from app.core.database import (
    get_db_connection,
    save_user,
    track_event,
    get_user_history,
    save_dropoff,
    save_cv_version,
    count_referrals,
    check_user_paid,
    create_order,
    create_donation_session,
    match_and_complete_order,
    match_and_complete_donation
)
from app.services.analytics_service import analytics_service
from app.services.whatsapp_service import log_to_supabase_messages
from app.services.ai_gateway import ai_gateway
from app.services.brain_engine import BrainEngine
from app.services.cloudflare_service import get_user_slug, generate_unique_slug, sync_profile_to_cloudflare_kv
from app.services.document_parser_service import parse_cv_document
from app.services.receipt_ocr_service import parse_receipt_image
from app.handlers.admin_handler import admin_handler
from app.handlers.career_page_flow import register_career_page_handlers, start_career_page_claim
from app.services.cv_flow_service import (
    TOTAL_STEPS,
    REQUIRED_REFERRALS,
    CLOSING_WORDS,
    get_progress_bar,
    clean_val,
    get_question_text,
    format_telegram_review_response,
    handle_cv_review_process,
    ai_career_chat_response,
    create_cv_docx,
    get_career_home_keyboard,
    get_donation_options_keyboard,
    process_and_send_cv
)

logger = logging.getLogger("TELEGRAM_HANDLERS")

QRIS_IMAGE_PATH = "assets/qris.jpg"
EBOOK_FILE_ID = os.getenv("EBOOK_FILE_ID", "YOUR_TELEGRAM_EBOOK_FILE_ID")

user_state = {}

def register_all_bot_handlers(dp: Dispatcher, bot: Bot):
    """Mendaftarkan seluruh command, callback, media, dan message handler ke Dispatcher."""
    register_career_page_handlers(dp)

    async def handle_admin_commands(message: types.Message):
        response = await admin_handler.handle_admin_command(message.from_user.id, message.text)
        await message.reply(response, parse_mode="Markdown")

    @dp.message_handler(commands=['start'])
    async def send_welcome(message: types.Message):
        user_id = message.from_user.id
        first_name = message.from_user.first_name or "Teman"
        await save_user(message.from_user)

        # Log ke Supabase
        await log_to_supabase_messages(
            sender=f"Telegram / {first_name}",
            text="/start",
            tenant_id="boontrack-career",
            channel="telegram",
            user_id=str(user_id),
            user_name=first_name
        )

        text_parts = message.text.split()
        args = text_parts[1] if len(text_parts) > 1 else "direct"

        meta_data = {}
        if args.startswith("ref_"):
            meta_data = {"utm_source": "referral", "referrer_id": args.replace("ref_", "")}
        elif args.startswith("CLK-"):
            meta_data = {"utm_source": "click_logs", "click_id": args}
            def _link_user_attribution():
                try:
                    conn = get_db_connection()
                    cur = conn.cursor()
                    cur.execute("""
                        UPDATE click_logs 
                        SET telegram_user_id = %s, event_name = 'start_bot'
                        WHERE click_id = %s
                    """, (user_id, args))
                    conn.commit()
                    cur.close()
                    conn.close()
                except Exception as e:
                    print(f"[Attribution Error] {e}", flush=True)

            asyncio.create_task(asyncio.to_thread(_link_user_attribution))
        else:
            meta_data = {"utm_source": args}
            asyncio.create_task(analytics_service.save_user_utm(user_id, args))

        asyncio.create_task(track_event(user_id, "start", meta=meta_data))
        progress, last_cv = await get_user_history(user_id)
        saved_data = progress.get("data", {}) if progress else {}
        user_name = saved_data.get("nama_panggilan") or message.from_user.first_name or "Teman"

        if progress is not None:
            last_step = progress.get("last_step", 0)

            if last_step == TOTAL_STEPS or last_step == 0:
                user_state[user_id] = {"step": 0, "data": saved_data, "meta": meta_data}
                home_msg = (
                    f"Halo lagi, <b>{user_name}</b>! 👋\n\n"
                    "Ada yang bisa saya bantu untuk persiapan kariermu hari ini?\n\n"
                    "👇 <i>Pilih opsi di bawah:</i>"
                )
                kbd = get_career_home_keyboard()
                await message.reply(home_msg, reply_markup=kbd, parse_mode="HTML")
                return

            if isinstance(last_step, int) and last_step > 0:
                user_state[user_id] = {"step": last_step, "data": saved_data, "meta": meta_data}
                kbd = InlineKeyboardMarkup(row_width=2)
                kbd.add(
                    InlineKeyboardButton("▶️ Lanjutkan CV", callback_data="resume_flow"),
                    InlineKeyboardButton("🔄 Mulai Baru", callback_data="restart_flow")
                )
                await message.reply(
                    f"Halo lagi, <b>{user_name}</b>! 👋\n\n"
                    f"Kemarin kita sempat menyusun CV sampai di <b>Langkah {last_step} dari {TOTAL_STEPS}</b>.\n\n"
                    "Mau kita tuntaskan sekarang agar CV kamu siap dipakai melamar kerja?",
                    reply_markup=kbd,
                    parse_mode="HTML"
                )
                return

        user_state[user_id] = {"step": "ONBOARDING_NAMA", "data": {}, "meta": meta_data}
        await save_dropoff(user_id, 0, {})

        msg_1 = (
            "<b>Saya BoonTrack Career Assistant.</b>\n"
            "Saya akan membantu meningkatkan peluang kamu dipanggil interview.\n\n"
            "Sebelum mulai...\n"
            "Boleh kenalan dulu?\n"
            "Ini dengan siapa?"
        )
        await message.reply(msg_1, parse_mode="HTML")

    @dp.callback_query_handler(lambda c: c.data == "trigger_cv_review")
    async def handle_trigger_cv_review(callback_query: types.CallbackQuery):
        user_id = callback_query.from_user.id
        first_name = callback_query.from_user.first_name or "Teman"

        try:
            await callback_query.answer()
        except Exception:
            pass

        if user_id not in user_state:
            user_state[user_id] = {"step": 0, "data": {}}

        user_data = user_state.get(user_id, {}).get("data", {})
        cv_text_summary = " ".join([str(v) for k, v in user_data.items() if str(v).strip()]).strip()
        position = str(user_data.get("target_position") or user_data.get(6) or user_data.get("6") or "General Professional")

        if cv_text_summary and len(cv_text_summary) > 20:
            from app.handlers.commands import render_free_cv_review
            await render_free_cv_review(user_id, bot, cv_text_summary, target_position=position)
            return

        user_state[user_id]["step"] = "WAITING_CV_INPUT"

        kbd_review = types.InlineKeyboardMarkup(row_width=1)
        kbd_review.add(
            types.InlineKeyboardButton("📝 Buat CV Baru via Chat", callback_data="home_create_cv"),
            types.InlineKeyboardButton("🏠 Menu Utama", callback_data="home_back_main")
        )
        msg_prompt = (
            "🔍 <b>DIAGNOSIS & REVIEW CV GRATIS (ATS COMPLIANT)</b>\n"
            "━━━━━━━━━━━━━━━━━━━━━━━━\n"
            "Silakan kirimkan CV kamu untuk dianalisis dan dinilai oleh AI:\n\n"
            "📁 <b>Upload File:</b> Kirim file CV kamu dalam format <b>Word (.docx)</b>, <b>PDF (.pdf)</b>, atau <b>.txt</b>\n"
            "✍️ <b>Ketik / Paste:</b> Atau salin ringkasan teks CV kamu langsung ke chat ini.\n\n"
            "<i>Mau buat CV ATS baru dari nol? Klik tombol di bawah:</i>"
        )

        await callback_query.message.answer(msg_prompt, reply_markup=kbd_review, parse_mode="HTML")
        return

    @dp.callback_query_handler(lambda c: True)
    async def handle_callback_navigation(callback_query: types.CallbackQuery):
        user_id = callback_query.from_user.id
        code = callback_query.data

        try:
            await callback_query.answer()
        except Exception:
            pass

        if code in ["home_back_main", "main_menu", "back_to_main"]:
            user_state[user_id] = {"step": 0, "data": {}}
            from app.handlers.commands import send_welcome
            await send_welcome(callback_query.message)
            return

        try:
            await callback_query.message.edit_reply_markup(reply_markup=None)
        except Exception:
            pass

        if user_id not in user_state:
            progress, _ = await get_user_history(user_id)
            saved_data = progress.get("data", {}) if progress else {}
            user_state[user_id] = {"step": 0, "data": saved_data}

        user_data = user_state[user_id].get("data", {})
        user_name = user_data.get("nama_panggilan", callback_query.from_user.first_name or "Teman")
        slug = get_user_slug(user_data, callback_query.from_user.first_name)

        if code == "trigger_cv_review":
            from app.handlers.commands import render_free_cv_review

            position = str(user_data.get("target_position") or user_data.get(6) or user_data.get("6") or "General Professional")
            cv_text_summary = " ".join([str(v) for k, v in user_data.items() if str(v).strip()]).strip()

            if cv_text_summary and len(cv_text_summary) > 20:
                await render_free_cv_review(user_id, bot, cv_text_summary, target_position=position)
                return
            else:
                user_state[user_id]["step"] = "WAITING_CV_INPUT"

                kbd_review = types.InlineKeyboardMarkup(row_width=1)
                kbd_review.add(
                    types.InlineKeyboardButton("📝 Buat CV Baru via Chat", callback_data="home_create_cv"),
                    types.InlineKeyboardButton("🏠 Menu Utama", callback_data="home_back_main")
                )
                msg_prompt = (
                    "🔍 <b>DIAGNOSIS & REVIEW CV GRATIS (ATS COMPLIANT)</b>\n"
                    "━━━━━━━━━━━━━━━━━━━━━━━━\n"
                    "Silakan kirimkan CV kamu untuk dianalisis dan dinilai oleh AI:\n\n"
                    "📁 <b>Upload File:</b> Kirim file CV dalam format <b>Word (.docx)</b>, <b>PDF (.pdf)</b>, atau <b>.txt</b>\n"
                    "✍️ <b>Ketik / Paste:</b> Atau salin ringkasan teks CV langsung ke chat ini.\n\n"
                    "<i>Mau buat CV ATS baru dari nol? Klik tombol di bawah:</i>"
                )
                await bot.send_message(user_id, msg_prompt, reply_markup=kbd_review, parse_mode="HTML")
                return


        elif code in ["don_5000", "don_10000", "don_25000"]:
            base_amt = 5000 if code == "don_5000" else (10000 if code == "don_10000" else 25000)
            unique_code = random.randint(100, 999)
            total_amt = base_amt + unique_code

            await create_donation_session(user_id, base_amt, unique_code, total_amt)

            don_msg = (
                f"🎉 <b>Terima kasih telah memilih BoonTrack!</b>\n\n"
                f"Tinggal satu langkah lagi untuk mengaktifkan <b>Career Page Profesional</b> milikmu dan tampil lebih menonjol di mata HRD/Klien.\n\n"
                f"🌐 <b>Contoh Tampilan Career Page:</b>\n"
                f"Lihat preview tampilan Career Page yang akan kamu dapatkan di sini:\n"
                f"👉 https://rayigemilang.boontrack.com\n\n"
                f"<i>✨ Format modern, recruiter-friendly, responsif di HP/laptop, dan <b>aktif seumur hidup (sekali bayar tanpa biaya langganan)</b>.</i>\n\n"
                f"💳 <b>Rincian Pembayaran:</b>\n"
                f"• <b>Item:</b> Aktivasi Career Page Personal (Lifetime Access)\n"
                f"• <b>Transfer Tepat:</b> <code>Rp{total_amt:,}</code> <i>(Wajib transfer sesuai hingga 3 digit terakhir)</i>\n"
                f"• <b>Rincian:</b> Rp{base_amt:,} + kode verifikasi Rp{unique_code}\n"
                f"• <b>Masa Aktif Web:</b> <b>Aktif Seumur Hidup</b>\n"
                f"• <b>Batas Waktu Bayar:</b> 15 Menit\n\n"
                f"📱 <b>Panduan Bayar via QRIS (Jika Pakai 1 HP):</b>\n"
                f"1. <b>Simpan QR:</b> <b>Screenshot layar ini</b> atau unduh gambar QRIS di atas.\n"
                f"2. <b>Buka E-Wallet / Mobile Banking:</b> (BCA, Mandiri, BRI, DANA, GoPay, OVO, ShopeePay, dll).\n"
                f"3. <b>Pilih Menu QRIS / Scan:</b> Buka scanner QRIS di aplikasimu.\n"
                f"4. <b>Upload dari Galeri:</b> Klik ikon galeri/foto di menu scanner.\n"
                f"5. <b>Pilih Screenshot QR:</b> Masukkan gambar QR tadi & pastikan nominalnya tepat <b>Rp{total_amt:,}</b>.\n"
                f"6. Selesaikan pembayaran.\n\n"
                f"⏳ <i>Sistem otomatis memverifikasi pembayaran tanpa perlu kirim bukti transfer. Setelah terdeteksi, bot akan langsung mengirimkan pilihan link subdomain personalmu!</i>"
            )

            kbd_qris = InlineKeyboardMarkup(row_width=1)
            kbd_qris.add(
                InlineKeyboardButton("⏳ Bayar Nanti (Kembali ke Menu Utama)", callback_data="home_back_main"),
                InlineKeyboardButton("❌ Batalkan Transaksi", callback_data="cancel_checkout")
            )

            possible_qris_paths = [QRIS_IMAGE_PATH, "/app/qris.jpg", "qris.jpg"]
            found_qris = next((p for p in possible_qris_paths if os.path.exists(p)), None)
            if found_qris:
                await bot.send_photo(chat_id=user_id, photo=InputFile(found_qris))
                await send_chunked_message(user_id, don_msg, reply_markup=kbd_qris, parse_mode="HTML")
            else:
                await send_chunked_message(user_id, don_msg, reply_markup=kbd_qris, parse_mode="HTML")

        elif code == "cancel_checkout":
            user_state[user_id]["step"] = 0
            kbd = await get_career_home_keyboard(user_id)
            await bot.send_message(user_id, "❌ <b>Transaksi dibatalkan.</b> Kembali ke menu utama:", reply_markup=kbd, parse_mode="HTML")

        elif code in ["cp_build_now", "cp_manage"]:
            is_paid = await check_user_paid(user_id)
            if not is_paid and code == "cp_manage":
                don_msg = (
                    f"🔒 <b>Website Career Page Belum Aktif</b>\n\n"
                    f"Kamu perlu mengaktifkan akses Career Page terlebih dahulu (Rp10.000) untuk mengakses menu ini."
                )
                await bot.send_message(user_id, don_msg, reply_markup=get_donation_options_keyboard(), parse_mode="HTML")
                return

            user_data["cp_status"] = "active"
            default_slug = await generate_unique_slug(user_data)
            user_data["temp_slug"] = default_slug
            user_state[user_id]["data"] = user_data

            kbd_post = InlineKeyboardMarkup(row_width=1)
            kbd_post.add(
                InlineKeyboardButton(f"✅ Pakai {default_slug}.boontrack.com", callback_data="cp_confirm_default_slug"),
                InlineKeyboardButton("✏️ Ketik Nama Custom Sendiri", callback_data="cp_change_slug_start")
            )

            await bot.send_message(
                user_id,
                "🎉 <b>Akses Career Page Aktif!</b>\n\n"
                "Mari tentukan nama link subdomain untuk Career Page milikmu:\n\n"
                f"<b>Rekomendasi Subdomain:</b>\n"
                f"👉 <code>{default_slug}.boontrack.com</code>\n\n"
                "Apakah kamu mau memakai nama rekomendasi di atas, atau ingin mengetik nama custom sendiri?",
                reply_markup=kbd_post,
                parse_mode="HTML"
            )

        elif code == "cp_confirm_default_slug":
            user_data = user_state[user_id].get("data", {})
            final_slug = user_data.get("temp_slug") or await generate_unique_slug(user_data)
            user_data["slug"] = final_slug
            user_state[user_id]["data"] = user_data

            await update_cloudflare_kv(final_slug, user_data)

            await bot.send_message(
                user_id,
                f"✅ <b>Career Page Berhasil Diterbitkan!</b>\n\n"
                f"🌐 Link portofolio kamu: https://{final_slug}.boontrack.com",
                parse_mode="HTML"
            )

        elif code == "cp_change_slug_start":
            user_state[user_id]["step"] = "CP_INPUT_CUSTOM_SLUG"
            await bot.send_message(
                user_id,
                "✏️ <b>Ketik nama subdomain (slug) baru yang kamu inginkan:</b>\n\n"
                "<i>Contoh: ketik <code>alldy-pro</code> untuk mendapatkan link https://alldy-pro.boontrack.com</i>\n"
                "<i>(Hanya huruf, angka, dan tanda hubung [-])</i>",
                parse_mode="HTML"
            )

        elif code == "cp_edit_slug":
            user_state[user_id]["step"] = "CP_INPUT_CUSTOM_SLUG"
            await bot.send_message(
                user_id,
                f"🔗 <b>Ubah Subdomain / Slug Website Kamu</b>\n\n"
                f"Subdomain kamu saat ini: <code>{slug}</code> (https://{slug}.boontrack.com)\n\n"
                f"Ketik nama subdomain kustom baru yang kamu inginkan:\n"
                f"<i>(Contoh: ratuhrd, ratu-official, rayigemilang)</i>",
                parse_mode="HTML"
            )

        elif code == "cp_edit_data":
            kbd_sections = InlineKeyboardMarkup(row_width=1)
            kbd_sections.add(
                InlineKeyboardButton("💼 Edit Posisi / Headline", callback_data="cp_edit_posisi_btn"),
                InlineKeyboardButton("📝 Edit Ringkasan Profil / Bio", callback_data="cp_edit_summary_btn"),
                InlineKeyboardButton("🏢 Edit Pengalaman Kerja / Proyek", callback_data="cp_edit_exp_btn"),
                InlineKeyboardButton("🛠️ Edit Keahlian / Skill Grid", callback_data="cp_edit_skills_btn"),
                InlineKeyboardButton("🔙 Batal / Kembali ke Menu Career Page", callback_data="cp_manage")
            )
            await bot.send_message(user_id, "✏️ <b>Pilih Bagian yang Ingin Kamu Isi atau Edit:</b>", reply_markup=kbd_sections, parse_mode="HTML")

        elif code == "cp_edit_posisi_btn":
            user_state[user_id]["step"] = "CP_EDIT_POSISI"
            await bot.send_message(user_id, "💼 <b>Edit Posisi / Headline Website</b>\n\nKetik judul posisi impianmu:\n<i>(Contoh: AI & Operations Workflow Optimization Specialist)</i>", parse_mode="HTML")

        elif code == "cp_edit_summary_btn":
            user_state[user_id]["step"] = "CP_EDIT_SUMMARY"
            await bot.send_message(user_id, "📝 <b>Edit Ringkasan Profil / Bio Website</b>\n\nKetik deskripsi singkat tentang dirimu (1-3 kalimat):\n<i>(Contoh: Membantu tim operasional memangkas waktu kerja manual dengan otomatisasi sistem)</i>", parse_mode="HTML")

        elif code == "cp_edit_exp_btn":
            user_state[user_id]["step"] = "CP_EDIT_EXP"
            await bot.send_message(user_id, "🏢 <b>Edit Pengalaman Kerja / Proyek Website</b>\n\nKetik detail pengalaman kerja atau portofolio utamamu:\n<i>(Contoh: Manager HRD di PT ABC (2022-Sekarang) - Memimpin tim 10 orang & merekrut 50+ karyawan)</i>", parse_mode="HTML")

        elif code == "cp_edit_skills_btn":
            user_state[user_id]["step"] = "CP_EDIT_SKILLS"
            await bot.send_message(user_id, "🛠️ <b>Edit Keahlian / Skill Website</b>\n\nKetik skill utama dipisahkan dengan koma:\n<i>(Contoh: Python, OpenAI API, Cloudflare Workers, SQL, Recruitment)</i>", parse_mode="HTML")

        elif code == "cp_edit_resume":
            user_state[user_id]["step"] = "CP_EDIT_RESUME"
            await bot.send_message(
                user_id,
                "📄 <b>Input / Update Link Resume PDF</b>\n\n"
                "Ketik atau paste link Google Drive / tautan publik PDF resume kamu di sini:\n"
                "<b>Contoh:</b> <code>https://drive.google.com/file/d/1A2b3C.../view?usp=sharing</code>\n\n"
                "<i>Ketik '-' jika ingin menyembunyikan tombol download resume.</i>",
                parse_mode="HTML"
            )

        elif code == "cp_import_cv":
            kbd_import = InlineKeyboardMarkup(row_width=1)
            kbd_import.add(
                InlineKeyboardButton("✅ Ya, Gunakan Semua Data CV", callback_data="cp_confirm_import"),
                InlineKeyboardButton("✏️ Batal, Pilih Edit Bagian Manual", callback_data="cp_edit_data")
            )
            await bot.send_message(
                user_id,
                "⚠️ <b>Konfirmasi Impor Data CV</b>\n\n"
                "Sistem akan menyalin ringkasan profil, kontak, dan keahlian dari draf CV ke website.\n"
                "Kamu tetap bisa mengedit bagian mana saja kapan pun!",
                reply_markup=kbd_import,
                parse_mode="HTML"
            )

        elif code == "cp_confirm_import":
            user_data["ringkasan_web"] = user_data.get("3", "")
            user_data["pengalaman_web"] = user_data.get("3", "")
            user_data["keahlian_web"] = user_data.get("6", "")

            await save_dropoff(user_id, TOTAL_STEPS, user_data)
            await update_cloudflare_kv(slug, user_data)

            kbd_done = InlineKeyboardMarkup(row_width=1)
            kbd_done.add(
                InlineKeyboardButton("🌐 Buka Website Live", url=f"https://{slug}.boontrack.com"),
                InlineKeyboardButton("🔙 Kembali ke Menu Career Page", callback_data="cp_manage"),
                InlineKeyboardButton("🏠 Menu Utama", callback_data="home_back_main")
            )
            await bot.send_message(
                user_id,
                f"✅ <b>Data CV Berhasil Diimpor & Disinkronkan!</b>\n\n"
                f"Tampilan web kamu sudah terbarui secara realtime di:\n"
                f"👉 https://{slug}.boontrack.com",
                reply_markup=kbd_done,
                parse_mode="HTML"
            )

        elif code == "cp_build_later":
            kbd = await get_career_home_keyboard(user_id)
            await bot.send_message(
                user_id,
                f"Siap, {user_name}! Akses pembuatan Career Page kamu sudah tersimpan aman.\n"
                f"Kapan saja kamu siap melengkapi datanya, tinggal klik menu <b>'🌐 Kelola Career Page Saya'</b> di Menu Utama! 👍",
                reply_markup=kbd,
                parse_mode="HTML"
            )

        elif code == "cp_upload_photo":
            user_state[user_id]["step"] = "WAITING_PHOTO"
            await bot.send_message(user_id, "📸 <b>Kirimkan foto profil terbaikmu ke chat ini ya!</b>\n<i>(Disarankan foto formal/semi-formal setengah badan)</i>", parse_mode="HTML")

        elif code == "cp_choose_theme":
            kbd_theme = InlineKeyboardMarkup(row_width=2)
            kbd_theme.add(
                InlineKeyboardButton("💛 Happy Gold", callback_data="theme_happy"),
                InlineKeyboardButton("💙 Modern Blue", callback_data="theme_blue"),
                InlineKeyboardButton("🖤 Dark Minimalist", callback_data="theme_dark"),
                InlineKeyboardButton("💚 Emerald Green", callback_data="theme_emerald"),
                InlineKeyboardButton("💜 Elegant Purple", callback_data="theme_purple")
            )
            await bot.send_message(user_id, "🎨 <b>Pilih tema warna favoritmu untuk Career Page:</b>", reply_markup=kbd_theme, parse_mode="HTML")

        elif code.startswith("theme_"):
            selected_theme = code.replace("theme_", "")
            user_data["theme"] = selected_theme

            await save_dropoff(user_id, TOTAL_STEPS, user_data)
            await update_cloudflare_kv(slug, user_data)

            kbd_done = InlineKeyboardMarkup(row_width=1)
            kbd_done.add(
                InlineKeyboardButton("🌐 Buka Website Live", url=f"https://{slug}.boontrack.com"),
                InlineKeyboardButton("🔙 Kembali ke Menu Career Page", callback_data="cp_manage"),
                InlineKeyboardButton("🏠 Menu Utama", callback_data="home_back_main")
            )
            await bot.send_message(
                user_id,
                f"🎨 <b>Tema berhasil diubah ke {selected_theme.capitalize()}!</b>\n\n"
                f"Cek perubahannya secara langsung di:\n"
                f"👉 https://{slug}.boontrack.com",
                reply_markup=kbd_done,
                parse_mode="HTML"
            )

        elif code == "cp_deploy_live":
            is_success = await update_cloudflare_kv(slug, user_data)
            if is_success:
                msg = (
                    f"🎉 <b>SELAMAT! Website Career Page Kamu Resmi Aktif!</b>\n\n"
                    f"👉 <b>Link Web Live:</b> https://{slug}.boontrack.com\n\n"
                    f"Website ini sudah siap kamu pajang di bio LinkedIn atau WhatsApp kamu! 🚀"
                )
            else:
                msg = "⚠️ Terjadi masalah sinkronisasi server KV. Pastikan konfigurasi `.env` sudah sesuai."

            kbd_done = InlineKeyboardMarkup(row_width=1)
            kbd_done.add(
                InlineKeyboardButton("🌐 Buka Website Live", url=f"https://{slug}.boontrack.com"),
                InlineKeyboardButton("🔙 Kembali ke Menu Career Page", callback_data="cp_manage")
            )
            await bot.send_message(user_id, msg, reply_markup=kbd_done, parse_mode="HTML")

        elif code == "home_digital_products":
            kbd_products = InlineKeyboardMarkup(row_width=1)
            kbd_products.add(
                InlineKeyboardButton("📘 Ebook Panduan Lolos Interview & Gaji (Rp49.000)", callback_data="buy_ebook_interview"),
                InlineKeyboardButton("🔙 Kembali ke Menu Utama", callback_data="home_back_main")
            )
            msg_catalog = (
                "🚀 <b>PROGRAM & PRODUK DIGITAL KARIR</b>\n\n"
                "Tingkatkan peluang dipanggil dan lolos kerja dengan panduan eksklusif dari BoonTrack:\n\n"
                "📖 <b>Ebook Panduan Lolos Interview & Negosiasi Gaji</b>\n"
                "• Rangkuman pertanyaan jebakan HR + cara jawabnya\n"
                "• Template surat lamaran & email melamar kerja\n"
                "• Strategi negosiasi gaji untuk Fresh Graduate & Exp\n\n"
                "👇 <i>Klik tombol di bawah untuk membeli secara otomatis:</i>"
            )
            await bot.send_message(user_id, msg_catalog, reply_markup=kbd_products, parse_mode="HTML")

        elif code == "buy_ebook_interview":
            base_price = 50000
            unique_code = random.randint(100, 999)
            total_amount = base_price - unique_code

            await create_order(user_id, "Ebook Interview", base_price, unique_code, total_amount)

            msg_checkout = (
                f"🛒 <b>CHECKOUT: Ebook Panduan Lolos Interview</b>\n\n"
                f"💵 Harga Normal: <s>Rp{base_price:,}</s>\n"
                f"🎉 <b>Total Transfer (Dapat Potongan):</b>\n"
                f"<code>{total_amount}</code> 👈 <i>(Tekan/salin angka ini)</i>\n\n"
                f"👇 <b>Cara Pembayaran:</b>\n"
                f"1. Scan QRIS di atas atau transfer via DANA Bisnis.\n"
                f"2. Masukkan nominal <b>PRESISI <code>{total_amount}</code></b> (sampai 3 digit terakhir).\n"
                f"3. Dalam 1-3 detik setelah transfer, Ebook otomatis terkirim di sini! 🚀\n\n"
                f"⏳ <i>Nominal unik ini berlaku selama 15 menit.</i>"
            )

            kbd_qris = InlineKeyboardMarkup(row_width=1)
            kbd_qris.add(
                InlineKeyboardButton("⏳ Bayar Nanti (Kembali ke Menu Utama)", callback_data="home_back_main"),
                InlineKeyboardButton("❌ Batalkan Transaksi", callback_data="cancel_checkout")
            )

            possible_qris_paths = [QRIS_IMAGE_PATH, "/app/qris.jpg", "qris.jpg"]
            found_qris = next((p for p in possible_qris_paths if os.path.exists(p)), None)
            if found_qris:
                await bot.send_photo(chat_id=user_id, photo=InputFile(found_qris), caption=msg_checkout, reply_markup=kbd_qris, parse_mode="HTML")
            else:
                await bot.send_message(user_id, msg_checkout, reply_markup=kbd_qris, parse_mode="HTML")

        elif code == "home_back_main":
            current_data = user_state.get(user_id, {}).get("data", {})
            user_state[user_id] = {"step": 0, "data": current_data}
            kbd = await get_career_home_keyboard(user_id)
            await bot.send_message(user_id, "👋 <b>Kembali ke Menu Utama:</b>", reply_markup=kbd, parse_mode="HTML")

        elif code == "restart_flow":
            user_state[user_id] = {"step": 0, "data": {}}
            kbd = await get_career_home_keyboard(user_id)
            await bot.send_message(user_id, "👋 <b>Menu Utama (Data Reset):</b>", reply_markup=kbd, parse_mode="HTML")

        elif code == "home_create_cv":
            old_name = user_data.get("nama_panggilan", callback_query.from_user.first_name or "")
            new_data = {"nama_panggilan": old_name} if old_name else {}
            user_state[user_id] = {"step": "ONBOARDING_NAMA", "data": new_data}
            await save_dropoff(user_id, 0, new_data)

            if old_name:
                msg_restart = f"Sip, {old_name}! Kita susun versi CV baru ya. 👍\n\nKamu mau tetap pakai nama panggilan <b>{old_name}</b> atau mau ganti nama baru?\n<i>(Ketik langsung nama panggilanmu di bawah untuk melanjutkan)</i>"
            else:
                msg_restart = "Sip! Kita susun versi CV baru ya. 👍\n\nBoleh kenalan dulu?\n<b>Ini dengan siapa?</b>"
            await bot.send_message(user_id, msg_restart, parse_mode="HTML")

        elif code == "home_check_ref":
            total_refs = await count_referrals(user_id)
            bot_info = await bot.get_me()
            user_ref_link = f"https://t.me/{bot_info.username}?start=ref_{user_id}"
            kbd = await get_career_home_keyboard(user_id)

            ref_msg = (
                "🎁 <b>REFERRAL & BONUS PORTFOLIO WEBSITE</b>\n\n"
                f"📊 <b>Progress Referral Kamu: {total_refs} / {REQUIRED_REFERRALS}</b>\n\n"
                f"Ajak {REQUIRED_REFERRALS} temanmu membuat CV di BoonTrack, dan kami akan buatkan **Website Portfolio Personal Gratis**!\n"
                "<i>Contoh Live:</i> https://rayigemilang.boontrack.com\n\n"
                f"👇 Bagikan link referral-mu ke teman:\n"
                f"<code>{user_ref_link}</code>"
            )
            await bot.send_message(user_id, ref_msg, reply_markup=kbd, parse_mode="HTML")

        elif code == "home_career_qa":
            user_state[user_id]["step"] = "CAREER_QA"
            qa_msg = "💬 <b>Tanya Seputar Dunia Kerja</b>\n\nKamu bisa tanyakan apa saja tentang persiapan kerja, tips interview, negosiasi gaji, atau kualifikasi posisi impianmu.\n\n<i>Ketik saja pertanyaanmu langsung di obrolan ini ya!</i> 👇"
            await bot.send_message(user_id, qa_msg, parse_mode="HTML")

        elif code in ["status_fresh", "status_exp"]:
            user_data["status_kerja"] = "Fresh Graduate" if code == "status_fresh" else "Berpengalaman"
            user_state[user_id]["step"] = "ONBOARDING_POSISI"
            await save_dropoff(user_id, 0, user_data)

            reassurance = f"Oke, {user_name}! Berarti kita punya strategi khusus untuk Fresh Graduate 👍\nNanti kita fokus menonjolkan pendidikan, project, organisasi, dan skill utama kamu.\n\n🎯 <b>Kamu saat ini ingin melamar posisi apa?</b>\n<i>(Contoh: Admin, Marketing, Software Engineer, Customer Service)</i>" if code == "status_fresh" else f"Sip, {user_name} 👍\nKita akan fokus menggali pengalaman dan pencapaian terbaikmu agar CV-nya makin menjual di mata HR.\n\n🎯 <b>Kamu saat ini ingin melamar posisi apa?</b>\n<i>(Contoh: Marketing Executive, Admin Operational, Graphic Designer)</i>"
            await bot.send_message(user_id, reassurance, parse_mode="HTML")

        elif code in ["lang_id", "lang_en", "lang_hybrid"]:
            target_lang = "ID" if code == "lang_id" else ("EN" if code == "lang_en" else "HYBRID")
            user_data["target_lang"] = target_lang

            msg_lang = "Siap! Percakapan dan CV kamu akan dibuat dalam <b>Bahasa Indonesia</b> 🇮🇩" if code == "lang_id" else ("Great! We will conduct our conversation and build your CV in <b>English</b> 🇬🇧" if code == "lang_en" else "Sip! Pilihan cerdas 🌐\nCV kamu akan dibuat dalam <b>English profesional</b>, tapi selama pengisian kamu <b>bebas cerita dalam Bahasa Indonesia</b>.\nNanti saya bantu terjemahkan dan rapikan! 😊")
            await bot.send_message(user_id, msg_lang, parse_mode="HTML")

            user_state[user_id]["step"] = 1
            await save_dropoff(user_id, 1, user_data)

            reassurance_msg = "Sip, kita mulai pelan-pelan ya 😊\n🔒 <i>Data kamu digunakan untuk membantu membuat dan menyimpan progres CV-mu. Kami tidak meminta data yang tidak diperlukan untuk proses ini.</i>\n\nKalau ada informasi yang belum kamu punya, beberapa bagian nanti bisa dilewati. Cerita saja seperti ngobrol biasa."
            await bot.send_message(user_id, reassurance_msg, parse_mode="HTML")

            status_kerja = user_data.get("status_kerja", "Berpengalaman")
            first_q = f"{get_progress_bar(1)}\n{get_question_text(1, target_lang, status_kerja)}"
            await bot.send_message(user_id, first_q, parse_mode="HTML")

        elif code == "skip_optional":
            current_step = user_state[user_id].get("step", 1)
            if isinstance(current_step, int):
                user_data[str(current_step)] = ""

                if current_step >= TOTAL_STEPS:
                    await process_and_send_cv(callback_query.message, user_id, user_data)
                else:
                    next_step = current_step + 1
                    user_state[user_id]["step"] = next_step
                    await save_dropoff(user_id, next_step, user_data)

                    target_lang = user_data.get("target_lang", "ID")
                    status_kerja = user_data.get("status_kerja", "Berpengalaman")
                    kbd = None
                    if next_step in [4, 7, 8, 9]:
                        kbd = InlineKeyboardMarkup().add(InlineKeyboardButton("⏩ Lewati Langkah Ini", callback_data="skip_optional"))

                    await bot.send_message(
                        user_id,
                        f"{get_progress_bar(next_step)}\n{get_question_text(next_step, target_lang, status_kerja)}",
                        reply_markup=kbd,
                        parse_mode="HTML"
                    )

        elif code == "resume_flow":
            state = user_state.get(user_id, {"step": 1, "data": {}})
            step = state["step"]
            target_lang = state.get("data", {}).get("target_lang", "ID")
            status_kerja = state.get("data", {}).get("status_kerja", "Berpengalaman")
            kbd = None
            if step in [4, 7, 8, 9]:
                kbd = InlineKeyboardMarkup().add(InlineKeyboardButton("⏩ Lewati Langkah Ini", callback_data="skip_optional"))
            await bot.send_message(
                user_id,
                f"Sip, mari kita lanjutkan! 👍\n\n{get_progress_bar(step)}\n{get_question_text(step, target_lang, status_kerja)}",
                reply_markup=kbd,
                parse_mode="HTML"
            )

    @dp.message_handler(content_types=['document'])
    async def handle_document_upload(message: types.Message):
        user_id = message.from_user.id
        first_name = message.from_user.first_name or "Teman"
        doc = message.document
        file_name = (doc.file_name or "").lower()

        # Log attachment
        await log_to_supabase_messages(
            sender=f"Telegram / {first_name}",
            text=f"[Mengirim Dokumen: {file_name}]",
            tenant_id="boontrack-career",
            channel="telegram",
            user_id=str(user_id),
            user_name=first_name
        )

        if not (file_name.endswith('.docx') or file_name.endswith('.pdf') or file_name.endswith('.txt')):
            await message.reply(
                "⚠️ Format file belum didukung.\n"
                "Silakan kirim dokumen CV dalam format <b>Word (.docx)</b>, <b>PDF (.pdf)</b>, atau <b>Teks (.txt)</b>.",
                parse_mode="HTML"
            )
            return

        wait_msg = await message.reply("⏳ <b>Dokumen diterima! Sedang membaca dan mengekstrak CV...</b>", parse_mode="HTML")

        try:
            import io
            from docx import Document

            file_bytes = io.BytesIO()
            await doc.download(destination_file=file_bytes)
            file_bytes.seek(0)

            extracted_text = ""

            if file_name.endswith('.docx'):
                word_doc = Document(file_bytes)
                full_text = []
                for para in word_doc.paragraphs:
                    if para.text.strip():
                        full_text.append(para.text.strip())
                for table in word_doc.tables:
                    for row in table.rows:
                        row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if row_data:
                            full_text.append(" | ".join(row_data))
                extracted_text = "\n".join(full_text)

            elif file_name.endswith('.txt'):
                extracted_text = file_bytes.read().decode('utf-8', errors='ignore')

            elif file_name.endswith('.pdf'):
                try:
                    import pypdf
                    reader = pypdf.PdfReader(file_bytes)
                    full_text = [page.extract_text() for page in reader.pages if page.extract_text()]
                    extracted_text = "\n".join(full_text)
                except Exception:
                    extracted_text = ""

            if not extracted_text or len(extracted_text.strip()) < 30:
                await bot.edit_message_text(
                    "⚠️ Teks di dalam dokumen tidak terbaca atau terlalu pendek (mungkin format gambar/scan).\n"
                    "Silakan kirim file DOCX teks atau salin langsung teks CV kamu ke chat ini.",
                    chat_id=user_id,
                    message_id=wait_msg.message_id,
                    parse_mode="HTML"
                )
                return

            try:
                await bot.edit_message_text(
                    "🤖 <b>AI sedang menganalisis skor ATS & detail perbaikan CV kamu...</b>",
                    chat_id=user_id,
                    message_id=wait_msg.message_id,
                    parse_mode="HTML"
                )
            except Exception:
                pass

            user_data = user_state.get(user_id, {}).get("data", {})
            position = str(user_data.get("target_position") or "General Professional")

            from app.handlers.commands import render_free_cv_review
            user_state[user_id]["step"] = 0
            await render_free_cv_review(user_id, bot, extracted_text, target_position=position)

            try:
                await bot.delete_message(chat_id=user_id, message_id=wait_msg.message_id)
            except Exception:
                pass

        except Exception as e:
            print(f"[Document Extraction Error]: {e}", flush=True)
            try:
                await bot.edit_message_text(
                    "❌ Terjadi kendala saat membaca file. Silakan coba lagi atau paste teks CV kamu langsung.",
                    chat_id=user_id,
                    message_id=wait_msg.message_id,
                    parse_mode="HTML"
                )
            except Exception:
                await message.reply("❌ Terjadi kendala saat membaca file. Silakan coba lagi atau paste teks CV kamu langsung.", parse_mode="HTML")

    @dp.message_handler(content_types=['photo'])
    async def handle_photo(message: types.Message):
        user_id = message.from_user.id
        first_name = message.from_user.first_name or "Teman"
        current_step = user_state.get(user_id, {}).get("step")

        await log_to_supabase_messages(
            sender=f"Telegram / {first_name}",
            text="[Mengirim Foto]",
            tenant_id="boontrack-career",
            channel="telegram",
            user_id=str(user_id),
            user_name=first_name
        )

        if current_step == "WAITING_PHOTO":
            photo = message.photo[-1]
            file_info = await bot.get_file(photo.file_id)
            photo_url = f"https://api.telegram.org/file/bot{BOT_TOKEN}/{file_info.file_path}"

            user_data = user_state[user_id].get("data", {})
            user_data["foto_url"] = photo_url
            user_state[user_id]["step"] = 0

            slug = get_user_slug(user_data, message.from_user.first_name)
            await save_dropoff(user_id, TOTAL_STEPS, user_data)
            await update_cloudflare_kv(slug, user_data)

            kbd_done = InlineKeyboardMarkup(row_width=1)
            kbd_done.add(
                InlineKeyboardButton("🌐 Buka Website Live", url=f"https://{slug}.boontrack.com"),
                InlineKeyboardButton("🔙 Kembali ke Menu Career Page", callback_data="cp_manage"),
                InlineKeyboardButton("🏠 Menu Utama", callback_data="home_back_main")
            )
            await message.reply(
                "📸 <b>Foto profil berhasil diupload & diperbarui di website!</b>\n\n"
                f"👉 <i>Cek hasilnya di:</i> https://{slug}.boontrack.com",
                reply_markup=kbd_done,
                parse_mode="HTML"
            )

    @dp.message_handler(commands=['cancel'])
    async def cancel_handler(message: types.Message):
        user_id = message.from_user.id
        first_name = message.from_user.first_name or "Teman"
        user_state[user_id] = {"step": 0, "data": {}}
        await save_dropoff(user_id, 0, {})

        await log_to_supabase_messages(
            sender=f"Telegram / {first_name}",
            text="/cancel",
            tenant_id="boontrack-career",
            channel="telegram",
            user_id=str(user_id),
            user_name=first_name
        )

        await message.reply("❌ <b>Proses pembuatan CV dibatalkan.</b>\n\nKetik /start kapan saja untuk kembali ke Menu Utama!", parse_mode="HTML")

    @dp.message_handler()
    async def handle_message(message: types.Message):
        t0 = time.perf_counter()
        user_id = message.from_user.id
        first_name = message.from_user.first_name or "Teman"
        text = (message.text or "").strip()

        current_step = user_state.get(user_id, {}).get('step', 0)

        # 1. CATAT SEMUA CHAT USER TELEGRAM KE SUPABASE SECARA INSTAN
        if text:
            await log_to_supabase_messages(
                sender=f"Telegram / {first_name}",
                text=text,
                tenant_id="boontrack-career",
                channel="telegram",
                user_id=str(user_id),
                user_name=first_name
            )

        t_db_start = time.perf_counter()
        if user_id not in user_state:
            progress, _ = await get_user_history(user_id)
            if progress and progress.get("last_step", 0) > 0:
                user_state[user_id] = {"step": progress["last_step"], "data": {}}
            else:
                user_state[user_id] = {"step": 0, "data": {}}
        t_db_end = time.perf_counter()

        if current_step == "WAITING_CV_INPUT":
            user_state[user_id]["step"] = 0
            from app.handlers.commands import render_free_cv_review
            user_data = user_state.get(user_id, {}).get("data", {})
            position = user_data.get("target_position", "General Professional")
            await render_free_cv_review(user_id, bot, text, target_position=position)
            return

        # PRIORITAS 1: ROUTING UTAMA KE AI COMPANION (QA / Chat Umum)
        if current_step == "CAREER_QA" or current_step == 0:
            words_count = len(text.strip().split())
            is_explicit_closing = any(re.search(rf"\b{re.escape(w)}\b", text.lower()) for w in CLOSING_WORDS)

            if is_explicit_closing and words_count <= 3:
                user_state[user_id]["step"] = 0
                closing_reply = "Siap! Kapan pun mau tanya lagi tinggal chat di sini."
                await message.reply(closing_reply, reply_markup=types.ReplyKeyboardRemove())
                await log_to_supabase_messages("BoonTrack AI", closing_reply, tenant_id="boontrack-career", channel="telegram", user_id=str(user_id), user_name=first_name)
                return

            asyncio.create_task(track_event(user_id, "career_ai_query", meta={"query": text}))
            await bot.send_chat_action(chat_id=user_id, action="typing")

            user_data = user_state.get(user_id, {}).get("data", {})

            t_ai_start = time.perf_counter()
            ai_reply = await ai_career_chat_response(text, user_data)
            t_ai_end = time.perf_counter()

            kbd_chat = InlineKeyboardMarkup(row_width=1)
            kbd_chat.add(
                InlineKeyboardButton("🏠 Menu Utama", callback_data="home_back_main")
            )

            t_send_start = time.perf_counter()
            await send_chunked_message(user_id, ai_reply, reply_markup=kbd_chat, parse_mode="HTML")
            t_send_end = time.perf_counter()

            # CATAT RESPON AI KE SUPABASE
            await log_to_supabase_messages(
                sender="BoonTrack AI",
                text=ai_reply,
                tenant_id="boontrack-career",
                channel="telegram",
                user_id=str(user_id),
                user_name=first_name
            )

            user_state[user_id]["step"] = "CAREER_QA"
            return

        # PRIORITAS 2: EDIT CAREER PAGE & FORM CV
        user_data = user_state.get(user_id, {}).get("data", {})
        slug = get_user_slug(user_data, message.from_user.first_name)

        if current_step == "CP_INPUT_CUSTOM_SLUG":
            clean_slug = re.sub(r'[^a-z0-9-]', '', text.lower())
            if not clean_slug or len(clean_slug) < 3:
                await message.reply("⚠️ Nama subdomain minimal 3 karakter, hanya huruf, angka, dan (-). Silakan coba lagi!")
                return

            if await check_kv_key_exists(clean_slug):
                saran_1 = f"{clean_slug}-pro"
                saran_2 = f"{clean_slug}1"
                await message.reply(
                    f"❌ <b>Subdomain Tidak Tersedia!</b>\n\n"
                    f"Subdomain <code>{clean_slug}.boontrack.com</code> sudah terdaftar oleh pengguna lain.\n\n"
                    f"💡 <b>Saran Subdomain Alternatif:</b>\n"
                    f"• <code>{saran_1}</code>\n"
                    f"• <code>{saran_2}</code>\n\n"
                    f"Silakan ketik nama subdomain/slug lain yang ingin kamu gunakan:",
                    parse_mode="HTML"
                )
                return

            user_data["custom_slug"] = clean_slug
            user_data["slug"] = clean_slug
            user_state[user_id]["data"] = user_data
            user_state[user_id]["step"] = 0

            await save_dropoff(user_id, TOTAL_STEPS, user_data)
            await update_cloudflare_kv(clean_slug, user_data)

            kbd_done = InlineKeyboardMarkup(row_width=1)
            kbd_done.add(
                InlineKeyboardButton("🌐 Buka Website Live", url=f"https://{clean_slug}.boontrack.com"),
                InlineKeyboardButton("🔙 Kembali ke Menu Career Page", callback_data="cp_manage"),
                InlineKeyboardButton("🏠 Menu Utama", callback_data="home_back_main")
            )
            await message.reply(
                f"✅ <b>Subdomain website berhasil disimpan ke:</b>\n"
                f"👉 <b>https://{clean_slug}.boontrack.com</b>",
                reply_markup=kbd_done,
                parse_mode="HTML"
            )
            return

        if current_step == "CP_EDIT_POSISI":
            user_data["target_position"] = text
            user_state[user_id]["step"] = 0
            await save_dropoff(user_id, TOTAL_STEPS, user_data)
            await update_cloudflare_kv(slug, user_data)

            kbd_done = InlineKeyboardMarkup(row_width=1)
            kbd_done.add(
                InlineKeyboardButton("🌐 Buka Website Live", url=f"https://{slug}.boontrack.com"),
                InlineKeyboardButton("🔙 Kembali ke Menu Career Page", callback_data="cp_manage"),
                InlineKeyboardButton("🏠 Menu Utama", callback_data="home_back_main")
            )
            await message.reply(f"✅ <b>Posisi berhasil diperbarui ke:</b> {text}\n\n👉 <i>Cek di:</i> https://{slug}.boontrack.com", reply_markup=kbd_done, parse_mode="HTML")
            return

        if current_step == "CP_EDIT_SUMMARY":
            user_data["ringkasan_web"] = text
            user_state[user_id]["step"] = 0
            await save_dropoff(user_id, TOTAL_STEPS, user_data)
            await update_cloudflare_kv(slug, user_data)

            kbd_done = InlineKeyboardMarkup(row_width=1)
            kbd_done.add(
                InlineKeyboardButton("🌐 Buka Website Live", url=f"https://{slug}.boontrack.com"),
                InlineKeyboardButton("🔙 Kembali ke Menu Career Page", callback_data="cp_manage"),
                InlineKeyboardButton("🏠 Menu Utama", callback_data="home_back_main")
            )
            await message.reply(f"✅ <b>Ringkasan Profil berhasil diperbarui!</b>\n\n👉 <i>Cek di:</i> https://{slug}.boontrack.com", reply_markup=kbd_done, parse_mode="HTML")
            return

        if current_step == "CP_EDIT_EXP":
            user_data["pengalaman_web"] = text
            user_state[user_id]["step"] = 0
            await save_dropoff(user_id, TOTAL_STEPS, user_data)
            await update_cloudflare_kv(slug, user_data)

            kbd_done = InlineKeyboardMarkup(row_width=1)
            kbd_done.add(
                InlineKeyboardButton("🌐 Buka Website Live", url=f"https://{slug}.boontrack.com"),
                InlineKeyboardButton("🔙 Kembali ke Menu Career Page", callback_data="cp_manage"),
                InlineKeyboardButton("🏠 Menu Utama", callback_data="home_back_main")
            )
            await message.reply(f"✅ <b>Pengalaman Kerja berhasil diperbarui!</b>\n\n👉 <i>Cek di:</i> https://{slug}.boontrack.com", reply_markup=kbd_done, parse_mode="HTML")
            return

        if current_step == "CP_EDIT_SKILLS":
            user_data["keahlian_web"] = text
            user_state[user_id]["step"] = 0
            await save_dropoff(user_id, TOTAL_STEPS, user_data)
            await update_cloudflare_kv(slug, user_data)

            kbd_done = InlineKeyboardMarkup(row_width=1)
            kbd_done.add(
                InlineKeyboardButton("🌐 Buka Website Live", url=f"https://{slug}.boontrack.com"),
                InlineKeyboardButton("🔙 Kembali ke Menu Career Page", callback_data="cp_manage"),
                InlineKeyboardButton("🏠 Menu Utama", callback_data="home_back_main")
            )
            await message.reply(f"✅ <b>Keahlian/Skill berhasil diperbarui!</b>\n\n👉 <i>Cek di:</i> https://{slug}.boontrack.com", reply_markup=kbd_done, parse_mode="HTML")
            return

        if current_step == "CP_EDIT_RESUME":
            if text.strip() == "-" or text.lower() == "kosong":
                user_data["resume_url"] = ""
            else:
                user_data["resume_url"] = text

            if user_id in user_state:
                user_state[user_id]["data"] = user_data
                user_state[user_id]["step"] = 0

                await save_dropoff(user_id, TOTAL_STEPS, user_data)
                await update_cloudflare_kv(slug, user_data)

                kbd_done = InlineKeyboardMarkup(row_width=1)
                kbd_done.add(
                    InlineKeyboardButton("🌐 Buka Website Live", url=f"https://{slug}.boontrack.com"),
                    InlineKeyboardButton("🔙 Kembali ke Menu Career Page", callback_data="cp_manage"),
                    InlineKeyboardButton("🏠 Menu Utama", callback_data="home_back_main")
                )
                await message.reply(f"✅ <b>Resume berhasil diperbarui!</b>\n\n👉 <i>Cek di:</i> https://{slug}.boontrack.com", reply_markup=kbd_done, parse_mode="HTML")
                return

        if current_step == "ONBOARDING_NAMA":
            user_data["nama_panggilan"] = text
            user_state[user_id]["step"] = "ONBOARDING_STATUS"
            await save_dropoff(user_id, 0, user_data)

            kbd_status = InlineKeyboardMarkup(row_width=1)
            kbd_status.add(
                InlineKeyboardButton("🔹 Fresh Graduate / Belum berpengalaman", callback_data="status_fresh"),
                InlineKeyboardButton("🔹 Sudah berpengalaman (Cari kerja baru)", callback_data="status_exp")
            )
            msg_2 = f"Halo {text} 😊\n\nBoleh saya tahu status kamu saat ini?"
            await message.reply(msg_2, reply_markup=kbd_status, parse_mode="HTML")
            return

        if current_step == "ONBOARDING_POSISI":
            user_data["target_position"] = text
            user_state[user_id]["step"] = "SELECT_LANGUAGE"
            await save_dropoff(user_id, 0, user_data)

            kbd_lang = InlineKeyboardMarkup(row_width=1)
            kbd_lang.add(
                InlineKeyboardButton("🌐 CV English (Ngobrol B. Indonesia)", callback_data="lang_hybrid"),
                InlineKeyboardButton("🇮🇩 CV Bahasa Indonesia", callback_data="lang_id"),
                InlineKeyboardButton("🇬🇧 Full English", callback_data="lang_en")
            )
            msg_insight = f"Oke, <b>{text}</b> 👍\n\nBerdasarkan posisi tersebut, kita akan susun CV yang menonjolkan kualifikasi yang paling dinilai rekruter.\n\nSebelum kita lanjut, CV kamu ingin dibuat dalam bahasa apa?"
            await message.reply(msg_insight, reply_markup=kbd_lang, parse_mode="HTML")
            return

        if current_step == "SELECT_LANGUAGE":
            await message.reply("Silakan <b>pilih salah satu bahasa di atas</b> ya 👆", parse_mode="HTML")
            return

        # TASK MODE: PENGISIAN CV LANGKAH BERTAHAP (1-9)
        if isinstance(current_step, int) and current_step > 0:
            target_lang = user_data.get("target_lang", "ID")
            status_kerja = user_data.get("status_kerja", "Berpengalaman")

            if current_step == 2:
                email_clean = text.strip().lower()
                if not re.match(r"^[^@]+@[^@]+\.[^@]+$", email_clean):
                    await message.reply("⚠️ <b>Format email belum sesuai.</b>\nMohon masukkan email yang valid (contoh: <code>nama@gmail.com</code>).", parse_mode="HTML")
                    return

            if current_step == 7:
                phone_digits = re.sub(r"\D", "", text)
                if len(phone_digits) < 10 or len(phone_digits) > 14:
                    kbd_skip = InlineKeyboardMarkup().add(InlineKeyboardButton("⏩ Lewati Langkah Ini", callback_data="skip_optional"))
                    await message.reply("⚠️ <b>Nomor HP/WhatsApp tidak valid.</b>\nNomor HP harus terdiri dari <b>10 sampai 14 digit</b> (contoh: <code>081234567890</code>).\n\nSilakan ketik ulang atau klik tombol di bawah untuk melewati:", reply_markup=kbd_skip, parse_mode="HTML")
                    return
                text = phone_digits

            user_data[str(current_step)] = text
            asyncio.create_task(track_event(user_id, f"step_{current_step}_completed"))

            if current_step < TOTAL_STEPS:
                next_step = current_step + 1
                user_state[user_id]["step"] = next_step
                await save_dropoff(user_id, next_step, user_data)

                kbd = None
                if next_step in [4, 7, 8, 9]:
                    kbd = InlineKeyboardMarkup().add(InlineKeyboardButton("⏩ Lewati Langkah Ini", callback_data="skip_optional"))

                step_q = f"{get_progress_bar(next_step)}\n{get_question_text(next_step, target_lang, status_kerja)}"
                await message.reply(step_q, reply_markup=kbd, parse_mode="HTML")
                await log_to_supabase_messages("BoonTrack AI", step_q, tenant_id="boontrack-career", channel="telegram", user_id=str(user_id), user_name=first_name)
            else:
                await process_and_send_cv(message, user_id, user_data)

