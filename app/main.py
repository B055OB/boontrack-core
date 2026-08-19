import sys
import os

from flask import app
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import asyncio
import json
import re
import random
import tempfile
import uuid
import time
import aiohttp
import logging
logger = logging.getLogger(__name__)
from datetime import datetime, timedelta
from typing import Optional, Dict
from dotenv import load_dotenv
import psycopg2
from psycopg2.extras import RealDictCursor
from pydantic import BaseModel, Field
from aiogram import Bot, Dispatcher, executor, types
from aiogram.utils.exceptions import ConflictError
from aiogram.types import InlineKeyboardMarkup, InlineKeyboardButton, InputFile
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from aiohttp import web
from sqlalchemy.ext.asyncio import create_async_engine, async_sessionmaker, AsyncSession

from app.repositories.session_repository import SessionRepository
from app.services.analytics_service import analytics_service
from app.services.ai_gateway import AIGateway
from app.services.brain_engine import BrainEngine
from app.handlers.admin_handler import admin_handler
from app.engines.cv_review_engine import cv_review_engine
from app.services.cv_review_service import cv_review_service
from app.routes.webchat import router as webchat_router
from app.modules.public_services.router import register_public_service_routes
from app.modules.commerce.router import commerce_routes
from app.handlers.career_page_flow import register_career_page_handlers, start_career_page_claim
from app.routes.whatsapp_career import register_whatsapp_career_routes
from app.routes.payment import register_payment_routes

# ============================================================
# B2B MULTI-TENANT ROUTERS & HANDLERS
# ============================================================
from app.telegram.router import register_telegram_routes
from app.whatsapp.router import register_whatsapp_routes
from app.reader.router import (
    pair_device_handler,
    refresh_token_handler,
    revoke_device_handler,
)

# ============================================================
# ASYNC DATABASE ENGINE & SESSION FACTORY
# ============================================================
DATABASE_URL = os.getenv("DATABASE_URL", "")
if DATABASE_URL.startswith("postgres://"):
    DATABASE_URL = DATABASE_URL.replace("postgres://", "postgresql+asyncpg://", 1)
elif DATABASE_URL.startswith("postgresql://"):
    DATABASE_URL = DATABASE_URL.replace("postgresql://", "postgresql+asyncpg://", 1)

engine = create_async_engine(DATABASE_URL, pool_pre_ping=True, echo=False)
async_session = async_sessionmaker(engine, expire_on_commit=False, class_=AsyncSession)

# --- WEB CHAT MVP SCHEMA & STATE ---
class WebChatRequest(BaseModel):
    session_id: str = Field(..., max_length=64)
    message: str = Field(..., max_length=500)
    utm_data: Optional[Dict[str, str]] = None
    click_id: Optional[str] = None

WEB_SESSION_COUNTS: Dict[str, int] = {}
MAX_WEB_MESSAGES = 7

# ==========================================
# 1. INITIALIZATION & FORMATTER
# ==========================================
load_dotenv()

# HELPER: CHUNKED MESSAGE UNTUK TELEGRAM (ANTI LIMIT 4096 KARAKTER)
async def send_chunked_message(chat_id: int, text: str, reply_markup=None, parse_mode="HTML"):
    MAX_CHUNK = 3800
    clean_text = (text or "").strip()
    
    if len(clean_text) <= MAX_CHUNK:
        await bot.send_message(chat_id, clean_text, reply_markup=reply_markup, parse_mode=parse_mode)
        return

    # Pecah berdasarkan paragraf
    lines = clean_text.split("\n")
    chunks = []
    current_chunk = ""

    for line in lines:
        if len(current_chunk) + len(line) + 1 > MAX_CHUNK:
            chunks.append(current_chunk.strip())
            current_chunk = line + "\n"
        else:
            current_chunk += line + "\n"

    if current_chunk.strip():
        chunks.append(current_chunk.strip())

    for idx, chunk in enumerate(chunks):
        is_last = (idx == len(chunks) - 1)
        k_markup = reply_markup if is_last else None
        await bot.send_message(chat_id, chunk, reply_markup=k_markup, parse_mode=parse_mode)

def format_telegram_review_response(data: dict, target_position: str) -> dict:
    scores = data.get("scores", {})
    confidence = data.get("confidence", {})
    
    msg = f"📊 <b>CV REVIEW DIAGNOSIS</b>\n"
    msg += f"🎯 Target: <b>{target_position}</b>\n\n"
    msg += f"📄 CV Quality        : <b>{scores.get('cv_quality', 0)}/100</b>\n"
    msg += f"🎯 Job Match         : <b>{scores.get('job_match', 0)}/100</b>\n"
    msg += f"💪 Evidence Strength : <b>{scores.get('evidence_strength', 0)}/100</b>\n"
    msg += f"───────────────\n"
    msg += f"📈 <b>Overall Score   : {data.get('overall_score', 0)}/100</b>\n\n"
  
    if data.get("strengths"):
        msg += "<b>💪 Kekuatan Utama:</b>\n"
        for s in data["strengths"]:
            msg += f"• {s}\n"
        msg += "\n"
        
    if data.get("weaknesses"):
        msg += "<b>⚠️ Celah Perbaikan:</b>\n"
        for w in data["weaknesses"]:
            msg += f"• {w}\n"
        msg += "\n"
        
    if data.get("action_plan"):
        msg += "<b>🎯 Prioritas Action Plan:</b>\n"
        for act in data["action_plan"][:3]:
            icon = "🔴" if act.get("priority") == "HIGH" else ("🟡" if act.get("priority") == "MEDIUM" else "🟢")
            msg += f"{icon} <b>{act.get('section')}</b>: {act.get('recommendation')}\n"
        msg += "\n"

    msg += f"🔍 <i>Confidence: {confidence.get('level', 'MEDIUM')} ({confidence.get('reason', '')})</i>\n"
    
    response = {
        "text": msg,
        "parse_mode": "HTML"
    }

    if data.get("is_locked"):
        msg += f"\n🔒 <i>{data.get('upgrade_cta')}</i>"
        response["text"] = msg
        response["reply_markup"] = {
            "inline_keyboard": [
                [{"text": "🚀 Buat Career Page Saya (Rp10.000)", "callback_data": "cp_build_now"}],
                [{"text": "🏠 Kembali ke Menu Utama", "callback_data": "home_back_main"}]
            ]
        }
        
    return response

async def handle_cv_review_process(user_id: int, target_position: str, cv_text: str, is_paid: bool = False):
    det_result = cv_review_engine.evaluate_cv(cv_text, target_position)
    prompt = cv_review_engine.build_llm_prompt(det_result, cv_text, target_position, is_paid)
    
    try:
        llm_raw_response = await ai_gateway.generate(
            user_message=prompt,
            context={"user_id": user_id, "feature": "cv_review"}
        )
        if llm_raw_response:
            llm_json = json.loads(llm_raw_response)
            det_result.update(llm_json)
    except Exception as e:
        print(f"[CV Review Engine] LLM Error / Timeout: {e}", flush=True)

    final_output = cv_review_engine.apply_access_control(det_result, is_paid)
    
    await cv_review_service.save_review(
        user_id=user_id,
        target_position=target_position,
        overall_score=final_output.get("overall_score", 0),
        quality_score=det_result["scores"]["cv_quality"],
        job_match_score=det_result["scores"]["job_match"],
        evidence_score=det_result["scores"]["evidence_strength"],
        review_json=final_output,
        confidence_level=det_result["confidence"]["level"]
    )
    
    return format_telegram_review_response(final_output, target_position)

# --- ENVIRONMENT CONFIGURATION ---
POSTGRES_HOST = os.getenv("POSTGRES_HOST", "postgres")
POSTGRES_PORT = os.getenv("POSTGRES_PORT", "5432")
POSTGRES_DB = os.getenv("POSTGRES_DB")
POSTGRES_USER = os.getenv("POSTGRES_USER")
POSTGRES_PASSWORD = os.getenv("POSTGRES_PASSWORD")

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "")
GROQ_API_KEY = os.getenv("GROQ_API_KEY", "")
OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY", "")
GITHUB_TOKEN = os.getenv("GITHUB_TOKEN", "")
BOT_TOKEN = os.getenv("TELEGRAM_BOT_TOKEN")
QRIS_IMAGE_PATH = "assets/qris.jpg"

EBOOK_FILE_ID = os.getenv("EBOOK_FILE_ID", "YOUR_TELEGRAM_EBOOK_FILE_ID")

CLOUDFLARE_ACCOUNT_ID = os.getenv("CLOUDFLARE_ACCOUNT_ID", "")
CLOUDFLARE_KV_NAMESPACE_ID = os.getenv("CLOUDFLARE_KV_NAMESPACE_ID", "")
CLOUDFLARE_API_TOKEN = os.getenv("CLOUDFLARE_API_TOKEN", "")

bot = Bot(token=BOT_TOKEN)
dp = Dispatcher(bot)

# --- INITIALIZE REPO, AI GATEWAY & BRAIN ENGINE ---
session_repo = SessionRepository()
ai_gateway = AIGateway()
brain_engine = BrainEngine(session_repo=session_repo, ai_gateway=ai_gateway)

user_state = {}
TOTAL_STEPS = 9
REQUIRED_REFERRALS = 5

CLOSING_WORDS = [
    "lanjut lagi", "nanti lanjut", "ntar lanjut", "okey", "oke deh", "okedeh",
    "bye", "dada", "dadah", "sampai jumpa", "sampe jumpa", "udah dulu", "udah ya",
    "siap makasih", "baik terima kasih", "sip makasih", "oke makasih", "ok makasih",
    "nanti aja", "ntar aja", "besok lagi", "ya makasih", "ok terima kasih", "makasih ya"
]

def get_progress_bar(step):
    return f"📍 <b>Langkah {step} dari {TOTAL_STEPS}</b>\n━━━━━━━━━━"

def get_db_connection():
    return psycopg2.connect(
        host=POSTGRES_HOST,
        port=POSTGRES_PORT,
        dbname=POSTGRES_DB,
        user=POSTGRES_USER,
        password=POSTGRES_PASSWORD
    )

def _init_db_sync():
    conn = get_db_connection()
    cur = conn.cursor()
    
    cur.execute("""
        CREATE TABLE IF NOT EXISTS users (
            telegram_id BIGINT PRIMARY KEY,
            username VARCHAR(255),
            first_name VARCHAR(255),
            last_name VARCHAR(255),
            language_code VARCHAR(10),
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        );
    """)
    cur.execute("""
        CREATE TABLE IF NOT EXISTS analytics (
            id SERIAL PRIMARY KEY,
            user_id BIGINT,
            event VARCHAR(100),
            meta JSONB,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        );
    """)
    cur.execute("""
        CREATE TABLE IF NOT EXISTS cv_documents (
            id SERIAL PRIMARY KEY,
            user_id BIGINT,
            version INT DEFAULT 1,
            position VARCHAR(255),
            data JSONB,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        );
    """)
    cur.execute("""
        CREATE TABLE IF NOT EXISTS user_progress (
            user_id BIGINT PRIMARY KEY,
            last_step INT,
            data JSONB,
            updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        );
    """)
    cur.execute("""
        CREATE TABLE IF NOT EXISTS product_orders (
            id SERIAL PRIMARY KEY,
            order_id VARCHAR(50) UNIQUE,
            telegram_id BIGINT,
            product_name VARCHAR(100),
            base_price INT,
            unique_code INT,
            total_amount INT UNIQUE,
            status VARCHAR(20) DEFAULT 'PENDING',
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            expires_at TIMESTAMP
        );
    """)
    cur.execute("""
        CREATE TABLE IF NOT EXISTS donation_sessions (
            id SERIAL PRIMARY KEY,
            donation_id VARCHAR(50) UNIQUE,
            telegram_id BIGINT,
            base_amount INT,
            unique_code INT,
            total_amount INT UNIQUE,
            status VARCHAR(20) DEFAULT 'PENDING',
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            expires_at TIMESTAMP
        );
    """)
    cur.execute("""
        CREATE TABLE IF NOT EXISTS click_logs (
            id SERIAL PRIMARY KEY,
            click_id VARCHAR(100) UNIQUE,
            source VARCHAR(50),
            utm_source VARCHAR(50),
            utm_medium VARCHAR(50),
            utm_campaign VARCHAR(100),
            utm_content VARCHAR(100),
            utm_term VARCHAR(100),
            telegram_user_id BIGINT,
            event_name VARCHAR(50) DEFAULT 'page_view',
            ip_address VARCHAR(50),
            user_agent TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        );
    """)
    cur.execute("""
        CREATE TABLE IF NOT EXISTS cv_reviews (
            id SERIAL PRIMARY KEY,
            user_id BIGINT NOT NULL,
            target_position VARCHAR(255) NOT NULL,
            cv_version INT DEFAULT 1,
            overall_score INT NOT NULL,
            quality_score INT NOT NULL,
            job_match_score INT NOT NULL,
            evidence_score INT NOT NULL,
            review_json JSONB NOT NULL,
            confidence_level VARCHAR(20) NOT NULL,
            created_at TIMESTAMPTZ DEFAULT CURRENT_TIMESTAMP
        );
    """)
    cur.execute("""
        CREATE TABLE IF NOT EXISTS ai_usage_logs (
            id SERIAL PRIMARY KEY,
            user_id BIGINT,
            provider VARCHAR(50) NOT NULL,
            feature VARCHAR(50) DEFAULT 'general',
            prompt_tokens INT DEFAULT 0,
            completion_tokens INT DEFAULT 0,
            total_tokens INT DEFAULT 0,
            status_code INT DEFAULT 200,
            is_error BOOLEAN DEFAULT FALSE,
            error_message TEXT,
            created_at TIMESTAMPTZ DEFAULT CURRENT_TIMESTAMP
        );
    """)
    conn.commit()
    cur.close()
    conn.close()

async def init_db():
    await asyncio.to_thread(_init_db_sync)

def _track_event_sync(user_id, event, meta=None):
    try:
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("INSERT INTO analytics (user_id, event, meta) VALUES (%s, %s, %s)",
                    (user_id, event, json.dumps(meta or {})))
        conn.commit()
        cur.close()
        conn.close()
    except Exception as e:
        print(f"Analytics DB Error: {e}", flush=True)

async def track_event(user_id, event, meta=None):
    await asyncio.to_thread(_track_event_sync, user_id, event, meta)

def _save_user_sync(user: types.User):
    try:
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("""
            INSERT INTO users (telegram_id, username, first_name, last_name, language_code)
            VALUES (%s, %s, %s, %s, %s)
            ON CONFLICT (telegram_id) DO UPDATE SET
                username = EXCLUDED.username,
                first_name = EXCLUDED.first_name,
                last_name = EXCLUDED.last_name;
        """, (user.id, user.username, user.first_name, user.last_name, user.language_code))
        conn.commit()
        cur.close()
        conn.close()
    except Exception as e:
        print(f"Save User DB Error: {e}", flush=True)

async def save_user(user: types.User):
    await asyncio.to_thread(_save_user_sync, user)

def _save_dropoff_sync(user_id, step, data):
    try:
        conn = get_db_connection()
        cur = conn.cursor()
        step_val = step if isinstance(step, int) else 0
        cur.execute("""
            INSERT INTO user_progress (user_id, last_step, data, updated_at)
            VALUES (%s, %s, %s, CURRENT_TIMESTAMP)
            ON CONFLICT (user_id) DO UPDATE SET
                last_step = EXCLUDED.last_step,
                data = EXCLUDED.data,
                updated_at = CURRENT_TIMESTAMP;
        """, (user_id, step_val, json.dumps(data)))
        conn.commit()
        cur.close()
        conn.close()
    except Exception as e:
        print(f"Dropoff DB Error: {e}", flush=True)

async def save_dropoff(user_id, step, data):
    await asyncio.to_thread(_save_dropoff_sync, user_id, step, data)

def _get_user_history_sync(user_id):
    try:
        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=RealDictCursor)
        cur.execute("SELECT last_step, data FROM user_progress WHERE user_id = %s", (user_id,))
        progress = cur.fetchone()
        cur.execute("SELECT version, position, data, created_at FROM cv_documents WHERE user_id = %s ORDER BY id DESC LIMIT 1", (user_id,))
        last_cv = cur.fetchone()
        cur.close()
        conn.close()
        return progress, last_cv
    except Exception as e:
        print(f"Get User History Error: {e}", flush=True)
        return None, None

async def get_user_history(user_id):
    return await asyncio.to_thread(_get_user_history_sync, user_id)

def _save_cv_version_sync(user_id, position, data):
    try:
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("SELECT COUNT(*) FROM cv_documents WHERE user_id = %s", (user_id,))
        count = cur.fetchone()[0]
        version = count + 1
        cur.execute("""
            INSERT INTO cv_documents (user_id, version, position, data)
            VALUES (%s, %s, %s, %s)
        """, (user_id, version, position, json.dumps(data)))
        conn.commit()
        cur.close()
        conn.close()
    except Exception as e:
        print(f"Save CV Version Error: {e}", flush=True)

async def save_cv_version(user_id, position, data):
    await asyncio.to_thread(_save_cv_version_sync, user_id, position, data)

def _count_referrals_sync(referrer_id):
    try:
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("""
            SELECT COUNT(DISTINCT user_id) 
            FROM analytics 
            WHERE event = 'start' AND meta->>'referrer_id' = %s
        """, (str(referrer_id),))
        count = cur.fetchone()[0]
        cur.close()
        conn.close()
        return count
    except Exception as e:
        print(f"Count Referrals Error: {e}", flush=True)
        return 0

async def count_referrals(referrer_id):
    return await asyncio.to_thread(_count_referrals_sync, referrer_id)

def _check_user_paid_sync(user_id):
    try:
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute(
            "SELECT id FROM donation_sessions WHERE telegram_id = %s AND status = 'VERIFIED' LIMIT 1;",
            (user_id,)
        )
        res = cur.fetchone()
        cur.close()
        conn.close()
        return bool(res)
    except Exception as e:
        print(f"Check User Paid Error: {e}", flush=True)
        return False

async def check_user_paid(user_id):
    return await asyncio.to_thread(_check_user_paid_sync, user_id)

def _create_order_sync(telegram_id, product_name, base_price, unique_code, total_amount):
    try:
        conn = get_db_connection()
        cur = conn.cursor()
        order_id = f"ORD-{int(datetime.now().timestamp())}"
        expires_at = datetime.now() + timedelta(minutes=15)
        cur.execute("""
            INSERT INTO product_orders (order_id, telegram_id, product_name, base_price, unique_code, total_amount, expires_at)
            VALUES (%s, %s, %s, %s, %s, %s, %s)
            ON CONFLICT (total_amount) DO UPDATE SET
                telegram_id = EXCLUDED.telegram_id,
                product_name = EXCLUDED.product_name,
                expires_at = EXCLUDED.expires_at,
                status = 'PENDING';
        """, (order_id, telegram_id, product_name, base_price, unique_code, total_amount, expires_at))
        conn.commit()
        cur.close()
        conn.close()
        return order_id
    except Exception as e:
        print(f"Create Order Error: {e}", flush=True)
        return None

async def create_order(telegram_id, product_name, base_price, unique_code, total_amount):
    return await asyncio.to_thread(_create_order_sync, telegram_id, product_name, base_price, unique_code, total_amount)

def _match_and_complete_order_sync(amount):
    try:
        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=RealDictCursor)
        cur.execute("""
            SELECT * FROM product_orders 
            WHERE total_amount = %s AND expires_at > CURRENT_TIMESTAMP
            LIMIT 1;
        """, (amount,))
        order = cur.fetchone()
        if order and order.get("status") == "PENDING":
            cur.execute("UPDATE product_orders SET status = 'PAID' WHERE id = %s;", (order["id"],))
            conn.commit()
        cur.close()
        conn.close()
        return order
    except Exception as e:
        print(f"Match Order Error: {e}", flush=True)
        return None

async def match_and_complete_order(amount):
    return await asyncio.to_thread(_match_and_complete_order_sync, amount)

def _create_donation_session_sync(telegram_id, base_amount, unique_code, total_amount):
    try:
        conn = get_db_connection()
        cur = conn.cursor()
        donation_id = f"DON-{int(datetime.now().timestamp())}"
        expires_at = datetime.now() + timedelta(minutes=15)
        cur.execute("""
            INSERT INTO donation_sessions (donation_id, telegram_id, base_amount, unique_code, total_amount, expires_at)
            VALUES (%s, %s, %s, %s, %s, %s)
            ON CONFLICT (total_amount) DO UPDATE SET
                telegram_id = EXCLUDED.telegram_id,
                base_amount = EXCLUDED.base_amount,
                expires_at = EXCLUDED.expires_at,
                status = 'PENDING';
        """, (donation_id, telegram_id, base_amount, unique_code, total_amount, expires_at))
        conn.commit()
        cur.close()
        conn.close()
        return donation_id
    except Exception as e:
        print(f"Create Donation Error: {e}", flush=True)
        return None

async def create_donation_session(telegram_id, base_amount, unique_code, total_amount):
    return await asyncio.to_thread(_create_donation_session_sync, telegram_id, base_amount, unique_code, total_amount)

def _match_and_complete_donation_sync(amount):
    try:
        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=RealDictCursor)
        cur.execute("""
            SELECT * FROM donation_sessions 
            WHERE total_amount = %s AND status = 'PENDING'
            ORDER BY created_at DESC 
            LIMIT 1;
        """, (amount,))
        donation = cur.fetchone()
        if donation:
            cur.execute("UPDATE donation_sessions SET status = 'VERIFIED' WHERE id = %s;", (donation["id"],))
            conn.commit()
        cur.close()
        conn.close()
        return donation
    except Exception as e:
        print(f"Match Donation Error: {e}", flush=True)
        return None

async def match_and_complete_donation(amount):
    return await asyncio.to_thread(_match_and_complete_donation_sync, amount)

# --- HELPER FUNCTIONS CLOUDFLARE KV & SLUG UNIK ---
async def check_kv_key_exists(slug: str) -> bool:
    if not CLOUDFLARE_API_TOKEN or not CLOUDFLARE_KV_NAMESPACE_ID or not CLOUDFLARE_ACCOUNT_ID:
        return False

    url = f"https://api.cloudflare.com/client/v4/accounts/{CLOUDFLARE_ACCOUNT_ID}/storage/kv/namespaces/{CLOUDFLARE_KV_NAMESPACE_ID}/values/{slug.lower()}"
    headers = {"Authorization": f"Bearer {CLOUDFLARE_API_TOKEN}"}

    try:
        async with aiohttp.ClientSession() as session:
            async with session.get(url, headers=headers) as resp:
                return resp.status == 200
    except Exception as e:
        print(f"[KV Check Error] {e}", flush=True)
        return False

async def generate_unique_slug(user_data: dict) -> str:
    custom_slug = user_data.get("custom_slug", "").strip().lower()
    if custom_slug:
        base_slug = re.sub(r'[^a-z0-9-]', '', custom_slug)
    else:
        raw_name = user_data.get("nama_panggilan", user_data.get("1", "user"))
        base_slug = re.sub(r'[^a-z0-9]', '', str(raw_name).lower().replace(" ", "")) or "user"

    slug = base_slug
    counter = 1

    while await check_kv_key_exists(slug):
        slug = f"{base_slug}{counter}"
        counter += 1

    return slug

def get_user_slug(user_data: dict, default_name: str = "") -> str | None:
    custom_slug = user_data.get("custom_slug", "").strip().lower()
    if custom_slug:
        return re.sub(r'[^a-z0-9-]', '', custom_slug)
    
    if user_data.get("slug"):
        return user_data.get("slug")

    if user_data.get("cp_status") != "active":
        return None

    raw_name = user_data.get("nama_panggilan", default_name or "user")
    clean_name = re.sub(r'[^a-z0-9]', '', str(raw_name).lower().replace(" ", ""))
    return clean_name or "user"

async def update_cloudflare_kv(slug: str | None, user_data: dict) -> bool:
    if not slug:
        print("[KV Info] Slug bernilai None/kosong. Skip update KV.", flush=True)
        return False

    if not CLOUDFLARE_API_TOKEN or not CLOUDFLARE_KV_NAMESPACE_ID or not CLOUDFLARE_ACCOUNT_ID:
        print("[KV Alert] Credentials Cloudflare belum lengkap di .env", flush=True)
        return False
        
    url = f"https://api.cloudflare.com/client/v4/accounts/{CLOUDFLARE_ACCOUNT_ID}/storage/kv/namespaces/{CLOUDFLARE_KV_NAMESPACE_ID}/values/{slug.lower()}"
    payload = {
        "nama": user_data.get("nama_panggilan", user_data.get("1", "Pelamar")),
        "posisi": user_data.get("target_position", "AI & Operations Workflow Optimization Specialist"),
        "email": user_data.get("2", ""),
        "telepon": user_data.get("7", ""),
        "ringkasan": user_data.get("ringkasan_web", user_data.get("3", "")),
        "pengalaman": user_data.get("pengalaman_web", user_data.get("3", "")),
        "pendidikan": user_data.get("5", ""),
        "keahlian": user_data.get("keahlian_web", user_data.get("6", "")),
        "foto": user_data.get("foto_url", ""),
        "resume_url": user_data.get("resume_url", ""),
        "theme": user_data.get("theme", "happy")
    }
    headers = {
        "Authorization": f"Bearer {CLOUDFLARE_API_TOKEN}",
        "Content-Type": "application/json"
    }
    try:
        async with aiohttp.ClientSession() as session:
            async with session.put(url, json=payload, headers=headers, timeout=aiohttp.ClientTimeout(total=5)) as resp:
                print(f"[KV Sync Status] Status: {resp.status} untuk slug: {slug.lower()}", flush=True)
                return resp.status == 200
    except Exception as e:
        print(f"[KV Sync Error] Gagal update Cloudflare KV: {e}", flush=True)
        return False

# --- HELPER FUNCTIONS AI CV GENERATOR ---
def clean_val(val):
    if not val:
        return ""
    v = str(val).strip().lower()
    if v in ["-", "skip", "lewati", "tidak ada", "ga ada", "ngga ada", "belum ada", "lupa", "kosong"]:
        return ""
    return str(val).strip()

def get_question_text(step, target_lang="ID", status_kerja="Berpengalaman"):
    is_full_en = target_lang == "EN"
    is_fresh = "fresh" in str(status_kerja).lower()
    
    if is_full_en:
        questions = {
            1: "👤 <b>What is your Full Name?</b>\n<i>(Recommended official name for the top of your CV)</i>",
            2: "📧 <b>Your Active Email?</b>\n<i>(Email regularly checked for recruiter replies)</i>",
            3: (
                "💼 <b>Your Work / Organization / Freelance Experience?</b>\n\n"
                "💬 <i>Describe it naturally, e.g.:\n"
                "'Cashier at Toko Makmur 2021-2023, then Sales Admin at PT ABC 2023-2024'\n"
                "I will format it into professional statements!</i>\n\n"
                "<i>(Type '-' or 'Fresh Grad' if none)</i>"
            ),
            4: (
                "🏆 <b>Any projects, organizations, competitions, or key accomplishments?</b>\n<i>(Be honest, I will refine the wording for you!)</i>"
                if is_fresh else
                "🏆 <b>What were your key responsibilities or accomplishments in that role?</b>\n<i>(Be honest, I will refine the wording for you!)</i>"
            ),
            5: "🎓 <b>Your Latest Education?</b>\n<i>(e.g., Bachelor of Management - Universitas Terbuka, 2023)</i>",
            6: "🛠️ <b>Your Top Skills / Expertise?</b>\n<i>(e.g., Ms. Excel, Customer Service, Canvassing, Python)</i>",
            7: "📱 <b>WhatsApp / Phone Number (Optional)</b>\n<i>Recruiters need contact info to reach you if you pass selection. You can enter your number here, or click 'Skip' to add it manually in Word later.</i>",
            8: "📍 <b>Current City of Residence?</b>\n<i>Optional if you prefer not to display your location on your CV right now.</i>",
            9: "🔗 <b>LinkedIn / Portfolio / GitHub Link?</b>\n<i>Optional. If you don't have one yet, feel free to skip!</i>"
        }
    else:
        questions = {
            1: "👤 <b>Siapa nama lengkapmu?</b>\n<i>(Nama resmi yang ingin dicantumkan di paling atas CV)</i>",
            2: "📧 <b>Email aktif yang bisa dihubungi recruiter?</b>\n<i>(contoh: nama@gmail.com)</i>",
            3: (
                "💼 <b>Pengalaman Kerja / Organisasi / Freelance terakhirmu?</b>\n\n"
                "💬 <i>Ceritakan santai saja seperti ke teman, contoh:\n"
                "'Kasir di Toko Makmur 2021-2023, lalu Admin Sales di PT ABC 2023-2024'\n"
                "Nanti saya bantu susun menjadi kalimat profesional!</i>\n\n"
                "<i>(Ketik '-' atau 'Fresh Grad' jika belum ada pengalaman)</i>"
            ),
            4: (
                "🏆 <b>Ada project, organisasi, lomba, magang, atau pencapaian yang pernah kamu lakukan?</b>\n"
                "<i>Ceritakan santai saja. Misalnya: 'Pernah bikin website untuk target kuliah' atau 'Aktif panitia kampus'.\n"
                "Kalau belum ada juga tidak masalah, tinggal tekan tombol Lewati 😊</i>"
                if is_fresh else
                "🏆 <b>Apa saja tugas utama atau pencapaianmu di pekerjaan tersebut?</b>\n"
                "<i>(Tulis apa adanya, tidak perlu dibuat-buat. Nanti saya rapikan!)</i>"
            ),
            5: "🎓 <b>Pendidikan terakhirmu?</b>\n<i>(contoh: S1 Manajemen - Universitas Terbuka, 2023)</i>",
            6: "🛠️ <b>Skill / Keahlian utama kamu?</b>\n<i>(contoh: Ms. Excel, Customer Service, Canvassing, Python)</i>",
            7: "📱 <b>Nomor WhatsApp / HP (Opsional)</b>\n<i>Rekruter butuh nomor kontak untuk menghubungi kamu jika lolos seleksi. Kamu bisa masukkan nomor HP di sini, atau tekan tombol [ ⏩ Lewati ] jika ingin mengisinya sendiri nanti di Word.</i>",
            8: "📍 <b>Kota Domisili saat ini?</b>\n<i>Tidak wajib diisi kalau kamu belum ingin mencantumkan lokasi di CV.</i>",
            9: "🔗 <b>Link LinkedIn / Portfolio / GitHub?</b>\n<i>Kalau belum punya, tidak masalah. Kamu bisa menyusul menambahkannya nanti.</i>"
        }
    return questions.get(step, "")

# --- AI CAREER COMPANION VIA BRAIN ENGINE & AI GATEWAY ---
async def ai_career_chat_response(user_query, user_context=None):
    user_context = user_context or {}
    try:
        response = await brain_engine.handle_message(
            user_message=user_query,
            context=user_context
        )
        if response:
            if isinstance(response, dict):
                return response.get("text", "")
            return str(response)
    except Exception as e:
        print(f"[BRAIN ENGINE ERROR]: {type(e).__name__}: {e}", flush=True)

    try:
        response = await ai_gateway.generate(
            user_message=user_query,
            context=user_context
        )
        if response:
            return response
    except Exception as e:
        print(f"[AI GATEWAY DIRECT ERROR]: {type(e).__name__}: {e}", flush=True)

    return "Maaf, staf kami yang menjawab untuk kebutuhan karir sedang tidak di tempat. Mungkin bisa coba lagi nanti ya 🙏"

def create_cv_docx(user_id, data):
    doc = Document()
    target_lang = data.get("target_lang", "ID")
    status_kerja = data.get("status_kerja", "Berpengalaman")
    is_en = target_lang in ["EN", "HYBRID"]
    
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    name = clean_val(data.get("1", "NAMA LENGKAP"))
    email = clean_val(data.get("2", ""))
    phone = clean_val(data.get("7", ""))
    domicile = clean_val(data.get("8", ""))
    linkedin = clean_val(data.get("9", ""))

    p_name = doc.add_paragraph()
    p_name.paragraph_format.space_after = Pt(2)
    r_name = p_name.add_run(name.upper())
    r_name.font.name = 'Calibri'
    r_name.font.size = Pt(16)
    r_name.font.bold = True
    r_name.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    contact_parts = [p for p in [email, phone, domicile, linkedin] if p]
    if contact_parts:
        p_contact = doc.add_paragraph()
        p_contact.paragraph_format.space_after = Pt(12)
        r_contact = p_contact.add_run(" | ".join(contact_parts))
        r_contact.font.name = 'Calibri'
        r_contact.font.size = Pt(10)
        r_contact.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    def add_section_header(title):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(10)
        h.paragraph_format.space_after = Pt(4)
        run = h.add_run(title.upper())
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
        
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '1F4E78')
        pBdr.append(bottom)
        h._p.get_or_add_pPr().append(pBdr)

    add_section_header("PROFESSIONAL SUMMARY" if is_en else "RINGKASAN PROFESIONAL")
    position_text = clean_val(data.get("target_position", "Profesional"))
    summary_text = f"Profesional yang berdedikasi dan berorientasi pada hasil dengan fokus pada bidang {position_text}. Memiliki kemampuan komunikasi yang baik serta siap memberikan kontribusi positif."
    p_sum = doc.add_paragraph(summary_text)
    p_sum.paragraph_format.space_after = Pt(8)
    for r in p_sum.runs:
        r.font.name = 'Calibri'
        r.font.size = Pt(10.5)

    exp = clean_val(data.get("3", ""))
    ach_raw = clean_val(data.get("4", ""))

    if exp:
        section_title = "ORGANIZATION & PROJECTS" if (is_en and "fresh" in str(status_kerja).lower()) else ("PROFESSIONAL EXPERIENCE" if is_en else "PENGALAMAN KERJA / ORGANISASI")
        add_section_header(section_title)
        raw_jobs = [j.strip() for j in re.split(r'[\n|]', exp) if j.strip()]
        
        for job_title in raw_jobs:
            if not job_title:
                continue
            
            p_job = doc.add_paragraph()
            p_job.paragraph_format.space_before = Pt(6)
            p_job.paragraph_format.space_after = Pt(2)
            r_job = p_job.add_run(job_title)
            r_job.font.name = 'Calibri'
            r_job.font.size = Pt(10.5)
            r_job.font.bold = True

            if ach_raw:
                for bullet in ach_raw.split("\n"):
                    b_text = bullet.strip().lstrip("-*• ")
                    if b_text:
                        p_b = doc.add_paragraph(style='List Bullet')
                        p_b.paragraph_format.space_after = Pt(2)
                        r_b = p_b.add_run(b_text)
                        r_b.font.name = 'Calibri'
                        r_b.font.size = Pt(10)

    edu = clean_val(data.get("5", ""))
    if edu:
        add_section_header("EDUCATION" if is_en else "PENDIDIKAN")
        p_edu = doc.add_paragraph(edu)
        p_edu.paragraph_format.space_after = Pt(8)
        for r in p_edu.runs:
            r.font.name = 'Calibri'
            r.font.size = Pt(10.5)

    skill = clean_val(data.get("6", ""))
    if skill:
        add_section_header("SKILLS & EXPERTISE" if is_en else "KEAHLIAN")
        for line in skill.split("\n"):
            line_str = line.strip()
            if line_str:
                p_skill = doc.add_paragraph(line_str)
                p_skill.paragraph_format.space_after = Pt(3)
                for r in p_skill.runs:
                    r.font.name = 'Calibri'
                    r.font.size = Pt(10)

    temp_dir = tempfile.gettempdir()
    file_path = os.path.join(temp_dir, f"CV_{user_id}.docx")
    doc.save(file_path)
    return file_path

def get_career_home_keyboard():
    kbd = types.InlineKeyboardMarkup(row_width=1)
    kbd.add(
        types.InlineKeyboardButton("📝 Buat / Edit CV Baru", callback_data="home_create_cv"),
        types.InlineKeyboardButton("🔍 Review CV Saya", callback_data="trigger_cv_review"),
        types.InlineKeyboardButton("🌐 Buat Career Page Profesional (Rp10.000)", callback_data="don_10000"),
        types.InlineKeyboardButton("📚 Ebook & Program Digital", callback_data="home_digital_products"),
        types.InlineKeyboardButton("🎁 Cek Referral Saya", callback_data="home_check_ref"),
        types.InlineKeyboardButton("💼 Tanya Seputar Dunia Kerja", callback_data="home_career_qa")
    )
    return kbd

def get_donation_options_keyboard():
    kbd = InlineKeyboardMarkup(row_width=1)
    kbd.add(
        InlineKeyboardButton("🌐 Buat Career Page Profesional (Rp10.000)", callback_data="don_10000"),
        InlineKeyboardButton("📣 Gratis via Invite 5 Teman (Referral)", callback_data="home_check_ref"),
        InlineKeyboardButton("⏩ Nanti Dulu / Cukup CV Word", callback_data="home_back_main")
    )
    return kbd

async def process_and_send_cv(message: types.Message, user_id: int, user_data: dict):
    user_state[user_id]["step"] = 0
    await save_dropoff(user_id, TOTAL_STEPS, user_data)
    
    processing_msg = await message.reply(
        "⏳ <b>Merapikan & menyusun CV kamu agar mudah dibaca sistem rekrutmen...</b>\n"
        "Mohon tunggu sekitar 15-20 detik ya!",
        parse_mode="HTML"
    )

    try:
        file_path = await asyncio.to_thread(create_cv_docx, user_id, user_data)
        position = clean_val(user_data.get("target_position", "General"))
        await save_cv_version(user_id, position, user_data)
        asyncio.create_task(track_event(user_id, "resume_generated", meta={"position": position}))

        document = InputFile(file_path)
        user_name = user_data.get("nama_panggilan", message.from_user.first_name or "Teman")

        await bot.send_document(
            chat_id=user_id,
            document=document,
            caption=f"🎉 <b>CV kamu sudah selesai, {user_name}!</b>\n\n"
                    "File dalam format Word (.docx) sudah saya kirim di atas. Bisa kamu edit kapan saja!",
            parse_mode="HTML"
        )
        
        try:
            await bot.delete_message(chat_id=user_id, message_id=processing_msg.message_id)
        except Exception:
            pass

        # CV Review Engine
        cv_text_summary = f"{user_data.get('3', '')} {user_data.get('4', '')} {user_data.get('6', '')}"
        is_paid = await check_user_paid(user_id)
        review_response = await handle_cv_review_process(user_id, position, cv_text_summary, is_paid)

        if isinstance(review_response, dict):
            await send_chunked_message(
                chat_id=user_id,
                text=review_response.get("text", ""),
                reply_markup=review_response.get("reply_markup"),
                parse_mode=review_response.get("parse_mode", "HTML")
            )
        else:
            await send_chunked_message(user_id, review_response, parse_mode="HTML")

        value_text = (
            "💡 <b>Tips Penting Sebelum Melamar:</b>\n\n"
            "1. <b>Subjek Email Jelas:</b> Gunakan format <code>[Posisi] - [Nama Kamu]</code> (Contoh: <i>Admin Operasional - Rayi Gemilang</i>)\n"
            "2. <b>Body Email Terisi:</b> Jangan biarkan pesan email kosong; sertakan Surat Lamaran/Cover Letter singkat.\n"
            "3. <b>Pencapaian Terukur:</b> Cantumkan angka atau pencapaian konkret saat wawancara nanti.\n\n"
            "CV ini sudah bisa kamu edit kapan saja di Word jika ada bagian yang ingin kamu sesuaikan kembali. 🚀"
        )
        await send_chunked_message(user_id, value_text, parse_mode="HTML")

        slug = get_user_slug(user_data, message.from_user.first_name)

        insight_text = (
            f"📊 <b>Career Insight untuk Posisi {position}:</b>\n\n"
            f"Berdasarkan data profilmu, kekuatan utamamu ada pada keahlian operasional & komunikasi. "
            f"Rekruter di bidang ini akan sangat menyukai portofolio interaktif yang bisa diakses langsung via link bio/LinkedIn.\n\n"
            f"Tampilkan CV, pengalaman, skill & portofolio kamu dalam satu halaman web profesional yang siap dibagikan ke rekruter."
        )

        if is_paid:
            kbd_paid = InlineKeyboardMarkup(row_width=1)
            kbd_paid.add(
                InlineKeyboardButton("🌐 Kelola Career Page Saya", callback_data="cp_manage"),
                InlineKeyboardButton("🔙 Kembali ke Menu Utama", callback_data="home_back_main")
            )
            monetize_text = (
                f"{insight_text}\n\n"
                f"👉 <i>Link Website Live Kamu:</i> https://{slug}.boontrack.com\n"
                f"Kamu bisa memperbarui foto, posisi, atau mengimpor data CV terbaru kapan saja!"
            )
            await send_chunked_message(user_id, monetize_text, reply_markup=kbd_paid, parse_mode="HTML")
        else:
            monetize_text = (
                f"{insight_text}\n\n"
                f"🌐 <b>Buat Career Page Profesional</b>\n"
                f"Contoh Live: <code>rayigemilang.boontrack.com</code>\n"
                f"<i>(Sekali aktivasi seumur hidup — Rp10.000)</i>"
            )
            await send_chunked_message(user_id, monetize_text, reply_markup=get_donation_options_keyboard(), parse_mode="HTML")

        if os.path.exists(file_path):
            os.remove(file_path)

        referrer_id = user_state.get(user_id, {}).get("meta", {}).get("referrer_id")
        if referrer_id:
            ref_count_referrer = await count_referrals(referrer_id)
            if ref_count_referrer >= REQUIRED_REFERRALS:
                reward_text = (
                    f"🎉 <b>SELAMAT! Target {REQUIRED_REFERRALS} Referral Kamu Tercapai!</b>\n\n"
                    f"{REQUIRED_REFERRALS} teman yang kamu ajak telah berhasil menyusun CV.\n"
                    "Kamu berhak klaim <b>Website Portfolio Personal Gratis</b>!\n\n"
                    "Ketik /claim_website untuk klaim websitemu! 🌐"
                )
                try:
                    await bot.send_message(chat_id=int(referrer_id), text=reward_text, parse_mode="HTML")
                except Exception as e:
                    print(f"Error send referral reward: {e}", flush=True)

    except Exception as e:
        print(f"Error Generate CV Flow: {e}", flush=True)
        await message.reply("❌ Terjadi kendala teknis. Silakan tekan /start untuk coba lagi!", parse_mode="HTML")

# COMMAND 
@dp.message_handler(commands=['analytics', 'admin'])
async def handle_admin_commands(message: types.Message):
    response = await admin_handler.handle_admin_command(message.from_user.id, message.text)
    await message.reply(response, parse_mode="Markdown")

@dp.message_handler(commands=['start'])
async def send_welcome(message: types.Message):
    user_id = message.from_user.id
    await save_user(message.from_user)
    
    text_parts = message.text.split()
    args = text_parts[1] if len(text_parts) > 1 else "direct"

    meta_data = {}
    if args.startswith("ref_"):
        meta_data = {"utm_source": "referral", "referrer_id": args.replace("ref_", "")}
    elif args.startswith("CLK-"):
        meta_data = {"utm_source": "click_logs", "click_id": args}
        def _link_user_attribution():
            try:
                conn = get_db_connection()
                cur = conn.cursor()
                cur.execute("""
                    UPDATE click_logs 
                    SET telegram_user_id = %s, event_name = 'start_bot'
                    WHERE click_id = %s
                """, (user_id, args))
                conn.commit()
                cur.close()
                conn.close()
            except Exception as e:
                print(f"[Attribution Error] {e}", flush=True)

        asyncio.create_task(asyncio.to_thread(_link_user_attribution))
    else:
        meta_data = {"utm_source": args}
        asyncio.create_task(analytics_service.save_user_utm(user_id, args))

    asyncio.create_task(track_event(user_id, "start", meta=meta_data))
    progress, last_cv = await get_user_history(user_id)
    saved_data = progress.get("data", {}) if progress else {}
    user_name = saved_data.get("nama_panggilan") or message.from_user.first_name or "Teman"

    if progress is not None:
        last_step = progress.get("last_step", 0)
        
        if last_step == TOTAL_STEPS or last_step == 0:
            user_state[user_id] = {"step": 0, "data": saved_data, "meta": meta_data}
            home_msg = (
                f"Halo lagi, <b>{user_name}</b>! 👋\n\n"
                "Ada yang bisa saya bantu untuk persiapan kariermu hari ini?\n\n"
                "👇 <i>Pilih opsi di bawah:</i>"
            )
            kbd = get_career_home_keyboard()
            await message.reply(home_msg, reply_markup=kbd, parse_mode="HTML")
            return

        if isinstance(last_step, int) and last_step > 0:
            user_state[user_id] = {"step": last_step, "data": saved_data, "meta": meta_data}
            kbd = InlineKeyboardMarkup(row_width=2)
            kbd.add(
                InlineKeyboardButton("▶️ Lanjutkan CV", callback_data="resume_flow"),
                InlineKeyboardButton("🔄 Mulai Baru", callback_data="restart_flow")
            )
            await message.reply(
                f"Halo lagi, <b>{user_name}</b>! 👋\n\n"
                f"Kemarin kita sempat menyusun CV sampai di <b>Langkah {last_step} dari {TOTAL_STEPS}</b>.\n\n"
                "Mau kita tuntaskan sekarang agar CV kamu siap dipakai melamar kerja?",
                reply_markup=kbd,
                parse_mode="HTML"
            )
            return

    user_state[user_id] = {"step": "ONBOARDING_NAMA", "data": {}, "meta": meta_data}
    await save_dropoff(user_id, 0, {})
    
    msg_1 = (
        "<b>Saya BoonTrack Career Assistant.</b>\n"
        "Saya akan membantu meningkatkan peluang kamu dipanggil interview.\n\n"
        "Sebelum mulai...\n"
        "Boleh kenalan dulu?\n"
        "Ini dengan siapa?"
    )
    await message.reply(msg_1, parse_mode="HTML")

@dp.callback_query_handler(lambda c: c.data == "trigger_cv_review")
async def handle_trigger_cv_review(callback_query: types.CallbackQuery):
    user_id = callback_query.from_user.id
    
    try:
        await callback_query.answer()
    except Exception:
        pass

    if user_id not in user_state:
        user_state[user_id] = {"step": 0, "data": {}}

    user_data = user_state.get(user_id, {}).get("data", {})
    cv_text_summary = " ".join([str(v) for k, v in user_data.items() if str(v).strip()]).strip()
    position = str(user_data.get("target_position") or user_data.get(6) or user_data.get("6") or "General Professional")

    # Jika user sudah punya riwayat isi CV di bot, langsung generate review
    if cv_text_summary and len(cv_text_summary) > 20:
        from app.handlers.commands import render_free_cv_review
        await render_free_cv_review(user_id, bot, cv_text_summary, target_position=position)
        return

    # Set state untuk input teks/data CV baru
    user_state[user_id]["step"] = "WAITING_CV_INPUT"

    kbd_review = types.InlineKeyboardMarkup(row_width=1)
    kbd_review.add(
        types.InlineKeyboardButton("📝 Buat CV Baru via Chat", callback_data="home_create_cv"),
        types.InlineKeyboardButton("🏠 Menu Utama", callback_data="home_back_main")
    )
    msg_prompt = (
        "🔍 <b>DIAGNOSIS & REVIEW CV GRATIS (ATS COMPLIANT)</b>\n"
        "━━━━━━━━━━━━━━━━━━━━━━━━\n"
        "Silakan kirimkan CV kamu untuk dianalisis dan dinilai oleh AI:\n\n"
        "📁 <b>Upload File:</b> Kirim file CV kamu dalam format <b>Word (.docx)</b>, <b>PDF (.pdf)</b>, atau <b>.txt</b>\n"
        "✍️ <b>Ketik / Paste:</b> Atau salin ringkasan teks CV kamu langsung ke chat ini.\n\n"
        "<i>Mau buat CV ATS baru dari nol? Klik tombol di bawah:</i>"
    )
    
    # Gunakan direct message answer agar langsung render di thread chat yang sama
    await callback_query.message.answer(msg_prompt, reply_markup=kbd_review, parse_mode="HTML")
    return

@dp.callback_query_handler(lambda c: True)
async def handle_callback_navigation(callback_query: types.CallbackQuery):
    user_id = callback_query.from_user.id
    code = callback_query.data
    
    # 1. Selalu jawab callback agar icon loading di tombol Telegram berhenti
    try:
        await callback_query.answer()
    except Exception:
        pass

    # 2. Navigasi Kembali ke Menu Utama
    if code in ["home_back_main", "main_menu", "back_to_main"]:
        user_state[user_id] = {"step": 0, "data": {}}
        from app.handlers.commands import send_welcome
        await send_welcome(callback_query.message)
        return

    # 3. Hilangkan markup lama dengan aman (abaikan jika sudah terhapus)
    try:
        await callback_query.message.edit_reply_markup(reply_markup=None)
    except Exception:
        pass

    if user_id not in user_state:
        progress, _ = await get_user_history(user_id)
        saved_data = progress.get("data", {}) if progress else {}
        user_state[user_id] = {"step": 0, "data": saved_data}
        
    user_data = user_state[user_id].get("data", {})
    user_name = user_data.get("nama_panggilan", callback_query.from_user.first_name or "Teman")
    slug = get_user_slug(user_data, callback_query.from_user.first_name)

    if code == "trigger_cv_review":
        from app.handlers.commands import render_free_cv_review
        
        position = str(user_data.get("target_position") or user_data.get(6) or user_data.get("6") or "General Professional")
        cv_text_summary = " ".join([str(v) for k, v in user_data.items() if str(v).strip()]).strip()
        
        if cv_text_summary and len(cv_text_summary) > 20:
            await render_free_cv_review(user_id, bot, cv_text_summary, target_position=position)
            return
        else:
            # Set state agar pesan berikutnya diproses sebagai input CV
            user_state[user_id]["step"] = "WAITING_CV_INPUT"
            
            kbd_review = types.InlineKeyboardMarkup(row_width=1)
            kbd_review.add(
                types.InlineKeyboardButton("📝 Buat CV Baru via Chat", callback_data="home_create_cv"),
                types.InlineKeyboardButton("🏠 Menu Utama", callback_data="home_back_main")
            )
            msg_prompt = (
                "🔍 <b>DIAGNOSIS & REVIEW CV GRATIS (ATS COMPLIANT)</b>\n"
                "━━━━━━━━━━━━━━━━━━━━━━━━\n"
                "Silakan kirimkan CV kamu untuk dianalisis dan dinilai oleh AI:\n\n"
                "📁 <b>Upload File:</b> Kirim file CV dalam format <b>Word (.docx)</b>, <b>PDF (.pdf)</b>, atau <b>.txt</b>\n"
                "✍️ <b>Ketik / Paste:</b> Atau salin ringkasan teks CV langsung ke chat ini.\n\n"
                "<i>Mau buat CV ATS baru dari nol? Klik tombol di bawah:</i>"
            )
            await bot.send_message(user_id, msg_prompt, reply_markup=kbd_review, parse_mode="HTML")
            return
    

    elif code in ["don_5000", "don_10000", "don_25000"]:
        base_amt = 5000 if code == "don_5000" else (10000 if code == "don_10000" else 25000)
        unique_code = random.randint(100, 999)
        total_amt = base_amt + unique_code
        
        await create_donation_session(user_id, base_amt, unique_code, total_amt)
        
        don_msg = (
            f"🎉 <b>Terima kasih telah memilih BoonTrack!</b>\n\n"
            f"Tinggal satu langkah lagi untuk mengaktifkan <b>Career Page Profesional</b> milikmu dan tampil lebih menonjol di mata HRD/Klien.\n\n"
            f"🌐 <b>Contoh Tampilan Career Page:</b>\n"
            f"Lihat preview tampilan Career Page yang akan kamu dapatkan di sini:\n"
            f"👉 https://rayigemilang.boontrack.com\n\n"
            f"<i>✨ Format modern, recruiter-friendly, responsif di HP/laptop, dan <b>aktif seumur hidup (sekali bayar tanpa biaya langganan)</b>.</i>\n\n"
            f"💳 <b>Rincian Pembayaran:</b>\n"
            f"• <b>Item:</b> Aktivasi Career Page Personal (Lifetime Access)\n"
            f"• <b>Transfer Tepat:</b> <code>Rp{total_amt:,}</code> <i>(Wajib transfer sesuai hingga 3 digit terakhir)</i>\n"
            f"• <b>Rincian:</b> Rp{base_amt:,} + kode verifikasi Rp{unique_code}\n"
            f"• <b>Masa Aktif Web:</b> <b>Aktif Seumur Hidup</b>\n"
            f"• <b>Batas Waktu Bayar:</b> 15 Menit\n\n"
            f"📱 <b>Panduan Bayar via QRIS (Jika Pakai 1 HP):</b>\n"
            f"1. <b>Simpan QR:</b> <b>Screenshot layar ini</b> atau unduh gambar QRIS di atas.\n"
            f"2. <b>Buka E-Wallet / Mobile Banking:</b> (BCA, Mandiri, BRI, DANA, GoPay, OVO, ShopeePay, dll).\n"
            f"3. <b>Pilih Menu QRIS / Scan:</b> Buka scanner QRIS di aplikasimu.\n"
            f"4. <b>Upload dari Galeri:</b> Klik ikon galeri/foto di menu scanner.\n"
            f"5. <b>Pilih Screenshot QR:</b> Masukkan gambar QR tadi & pastikan nominalnya tepat <b>Rp{total_amt:,}</b>.\n"
            f"6. Selesaikan pembayaran.\n\n"
            f"⏳ <i>Sistem otomatis memverifikasi pembayaran tanpa perlu kirim bukti transfer. Setelah terdeteksi, bot akan langsung mengirimkan pilihan link subdomain personalmu!</i>"
        )
        
        kbd_qris = InlineKeyboardMarkup(row_width=1)
        kbd_qris.add(
            InlineKeyboardButton("⏳ Bayar Nanti (Kembali ke Menu Utama)", callback_data="home_back_main"),
            InlineKeyboardButton("❌ Batalkan Transaksi", callback_data="cancel_checkout")
        )

        possible_qris_paths = [QRIS_IMAGE_PATH, "/app/qris.jpg", "qris.jpg"]
        found_qris = next((p for p in possible_qris_paths if os.path.exists(p)), None)
        if found_qris:
            await bot.send_photo(chat_id=user_id, photo=InputFile(found_qris))
            await send_chunked_message(user_id, don_msg, reply_markup=kbd_qris, parse_mode="HTML")
        else:
            await send_chunked_message(user_id, don_msg, reply_markup=kbd_qris, parse_mode="HTML")
            
    elif code == "cancel_checkout":
        user_state[user_id]["step"] = 0
        kbd = await get_career_home_keyboard(user_id)
        await bot.send_message(user_id, "❌ <b>Transaksi dibatalkan.</b> Kembali ke menu utama:", reply_markup=kbd, parse_mode="HTML")

    elif code in ["cp_build_now", "cp_manage"]:
        is_paid = await check_user_paid(user_id)
        if not is_paid and code == "cp_manage":
            don_msg = (
                f"🔒 <b>Website Career Page Belum Aktif</b>\n\n"
                f"Kamu perlu mengaktifkan akses Career Page terlebih dahulu (Rp10.000) untuk mengakses menu ini."
            )
            await bot.send_message(user_id, don_msg, reply_markup=get_donation_options_keyboard(), parse_mode="HTML")
            return

        user_data["cp_status"] = "active"
        default_slug = await generate_unique_slug(user_data)
        user_data["temp_slug"] = default_slug
        user_state[user_id]["data"] = user_data

        kbd_post = InlineKeyboardMarkup(row_width=1)
        kbd_post.add(
            InlineKeyboardButton(f"✅ Pakai {default_slug}.boontrack.com", callback_data="cp_confirm_default_slug"),
            InlineKeyboardButton("✏️ Ketik Nama Custom Sendiri", callback_data="cp_change_slug_start")
        )

        await bot.send_message(
            user_id,
            "🎉 <b>Akses Career Page Aktif!</b>\n\n"
            "Mari tentukan nama link subdomain untuk Career Page milikmu:\n\n"
            f"<b>Rekomendasi Subdomain:</b>\n"
            f"👉 <code>{default_slug}.boontrack.com</code>\n\n"
            "Apakah kamu mau memakai nama rekomendasi di atas, atau ingin mengetik nama custom sendiri?",
            reply_markup=kbd_post,
            parse_mode="HTML"
        )

    elif code == "cp_confirm_default_slug":
        user_data = user_state[user_id].get("data", {})
        final_slug = user_data.get("temp_slug") or await generate_unique_slug(user_data)
        user_data["slug"] = final_slug
        user_state[user_id]["data"] = user_data

        await update_cloudflare_kv(final_slug, user_data)

        await bot.send_message(
            user_id,
            f"✅ <b>Career Page Berhasil Diterbitkan!</b>\n\n"
            f"🌐 Link portofolio kamu: https://{final_slug}.boontrack.com",
            parse_mode="HTML"
        )

    elif code == "cp_change_slug_start":
        user_state[user_id]["step"] = "CP_INPUT_CUSTOM_SLUG"
        await bot.send_message(
            user_id,
            "✏️ <b>Ketik nama subdomain (slug) baru yang kamu inginkan:</b>\n\n"
            "<i>Contoh: ketik <code>alldy-pro</code> untuk mendapatkan link https://alldy-pro.boontrack.com</i>\n"
            "<i>(Hanya huruf, angka, dan tanda hubung [-])</i>",
            parse_mode="HTML"
        )

    elif code == "cp_edit_slug":
        user_state[user_id]["step"] = "CP_INPUT_CUSTOM_SLUG"
        await bot.send_message(
            user_id,
            f"🔗 <b>Ubah Subdomain / Slug Website Kamu</b>\n\n"
            f"Subdomain kamu saat ini: <code>{slug}</code> (https://{slug}.boontrack.com)\n\n"
            f"Ketik nama subdomain kustom baru yang kamu inginkan:\n"
            f"<i>(Contoh: ratuhrd, ratu-official, rayigemilang)</i>",
            parse_mode="HTML"
        )

    elif code == "cp_edit_data":
        kbd_sections = InlineKeyboardMarkup(row_width=1)
        kbd_sections.add(
            InlineKeyboardButton("💼 Edit Posisi / Headline", callback_data="cp_edit_posisi_btn"),
            InlineKeyboardButton("📝 Edit Ringkasan Profil / Bio", callback_data="cp_edit_summary_btn"),
            InlineKeyboardButton("🏢 Edit Pengalaman Kerja / Proyek", callback_data="cp_edit_exp_btn"),
            InlineKeyboardButton("🛠️ Edit Keahlian / Skill Grid", callback_data="cp_edit_skills_btn"),
            InlineKeyboardButton("🔙 Batal / Kembali ke Menu Career Page", callback_data="cp_manage")
        )
        await bot.send_message(user_id, "✏️ <b>Pilih Bagian yang Ingin Kamu Isi atau Edit:</b>", reply_markup=kbd_sections, parse_mode="HTML")

    elif code == "cp_edit_posisi_btn":
        user_state[user_id]["step"] = "CP_EDIT_POSISI"
        await bot.send_message(user_id, "💼 <b>Edit Posisi / Headline Website</b>\n\nKetik judul posisi impianmu:\n<i>(Contoh: AI & Operations Workflow Optimization Specialist)</i>", parse_mode="HTML")

    elif code == "cp_edit_summary_btn":
        user_state[user_id]["step"] = "CP_EDIT_SUMMARY"
        await bot.send_message(user_id, "📝 <b>Edit Ringkasan Profil / Bio Website</b>\n\nKetik deskripsi singkat tentang dirimu (1-3 kalimat):\n<i>(Contoh: Membantu tim operasional memangkas waktu kerja manual dengan otomatisasi sistem)</i>", parse_mode="HTML")

    elif code == "cp_edit_exp_btn":
        user_state[user_id]["step"] = "CP_EDIT_EXP"
        await bot.send_message(user_id, "🏢 <b>Edit Pengalaman Kerja / Proyek Website</b>\n\nKetik detail pengalaman kerja atau portofolio utamamu:\n<i>(Contoh: Manager HRD di PT ABC (2022-Sekarang) - Memimpin tim 10 orang & merekrut 50+ karyawan)</i>", parse_mode="HTML")

    elif code == "cp_edit_skills_btn":
        user_state[user_id]["step"] = "CP_EDIT_SKILLS"
        await bot.send_message(user_id, "🛠️ <b>Edit Keahlian / Skill Website</b>\n\nKetik skill utama dipisahkan dengan koma:\n<i>(Contoh: Python, OpenAI API, Cloudflare Workers, SQL, Recruitment)</i>", parse_mode="HTML")

    elif code == "cp_edit_resume":
        user_state[user_id]["step"] = "CP_EDIT_RESUME"
        await bot.send_message(
            user_id,
            "📄 <b>Input / Update Link Resume PDF</b>\n\n"
            "Ketik atau paste link Google Drive / tautan publik PDF resume kamu di sini:\n"
            "<b>Contoh:</b> <code>https://drive.google.com/file/d/1A2b3C.../view?usp=sharing</code>\n\n"
            "<i>Ketik '-' jika ingin menyembunyikan tombol download resume.</i>",
            parse_mode="HTML"
        )

    elif code == "cp_import_cv":
        kbd_import = InlineKeyboardMarkup(row_width=1)
        kbd_import.add(
            InlineKeyboardButton("✅ Ya, Gunakan Semua Data CV", callback_data="cp_confirm_import"),
            InlineKeyboardButton("✏️ Batal, Pilih Edit Bagian Manual", callback_data="cp_edit_data")
        )
        await bot.send_message(
            user_id,
            "⚠️ <b>Konfirmasi Impor Data CV</b>\n\n"
            "Sistem akan menyalin ringkasan profil, kontak, dan keahlian dari draf CV ke website.\n"
            "Kamu tetap bisa mengedit bagian mana saja kapan pun!",
            reply_markup=kbd_import,
            parse_mode="HTML"
        )

    elif code == "cp_confirm_import":
        user_data["ringkasan_web"] = user_data.get("3", "")
        user_data["pengalaman_web"] = user_data.get("3", "")
        user_data["keahlian_web"] = user_data.get("6", "")
        
        await save_dropoff(user_id, TOTAL_STEPS, user_data)
        await update_cloudflare_kv(slug, user_data)
        
        kbd_done = InlineKeyboardMarkup(row_width=1)
        kbd_done.add(
            InlineKeyboardButton("🌐 Buka Website Live", url=f"https://{slug}.boontrack.com"),
            InlineKeyboardButton("🔙 Kembali ke Menu Career Page", callback_data="cp_manage"),
            InlineKeyboardButton("🏠 Menu Utama", callback_data="home_back_main")
        )
        await bot.send_message(
            user_id,
            f"✅ <b>Data CV Berhasil Diimpor & Disinkronkan!</b>\n\n"
            f"Tampilan web kamu sudah terbarui secara realtime di:\n"
            f"👉 https://{slug}.boontrack.com",
            reply_markup=kbd_done,
            parse_mode="HTML"
        )

    elif code == "cp_build_later":
        kbd = await get_career_home_keyboard(user_id)
        await bot.send_message(
            user_id,
            f"Siap, {user_name}! Akses pembuatan Career Page kamu sudah tersimpan aman.\n"
            f"Kapan saja kamu siap melengkapi datanya, tinggal klik menu <b>'🌐 Kelola Career Page Saya'</b> di Menu Utama! 👍",
            reply_markup=kbd,
            parse_mode="HTML"
        )

    elif code == "cp_upload_photo":
        user_state[user_id]["step"] = "WAITING_PHOTO"
        await bot.send_message(user_id, "📸 <b>Kirimkan foto profil terbaikmu ke chat ini ya!</b>\n<i>(Disarankan foto formal/semi-formal setengah badan)</i>", parse_mode="HTML")

    elif code == "cp_choose_theme":
        kbd_theme = InlineKeyboardMarkup(row_width=2)
        kbd_theme.add(
            InlineKeyboardButton("💛 Happy Gold", callback_data="theme_happy"),
            InlineKeyboardButton("💙 Modern Blue", callback_data="theme_blue"),
            InlineKeyboardButton("🖤 Dark Minimalist", callback_data="theme_dark"),
            InlineKeyboardButton("💚 Emerald Green", callback_data="theme_emerald"),
            InlineKeyboardButton("💜 Elegant Purple", callback_data="theme_purple")
        )
        await bot.send_message(user_id, "🎨 <b>Pilih tema warna favoritmu untuk Career Page:</b>", reply_markup=kbd_theme, parse_mode="HTML")

    elif code.startswith("theme_"):
        selected_theme = code.replace("theme_", "")
        user_data["theme"] = selected_theme
        
        await save_dropoff(user_id, TOTAL_STEPS, user_data)
        await update_cloudflare_kv(slug, user_data)
        
        kbd_done = InlineKeyboardMarkup(row_width=1)
        kbd_done.add(
            InlineKeyboardButton("🌐 Buka Website Live", url=f"https://{slug}.boontrack.com"),
            InlineKeyboardButton("🔙 Kembali ke Menu Career Page", callback_data="cp_manage"),
            InlineKeyboardButton("🏠 Menu Utama", callback_data="home_back_main")
        )
        await bot.send_message(
            user_id,
            f"🎨 <b>Tema berhasil diubah ke {selected_theme.capitalize()}!</b>\n\n"
            f"Cek perubahannya secara langsung di:\n"
            f"👉 https://{slug}.boontrack.com",
            reply_markup=kbd_done,
            parse_mode="HTML"
        )

    elif code == "cp_deploy_live":
        is_success = await update_cloudflare_kv(slug, user_data)
        if is_success:
            msg = (
                f"🎉 <b>SELAMAT! Website Career Page Kamu Resmi Aktif!</b>\n\n"
                f"👉 <b>Link Web Live:</b> https://{slug}.boontrack.com\n\n"
                f"Website ini sudah siap kamu pajang di bio LinkedIn atau WhatsApp kamu! 🚀"
            )
        else:
            msg = "⚠️ Terjadi masalah sinkronisasi server KV. Pastikan konfigurasi `.env` sudah sesuai."

        kbd_done = InlineKeyboardMarkup(row_width=1)
        kbd_done.add(
            InlineKeyboardButton("🌐 Buka Website Live", url=f"https://{slug}.boontrack.com"),
            InlineKeyboardButton("🔙 Kembali ke Menu Career Page", callback_data="cp_manage")
        )
        await bot.send_message(user_id, msg, reply_markup=kbd_done, parse_mode="HTML")

    elif code == "home_digital_products":
        kbd_products = InlineKeyboardMarkup(row_width=1)
        kbd_products.add(
            InlineKeyboardButton("📘 Ebook Panduan Lolos Interview & Gaji (Rp49.000)", callback_data="buy_ebook_interview"),
            InlineKeyboardButton("🔙 Kembali ke Menu Utama", callback_data="home_back_main")
        )
        msg_catalog = (
            "🚀 <b>PROGRAM & PRODUK DIGITAL KARIR</b>\n\n"
            "Tingkatkan peluang dipanggil dan lolos kerja dengan panduan eksklusif dari BoonTrack:\n\n"
            "📖 <b>Ebook Panduan Lolos Interview & Negosiasi Gaji</b>\n"
            "• Rangkuman pertanyaan jebakan HR + cara jawabnya\n"
            "• Template surat lamaran & email melamar kerja\n"
            "• Strategi negosiasi gaji untuk Fresh Graduate & Exp\n\n"
            "👇 <i>Klik tombol di bawah untuk membeli secara otomatis:</i>"
        )
        await bot.send_message(user_id, msg_catalog, reply_markup=kbd_products, parse_mode="HTML")

    elif code == "buy_ebook_interview":
        base_price = 50000
        unique_code = random.randint(100, 999)
        total_amount = base_price - unique_code
        
        await create_order(user_id, "Ebook Interview", base_price, unique_code, total_amount)
        
        msg_checkout = (
            f"🛒 <b>CHECKOUT: Ebook Panduan Lolos Interview</b>\n\n"
            f"💵 Harga Normal: <s>Rp{base_price:,}</s>\n"
            f"🎉 <b>Total Transfer (Dapat Potongan):</b>\n"
            f"<code>{total_amount}</code> 👈 <i>(Tekan/salin angka ini)</i>\n\n"
            f"👇 <b>Cara Pembayaran:</b>\n"
            f"1. Scan QRIS di atas atau transfer via DANA Bisnis.\n"
            f"2. Masukkan nominal <b>PRESISI <code>{total_amount}</code></b> (sampai 3 digit terakhir).\n"
            f"3. Dalam 1-3 detik setelah transfer, Ebook otomatis terkirim di sini! 🚀\n\n"
            f"⏳ <i>Nominal unik ini berlaku selama 15 menit.</i>"
        )
        
        kbd_qris = InlineKeyboardMarkup(row_width=1)
        kbd_qris.add(
            InlineKeyboardButton("⏳ Bayar Nanti (Kembali ke Menu Utama)", callback_data="home_back_main"),
            InlineKeyboardButton("❌ Batalkan Transaksi", callback_data="cancel_checkout")
        )

        possible_qris_paths = [QRIS_IMAGE_PATH, "/app/qris.jpg", "qris.jpg"]
        found_qris = next((p for p in possible_qris_paths if os.path.exists(p)), None)
        if found_qris:
            await bot.send_photo(chat_id=user_id, photo=InputFile(found_qris), caption=msg_checkout, reply_markup=kbd_qris, parse_mode="HTML")
        else:
            await bot.send_message(user_id, msg_checkout, reply_markup=kbd_qris, parse_mode="HTML")

    elif code == "home_back_main":
        current_data = user_state.get(user_id, {}).get("data", {})
        user_state[user_id] = {"step": 0, "data": current_data}
        kbd = await get_career_home_keyboard(user_id)
        await bot.send_message(user_id, "👋 <b>Kembali ke Menu Utama:</b>", reply_markup=kbd, parse_mode="HTML")

    elif code == "restart_flow":
        user_state[user_id] = {"step": 0, "data": {}}
        kbd = await get_career_home_keyboard(user_id)
        await bot.send_message(user_id, "👋 <b>Menu Utama (Data Reset):</b>", reply_markup=kbd, parse_mode="HTML")

    elif code == "home_create_cv":
        old_name = user_data.get("nama_panggilan", callback_query.from_user.first_name or "")
        new_data = {"nama_panggilan": old_name} if old_name else {}
        user_state[user_id] = {"step": "ONBOARDING_NAMA", "data": new_data}
        await save_dropoff(user_id, 0, new_data)
        
        if old_name:
            msg_restart = f"Sip, {old_name}! Kita susun versi CV baru ya. 👍\n\nKamu mau tetap pakai nama panggilan <b>{old_name}</b> atau mau ganti nama baru?\n<i>(Ketik langsung nama panggilanmu di bawah untuk melanjutkan)</i>"
        else:
            msg_restart = "Sip! Kita susun versi CV baru ya. 👍\n\nBoleh kenalan dulu?\n<b>Ini dengan siapa?</b>"
        await bot.send_message(user_id, msg_restart, parse_mode="HTML")

    elif code == "home_check_ref":
        total_refs = await count_referrals(user_id)
        bot_info = await bot.get_me()
        user_ref_link = f"https://t.me/{bot_info.username}?start=ref_{user_id}"
        kbd = await get_career_home_keyboard(user_id)
        
        ref_msg = (
            "🎁 <b>REFERRAL & BONUS PORTFOLIO WEBSITE</b>\n\n"
            f"📊 <b>Progress Referral Kamu: {total_refs} / {REQUIRED_REFERRALS}</b>\n\n"
            f"Ajak {REQUIRED_REFERRALS} temanmu membuat CV di BoonTrack, dan kami akan buatkan **Website Portfolio Personal Gratis**!\n"
            "<i>Contoh Live:</i> https://rayigemilang.boontrack.com\n\n"
            f"👇 Bagikan link referral-mu ke teman:\n"
            f"<code>{user_ref_link}</code>"
        )
        await bot.send_message(user_id, ref_msg, reply_markup=kbd, parse_mode="HTML")

    elif code == "home_career_qa":
        user_state[user_id]["step"] = "CAREER_QA"
        qa_msg = "💬 <b>Tanya Seputar Dunia Kerja</b>\n\nKamu bisa tanyakan apa saja tentang persiapan kerja, tips interview, negosiasi gaji, atau kualifikasi posisi impianmu.\n\n<i>Ketik saja pertanyaanmu langsung di obrolan ini ya!</i> 👇"
        await bot.send_message(user_id, qa_msg, parse_mode="HTML")

    elif code in ["status_fresh", "status_exp"]:
        user_data["status_kerja"] = "Fresh Graduate" if code == "status_fresh" else "Berpengalaman"
        user_state[user_id]["step"] = "ONBOARDING_POSISI"
        await save_dropoff(user_id, 0, user_data)
        
        reassurance = f"Oke, {user_name}! Berarti kita punya strategi khusus untuk Fresh Graduate 👍\nNanti kita fokus menonjolkan pendidikan, project, organisasi, dan skill utama kamu.\n\n🎯 <b>Kamu saat ini ingin melamar posisi apa?</b>\n<i>(Contoh: Admin, Marketing, Software Engineer, Customer Service)</i>" if code == "status_fresh" else f"Sip, {user_name} 👍\nKita akan fokus menggali pengalaman dan pencapaian terbaikmu agar CV-nya makin menjual di mata HR.\n\n🎯 <b>Kamu saat ini ingin melamar posisi apa?</b>\n<i>(Contoh: Marketing Executive, Admin Operational, Graphic Designer)</i>"
        await bot.send_message(user_id, reassurance, parse_mode="HTML")

    elif code in ["lang_id", "lang_en", "lang_hybrid"]:
        target_lang = "ID" if code == "lang_id" else ("EN" if code == "lang_en" else "HYBRID")
        user_data["target_lang"] = target_lang
        
        msg_lang = "Siap! Percakapan dan CV kamu akan dibuat dalam <b>Bahasa Indonesia</b> 🇮🇩" if code == "lang_id" else ("Great! We will conduct our conversation and build your CV in <b>English</b> 🇬🇧" if code == "lang_en" else "Sip! Pilihan cerdas 🌐\nCV kamu akan dibuat dalam <b>English profesional</b>, tapi selama pengisian kamu <b>bebas cerita dalam Bahasa Indonesia</b>.\nNanti saya bantu terjemahkan dan rapikan! 😊")
        await bot.send_message(user_id, msg_lang, parse_mode="HTML")
        
        user_state[user_id]["step"] = 1
        await save_dropoff(user_id, 1, user_data)
        
        reassurance_msg = "Sip, kita mulai pelan-pelan ya 😊\n🔒 <i>Data kamu digunakan untuk membantu membuat dan menyimpan progres CV-mu. Kami tidak meminta data yang tidak diperlukan untuk proses ini.</i>\n\nKalau ada informasi yang belum kamu punya, beberapa bagian nanti bisa dilewati. Cerita saja seperti ngobrol biasa."
        await bot.send_message(user_id, reassurance_msg, parse_mode="HTML")
        
        status_kerja = user_data.get("status_kerja", "Berpengalaman")
        first_q = f"{get_progress_bar(1)}\n{get_question_text(1, target_lang, status_kerja)}"
        await bot.send_message(user_id, first_q, parse_mode="HTML")

    elif code == "skip_optional":
        current_step = user_state[user_id].get("step", 1)
        if isinstance(current_step, int):
            user_data[str(current_step)] = ""
            
            if current_step >= TOTAL_STEPS:
                await process_and_send_cv(callback_query.message, user_id, user_data)
            else:
                next_step = current_step + 1
                user_state[user_id]["step"] = next_step
                await save_dropoff(user_id, next_step, user_data)
                
                target_lang = user_data.get("target_lang", "ID")
                status_kerja = user_data.get("status_kerja", "Berpengalaman")
                kbd = None
                if next_step in [4, 7, 8, 9]:
                    kbd = InlineKeyboardMarkup().add(InlineKeyboardButton("⏩ Lewati Langkah Ini", callback_data="skip_optional"))
                    
                await bot.send_message(
                    user_id,
                    f"{get_progress_bar(next_step)}\n{get_question_text(next_step, target_lang, status_kerja)}",
                    reply_markup=kbd,
                    parse_mode="HTML"
                )

    elif code == "resume_flow":
        state = user_state.get(user_id, {"step": 1, "data": {}})
        step = state["step"]
        target_lang = state.get("data", {}).get("target_lang", "ID")
        status_kerja = state.get("data", {}).get("status_kerja", "Berpengalaman")
        kbd = None
        if step in [4, 7, 8, 9]:
            kbd = InlineKeyboardMarkup().add(InlineKeyboardButton("⏩ Lewati Langkah Ini", callback_data="skip_optional"))
        await bot.send_message(
            user_id,
            f"Sip, mari kita lanjutkan! 👍\n\n{get_progress_bar(step)}\n{get_question_text(step, target_lang, status_kerja)}",
            reply_markup=kbd,
            parse_mode="HTML"
        )

# PHOTO HANDLER FOR CAREER PAGE
# ============================================================
# HANDLER UPLOAD FILE DOKUMEN CV (.DOCX, .PDF, .TXT)
# ============================================================
@dp.message_handler(content_types=['document'])
async def handle_document_upload(message: types.Message):
    user_id = message.from_user.id
    doc = message.document
    file_name = (doc.file_name or "").lower()

    if not (file_name.endswith('.docx') or file_name.endswith('.pdf') or file_name.endswith('.txt')):
        await message.reply(
            "⚠️ Format file belum didukung.\n"
            "Silakan kirim dokumen CV dalam format <b>Word (.docx)</b>, <b>PDF (.pdf)</b>, atau <b>Teks (.txt)</b>.",
            parse_mode="HTML"
        )
        return

    # Notifikasi awal agar user tahu proses sedang berjalan
    wait_msg = await message.reply("⏳ <b>Dokumen diterima! Sedang membaca dan mengekstrak CV...</b>", parse_mode="HTML")

    try:
        import io
        from docx import Document

        # Download langsung ke memory buffer
        file_bytes = io.BytesIO()
        await doc.download(destination_file=file_bytes)
        file_bytes.seek(0)

        extracted_text = ""

        # Ekstraksi DOCX
        if file_name.endswith('.docx'):
            word_doc = Document(file_bytes)
            full_text = []
            for para in word_doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())
            for table in word_doc.tables:
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_data:
                        full_text.append(" | ".join(row_data))
            extracted_text = "\n".join(full_text)

        # Ekstraksi TXT
        elif file_name.endswith('.txt'):
            extracted_text = file_bytes.read().decode('utf-8', errors='ignore')

        # Ekstraksi PDF
        elif file_name.endswith('.pdf'):
            try:
                import pypdf
                reader = pypdf.PdfReader(file_bytes)
                full_text = [page.extract_text() for page in reader.pages if page.extract_text()]
                extracted_text = "\n".join(full_text)
            except Exception:
                extracted_text = ""

        if not extracted_text or len(extracted_text.strip()) < 30:
            await bot.edit_message_text(
                "⚠️ Teks di dalam dokumen tidak terbaca atau terlalu pendek (mungkin format gambar/scan).\n"
                "Silakan kirim file DOCX teks atau salin langsung teks CV kamu ke chat ini.",
                chat_id=user_id,
                message_id=wait_msg.message_id,
                parse_mode="HTML"
            )
            return

        # Update status saat AI mulai evaluasi
        try:
            await bot.edit_message_text(
                "🤖 <b>AI sedang menganalisis skor ATS & detail perbaikan CV kamu...</b>",
                chat_id=user_id,
                message_id=wait_msg.message_id,
                parse_mode="HTML"
            )
        except Exception:
            pass

        # Eksekusi Review ATS Engine
        user_data = user_state.get(user_id, {}).get("data", {})
        position = str(user_data.get("target_position") or "General Professional")
        
        from app.handlers.commands import render_free_cv_review
        user_state[user_id]["step"] = 0
        await render_free_cv_review(user_id, bot, extracted_text, target_position=position)

        # Hapus loading setelah hasil review selesai terkirim
        try:
            await bot.delete_message(chat_id=user_id, message_id=wait_msg.message_id)
        except Exception:
            pass

    except Exception as e:
        print(f"[Document Extraction Error]: {e}", flush=True)
        try:
            await bot.edit_message_text(
                "❌ Terjadi kendala saat membaca file. Silakan coba lagi atau paste teks CV kamu langsung.",
                chat_id=user_id,
                message_id=wait_msg.message_id,
                parse_mode="HTML"
            )
        except Exception:
            await message.reply("❌ Terjadi kendala saat membaca file. Silakan coba lagi atau paste teks CV kamu langsung.", parse_mode="HTML")

@dp.message_handler(content_types=['photo'])
async def handle_photo(message: types.Message):
    user_id = message.from_user.id
    current_step = user_state.get(user_id, {}).get("step")
    
    if current_step == "WAITING_PHOTO":
        photo = message.photo[-1]
        file_info = await bot.get_file(photo.file_id)
        photo_url = f"https://api.telegram.org/file/bot{BOT_TOKEN}/{file_info.file_path}"
        
        user_data = user_state[user_id].get("data", {})
        user_data["foto_url"] = photo_url
        user_state[user_id]["step"] = 0
        
        slug = get_user_slug(user_data, message.from_user.first_name)
        await save_dropoff(user_id, TOTAL_STEPS, user_data)
        await update_cloudflare_kv(slug, user_data)
        
        kbd_done = InlineKeyboardMarkup(row_width=1)
        kbd_done.add(
            InlineKeyboardButton("🌐 Buka Website Live", url=f"https://{slug}.boontrack.com"),
            InlineKeyboardButton("🔙 Kembali ke Menu Career Page", callback_data="cp_manage"),
            InlineKeyboardButton("🏠 Menu Utama", callback_data="home_back_main")
        )
        await message.reply(
            "📸 <b>Foto profil berhasil diupload & diperbarui di website!</b>\n\n"
            f"👉 <i>Cek hasilnya di:</i> https://{slug}.boontrack.com",
            reply_markup=kbd_done,
            parse_mode="HTML"
        )

@dp.message_handler(commands=['cancel'])
async def cancel_handler(message: types.Message):
    user_id = message.from_user.id
    user_state[user_id] = {"step": 0, "data": {}}
    await save_dropoff(user_id, 0, {})
    await message.reply("❌ <b>Proses pembuatan CV dibatalkan.</b>\n\nKetik /start kapan saja untuk kembali ke Menu Utama!", parse_mode="HTML")

@dp.message_handler()
async def handle_message(message: types.Message):
    t0 = time.perf_counter()
    user_id = message.from_user.id
    text = (message.text or "").strip()

    current_step = user_state.get(user_id, {}).get('step', 0)

    print(
        f"[DEBUG HANDLER] Pesan Masuk | "
        f"User: {user_id} | "
        f"Text: {text} | "
        f"Current Step: {current_step}",
        flush=True
    )

    t_db_start = time.perf_counter()
    if user_id not in user_state:
        progress, _ = await get_user_history(user_id)
        if progress and progress.get("last_step", 0) > 0:
            user_state[user_id] = {"step": progress["last_step"], "data": {}}
        else:
            user_state[user_id] = {"step": 0, "data": {}}
    t_db_end = time.perf_counter()

    # PRIORITAS 1: ROUTING UTAMA KE AI COMPANION (QA / Chat Umum)
    # PRIORITAS REVIEW CV INSTAN (DARI PASTE TEKS)
    if current_step == "WAITING_CV_INPUT":
        user_state[user_id]["step"] = 0
        from app.handlers.commands import render_free_cv_review
        user_data = user_state.get(user_id, {}).get("data", {})
        position = user_data.get("target_position", "General Professional")
        await render_free_cv_review(user_id, bot, text, target_position=position)
        return

    # PRIORITAS 1: ROUTING UTAMA KE AI COMPANION (QA / Chat Umum)
    if current_step == "CAREER_QA" or current_step == 0:
        # Cek closing words hanya jika kalimatnya sangat pendek (maksimal 3 kata)
        words_count = len(text.strip().split())
        is_explicit_closing = any(re.search(rf"\b{re.escape(w)}\b", text.lower()) for w in CLOSING_WORDS)
        
        if is_explicit_closing and words_count <= 3:
            user_state[user_id]["step"] = 0
            await message.reply("Siap! Kapan pun mau tanya lagi tinggal chat di sini.", reply_markup=types.ReplyKeyboardRemove())
            return

        # Background task agar non-blocking
        asyncio.create_task(track_event(user_id, "career_ai_query", meta={"query": text}))
        await bot.send_chat_action(chat_id=user_id, action="typing")
        
        user_data = user_state.get(user_id, {}).get("data", {})
        
        t_ai_start = time.perf_counter()
        ai_reply = await ai_career_chat_response(text, user_data)
        t_ai_end = time.perf_counter()
        
        kbd_chat = InlineKeyboardMarkup(row_width=1)
        kbd_chat.add(
            InlineKeyboardButton("🏠 Menu Utama", callback_data="home_back_main")
        )
        
        t_send_start = time.perf_counter()
        await send_chunked_message(user_id, ai_reply, reply_markup=kbd_chat, parse_mode="HTML")
        t_send_end = time.perf_counter()

        # Kunci state agar sesi chat mengalir terus secara otomatis
        user_state[user_id]["step"] = "CAREER_QA"

        # LOGGING PROFILING
        db_ms = (t_db_end - t_db_start) * 1000
        ai_ms = (t_ai_end - t_ai_start) * 1000
        send_ms = (t_send_end - t_send_start) * 1000
        total_ms = (t_send_end - t0) * 1000

        print(
            f"[PERF] User: {user_id} | "
            f"DB: {db_ms:.1f}ms | "
            f"AI Call: {ai_ms:.1f}ms | "
            f"Telegram Send: {send_ms:.1f}ms | "
            f"TOTAL: {total_ms:.1f}ms",
            flush=True
        )
        return

    # PRIORITAS 2: EDIT CAREER PAGE & FORM CV
    user_data = user_state.get(user_id, {}).get("data", {})
    slug = get_user_slug(user_data, message.from_user.first_name)

    if current_step == "CP_INPUT_CUSTOM_SLUG":
        clean_slug = re.sub(r'[^a-z0-9-]', '', text.lower())
        if not clean_slug or len(clean_slug) < 3:
            await message.reply("⚠️ Nama subdomain minimal 3 karakter, hanya huruf, angka, dan (-). Silakan coba lagi!")
            return
            
        if await check_kv_key_exists(clean_slug):
            saran_1 = f"{clean_slug}-pro"
            saran_2 = f"{clean_slug}1"
            await message.reply(
                f"❌ <b>Subdomain Tidak Tersedia!</b>\n\n"
                f"Subdomain <code>{clean_slug}.boontrack.com</code> sudah terdaftar oleh pengguna lain.\n\n"
                f"💡 <b>Saran Subdomain Alternatif:</b>\n"
                f"• <code>{saran_1}</code>\n"
                f"• <code>{saran_2}</code>\n\n"
                f"Silakan ketik nama subdomain/slug lain yang ingin kamu gunakan:",
                parse_mode="HTML"
            )
            return

        user_data["custom_slug"] = clean_slug
        user_data["slug"] = clean_slug
        user_state[user_id]["data"] = user_data
        user_state[user_id]["step"] = 0

        await save_dropoff(user_id, TOTAL_STEPS, user_data)
        await update_cloudflare_kv(clean_slug, user_data)
        
        kbd_done = InlineKeyboardMarkup(row_width=1)
        kbd_done.add(
            InlineKeyboardButton("🌐 Buka Website Live", url=f"https://{clean_slug}.boontrack.com"),
            InlineKeyboardButton("🔙 Kembali ke Menu Career Page", callback_data="cp_manage"),
            InlineKeyboardButton("🏠 Menu Utama", callback_data="home_back_main")
        )
        await message.reply(
            f"✅ <b>Subdomain website berhasil disimpan ke:</b>\n"
            f"👉 <b>https://{clean_slug}.boontrack.com</b>",
            reply_markup=kbd_done,
            parse_mode="HTML"
        )
        return

    if current_step == "CP_EDIT_POSISI":
        user_data["target_position"] = text
        user_state[user_id]["step"] = 0
        await save_dropoff(user_id, TOTAL_STEPS, user_data)
        await update_cloudflare_kv(slug, user_data)
        
        kbd_done = InlineKeyboardMarkup(row_width=1)
        kbd_done.add(
            InlineKeyboardButton("🌐 Buka Website Live", url=f"https://{slug}.boontrack.com"),
            InlineKeyboardButton("🔙 Kembali ke Menu Career Page", callback_data="cp_manage"),
            InlineKeyboardButton("🏠 Menu Utama", callback_data="home_back_main")
        )
        await message.reply(f"✅ <b>Posisi berhasil diperbarui ke:</b> {text}\n\n👉 <i>Cek di:</i> https://{slug}.boontrack.com", reply_markup=kbd_done, parse_mode="HTML")
        return

    if current_step == "CP_EDIT_SUMMARY":
        user_data["ringkasan_web"] = text
        user_state[user_id]["step"] = 0
        await save_dropoff(user_id, TOTAL_STEPS, user_data)
        await update_cloudflare_kv(slug, user_data)
        
        kbd_done = InlineKeyboardMarkup(row_width=1)
        kbd_done.add(
            InlineKeyboardButton("🌐 Buka Website Live", url=f"https://{slug}.boontrack.com"),
            InlineKeyboardButton("🔙 Kembali ke Menu Career Page", callback_data="cp_manage"),
            InlineKeyboardButton("🏠 Menu Utama", callback_data="home_back_main")
        )
        await message.reply(f"✅ <b>Ringkasan Profil berhasil diperbarui!</b>\n\n👉 <i>Cek di:</i> https://{slug}.boontrack.com", reply_markup=kbd_done, parse_mode="HTML")
        return

    if current_step == "CP_EDIT_EXP":
        user_data["pengalaman_web"] = text
        user_state[user_id]["step"] = 0
        await save_dropoff(user_id, TOTAL_STEPS, user_data)
        await update_cloudflare_kv(slug, user_data)
        
        kbd_done = InlineKeyboardMarkup(row_width=1)
        kbd_done.add(
            InlineKeyboardButton("🌐 Buka Website Live", url=f"https://{slug}.boontrack.com"),
            InlineKeyboardButton("🔙 Kembali ke Menu Career Page", callback_data="cp_manage"),
            InlineKeyboardButton("🏠 Menu Utama", callback_data="home_back_main")
        )
        await message.reply(f"✅ <b>Pengalaman Kerja berhasil diperbarui!</b>\n\n👉 <i>Cek di:</i> https://{slug}.boontrack.com", reply_markup=kbd_done, parse_mode="HTML")
        return

    if current_step == "CP_EDIT_SKILLS":
        user_data["keahlian_web"] = text
        user_state[user_id]["step"] = 0
        await save_dropoff(user_id, TOTAL_STEPS, user_data)
        await update_cloudflare_kv(slug, user_data)
        
        kbd_done = InlineKeyboardMarkup(row_width=1)
        kbd_done.add(
            InlineKeyboardButton("🌐 Buka Website Live", url=f"https://{slug}.boontrack.com"),
            InlineKeyboardButton("🔙 Kembali ke Menu Career Page", callback_data="cp_manage"),
            InlineKeyboardButton("🏠 Menu Utama", callback_data="home_back_main")
        )
        await message.reply(f"✅ <b>Keahlian/Skill berhasil diperbarui!</b>\n\n👉 <i>Cek di:</i> https://{slug}.boontrack.com", reply_markup=kbd_done, parse_mode="HTML")
        return

    if current_step == "CP_EDIT_RESUME":
        if text.strip() == "-" or text.lower() == "kosong":
            user_data["resume_url"] = ""
        else:
            user_data["resume_url"] = text

        if user_id in user_state:
            user_state[user_id]["data"] = user_data
            user_state[user_id]["step"] = 0
            
            await save_dropoff(user_id, TOTAL_STEPS, user_data)
            await update_cloudflare_kv(slug, user_data)
            
            kbd_done = InlineKeyboardMarkup(row_width=1)
            kbd_done.add(
                InlineKeyboardButton("🌐 Buka Website Live", url=f"https://{slug}.boontrack.com"),
                InlineKeyboardButton("🔙 Kembali ke Menu Career Page", callback_data="cp_manage"),
                InlineKeyboardButton("🏠 Menu Utama", callback_data="home_back_main")
            )
            await message.reply(f"✅ <b>Resume berhasil diperbarui!</b>\n\n👉 <i>Cek di:</i> https://{slug}.boontrack.com", reply_markup=kbd_done, parse_mode="HTML")
            return

    if current_step == "ONBOARDING_NAMA":
        user_data["nama_panggilan"] = text
        user_state[user_id]["step"] = "ONBOARDING_STATUS"
        await save_dropoff(user_id, 0, user_data)
        
        kbd_status = InlineKeyboardMarkup(row_width=1)
        kbd_status.add(
            InlineKeyboardButton("🔹 Fresh Graduate / Belum berpengalaman", callback_data="status_fresh"),
            InlineKeyboardButton("🔹 Sudah berpengalaman (Cari kerja baru)", callback_data="status_exp")
        )
        msg_2 = f"Halo {text} 😊\n\nBoleh saya tahu status kamu saat ini?"
        await message.reply(msg_2, reply_markup=kbd_status, parse_mode="HTML")
        return

    if current_step == "ONBOARDING_POSISI":
        user_data["target_position"] = text
        user_state[user_id]["step"] = "SELECT_LANGUAGE"
        await save_dropoff(user_id, 0, user_data)
        
        kbd_lang = InlineKeyboardMarkup(row_width=1)
        kbd_lang.add(
            InlineKeyboardButton("🌐 CV English (Ngobrol B. Indonesia)", callback_data="lang_hybrid"),
            InlineKeyboardButton("🇮🇩 CV Bahasa Indonesia", callback_data="lang_id"),
            InlineKeyboardButton("🇬🇧 Full English", callback_data="lang_en")
        )
        msg_insight = f"Oke, <b>{text}</b> 👍\n\nBerdasarkan posisi tersebut, kita akan susun CV yang menonjolkan kualifikasi yang paling dinilai rekruter.\n\nSebelum kita lanjut, CV kamu ingin dibuat dalam bahasa apa?"
        await message.reply(msg_insight, reply_markup=kbd_lang, parse_mode="HTML")
        return

    if current_step == "SELECT_LANGUAGE":
        await message.reply("Silakan <b>pilih salah satu bahasa di atas</b> ya 👆", parse_mode="HTML")
        return

    # TASK MODE: PENGISIAN CV LANGKAH BERTAHAP (1-9)
    if isinstance(current_step, int) and current_step > 0:
        target_lang = user_data.get("target_lang", "ID")
        status_kerja = user_data.get("status_kerja", "Berpengalaman")

        if current_step == 2:
            email_clean = text.strip().lower()
            if not re.match(r"^[^@]+@[^@]+\.[^@]+$", email_clean):
                await message.reply("⚠️ <b>Format email belum sesuai.</b>\nMohon masukkan email yang valid (contoh: <code>nama@gmail.com</code>).", parse_mode="HTML")
                return

        if current_step == 7:
            phone_digits = re.sub(r"\D", "", text)
            if len(phone_digits) < 10 or len(phone_digits) > 14:
                kbd_skip = InlineKeyboardMarkup().add(InlineKeyboardButton("⏩ Lewati Langkah Ini", callback_data="skip_optional"))
                await message.reply("⚠️ <b>Nomor HP/WhatsApp tidak valid.</b>\nNomor HP harus terdiri dari <b>10 sampai 14 digit</b> (contoh: <code>081234567890</code>).\n\nSilakan ketik ulang atau klik tombol di bawah untuk melewati:", reply_markup=kbd_skip, parse_mode="HTML")
                return
            text = phone_digits

        user_data[str(current_step)] = text
        asyncio.create_task(track_event(user_id, f"step_{current_step}_completed"))

        if current_step < TOTAL_STEPS:
            next_step = current_step + 1
            user_state[user_id]["step"] = next_step
            await save_dropoff(user_id, next_step, user_data)
            
            kbd = None
            if next_step in [4, 7, 8, 9]:
                kbd = InlineKeyboardMarkup().add(InlineKeyboardButton("⏩ Lewati Langkah Ini", callback_data="skip_optional"))

            await message.reply(
                f"{get_progress_bar(next_step)}\n{get_question_text(next_step, target_lang, status_kerja)}",
                reply_markup=kbd,
                parse_mode="HTML"
            )
        else:
            await process_and_send_cv(message, user_id, user_data)

# --- WEB TRACKER & DANA WEBHOOK HANDLERS ---
async def tracker_handler(request):
    try:
        source = request.match_info.get('source', 'direct')
        utm_source = request.query.get('utm_source', source)
        utm_medium = request.query.get('utm_medium', 'organic')
        utm_campaign = request.query.get('utm_campaign', 'general')
        utm_content = request.query.get('utm_content', '')
        utm_term = request.query.get('utm_term', '')
        
        ip = request.headers.get('CF-Connecting-IP') or request.remote
        user_agent = request.headers.get('User-Agent', '')
        
        click_id = f"CLK-{int(datetime.now().timestamp())}-{uuid.uuid4().hex[:6]}"

        def _log_click():
            conn = get_db_connection()
            cur = conn.cursor()
            cur.execute("""
                INSERT INTO click_logs 
                (click_id, source, utm_source, utm_medium, utm_campaign, utm_content, utm_term, event_name, ip_address, user_agent)
                VALUES (%s, %s, %s, %s, %s, %s, %s, 'click_link', %s, %s)
            """, (click_id, source, utm_source, utm_medium, utm_campaign, utm_content, utm_term, ip, user_agent))
            conn.commit()
            cur.close()
            conn.close()

        asyncio.create_task(asyncio.to_thread(_log_click))

        target_bot_url = f"https://t.me/boontrackbot?start={click_id}"
        return web.HTTPFound(location=target_bot_url)

    except Exception as e:
        print(f"[Tracker Error] {e}", flush=True)
        return web.HTTPFound(location="https://t.me/boontrackbot")

async def handle_web_chat_http(request):
    cors_headers = {
        "Access-Control-Allow-Origin": "*",
        "Access-Control-Allow-Methods": "POST, OPTIONS, GET",
        "Access-Control-Allow-Headers": "Content-Type, Authorization, X-Requested-With",
    }

    if request.method == "OPTIONS":
        return web.Response(status=200, headers=cors_headers)

    try:
        data = await request.json()
    except Exception:
        return web.json_response({"status": "error", "message": "Invalid JSON"}, status=400, headers=cors_headers)

    session_id = str(data.get("session_id", "")).strip()
    user_msg = str(data.get("message", "")).strip()
    utm_data = data.get("utm_data") or {}
    click_id = data.get("click_id")

    if not user_msg:
        return web.json_response({"status": "error", "message": "Pesan tidak boleh kosong"}, status=400, headers=cors_headers)

    current_count = WEB_SESSION_COUNTS.get(session_id, 0)

    # 1. Log Attribution jika ada UTM (Non-blocking)
    if utm_data and current_count == 0:
        def _log_utm():
            try:
                conn = get_db_connection()
                cur = conn.cursor()
                cur.execute("""
                    INSERT INTO click_logs (click_id, utm_source, utm_medium, utm_campaign, utm_content, created_at)
                    VALUES (%s, %s, %s, %s, %s, NOW())
                    ON CONFLICT (click_id) DO NOTHING
                """, (
                    click_id or session_id,
                    utm_data.get("utm_source", "web_direct"),
                    utm_data.get("utm_medium", "web_chat"),
                    utm_data.get("utm_campaign", "none"),
                    utm_data.get("utm_content", "none")
                ))
                conn.commit()
                cur.close()
                conn.close()
            except Exception as e:
                print(f"[WEB CHAT UTM LOG ERROR] {e}", flush=True)

        asyncio.create_task(asyncio.to_thread(_log_utm))

    # 2. Limit Kuota Percakapan Gratis
    if current_count >= MAX_WEB_MESSAGES:
        return web.json_response({
            "status": "limit_reached",
            "reply": "Kamu sudah mencapai batas konsultasi awal gratis di web. Mau lanjutkan konsultasi mendalam dan review CV lengkap?",
            "cta": {
                "type": "telegram",
                "label": "🚀 Lanjutkan di Telegram",
                "url": f"https://t.me/BoonTrackBot?start={click_id or session_id}"
            }
        }, headers=cors_headers)

    # 3. Panggil AI Companion dengan instruksi respons singkat
    try:
        web_context_prompt = (
            f"[Instruksi Khusus Web Chat: Berikan respons yang SANGAT SINGKAT, padat, dan to-the-point "
            f"(maksimal 2-3 kalimat pendek). Berikan 1 poin insight kunci tanpa bertele-tele.]\n\n"
            f"Pesan User: {user_msg}"
        )

        ai_reply = await ai_career_chat_response(
            user_query=web_context_prompt,
            user_context={"session_id": session_id, "source": "web_chat"}
        )
    except Exception as e:
        print(f"[WEB CHAT AI ERROR] {e}", flush=True)
        ai_reply = "Saya siap bantu carikan solusinya! Boleh ceritakan posisi apa yang ingin kamu lamar saat ini?"

    WEB_SESSION_COUNTS[session_id] = current_count + 1
    updated_count = WEB_SESSION_COUNTS[session_id]

    # Deteksi jika user bertanya tentang Telegram
    user_msg_lower = (user_msg if 'user_msg' in locals() else "").lower()
    if "telegram" in user_msg_lower or "link" in user_msg_lower:
        ai_reply = "Kamu bisa langsung lanjut konsultasi penuh dan pembuatan CV di bot resmi kami di https://t.me/boontrackbot atau klik tombol hijau di bawah ya!"
    elif updated_count in [2, 3]:
        ai_reply += "\n\n👉 *Untuk pembahasan lebih lengkap dan panduan detailnya, silakan lanjut di Telegram ya!*"

    # 4. Dynamic CTA
    cta_data = None
    if updated_count >= 3:
        cta_data = {
            "type": "telegram",
            "label": "🚀 Lanjutkan Konsultasi Penuh di Telegram",
            "url": f"https://t.me/BoonTrackBot?start={click_id or session_id}"
        }

    return web.json_response({
        "status": "success",
        "reply": ai_reply,
        "messages_used": updated_count,
        "messages_limit": MAX_WEB_MESSAGES,
        "cta": cta_data
    }, headers=cors_headers)

async def dana_webhook_handler(request):
    try:
        data = await request.json()
        print(f"[DANA RAW INCOMING]: {data}", flush=True)

        source = str(
            data.get("source", "") 
            or data.get("app", "") 
            or data.get("package_name", "") 
            or data.get("title", "") 
            or data.get("sender", "")
        ).lower()

        message = str(
            data.get("message", "") 
            or data.get("text", "") 
            or data.get("content", "") 
            or data.get("notification", "") 
            or data.get("body", "")
        )

        full_payload_str = (source + " " + message).lower()

        if "dana" not in full_payload_str:
            print(f"[DANA IGNORED] Not DANA related payload: {full_payload_str}", flush=True)
            return web.json_response({"status": "ignored", "reason": "not_dana"}, status=200)

        clean_text = message.replace(".", "").replace(",", "").replace("Rp", "Rp ").replace("rp", "Rp ")
        match = re.search(r"Rp\s*(\d+)", clean_text, re.IGNORECASE) or re.search(r"(\d{4,8})", clean_text)

        if match:
            incoming_amount = int(match.group(1))
            print(f"[DANA MATCHED AMOUNT]: Rp{incoming_amount:,}", flush=True)

            # 1. Cek Order E-book
            order = await match_and_complete_order(incoming_amount)
            if order:
                if order.get("status") == "PAID":
                    return web.json_response({"status": "already_fulfilled"}, status=200)

                buyer_id = order["telegram_id"]
                product = order["product_name"]

                caption_text = (
                    f"🎉 <b>Pembayaran Terkonfirmasi! (Rp{incoming_amount:,})</b>\n\n"
                    f"Terima kasih telah membeli <b>{product}</b>.\n"
                    f"File E-book kamu terlampir langsung di bawah ini. Selamat membaca dan sukses terus!"
                )
                try:
                    await bot.send_document(
                        chat_id=buyer_id,
                        document=EBOOK_FILE_ID,
                        caption=caption_text,
                        protect_content=True,
                        parse_mode="HTML"
                    )
                except Exception as doc_err:
                    print(f"[Document Send Error]: {doc_err}", flush=True)
                    fallback_msg = f"{caption_text}\n\n👉 Link Akses Alternative: https://cvats.boontrack.com/ebook-interview-boontrack.pdf"
                    await bot.send_message(chat_id=buyer_id, text=fallback_msg, parse_mode="HTML")

                return web.json_response({"status": "success_order", "order_id": order["order_id"]}, status=200)

            # 2. Cek Donasi / Aktivasi Career Page
            donation = await match_and_complete_donation(incoming_amount)
            if donation:
                if donation.get("status") == "VERIFIED":
                    return web.json_response({"status": "already_verified"}, status=200)

                donor_id = donation["telegram_id"]
                
                if donor_id not in user_state:
                    user_state[donor_id] = {"data": {}}
                    
                user_data = user_state[donor_id].setdefault("data", {})
                user_data["cp_status"] = "active"

                default_slug = await generate_unique_slug(user_data)
                user_data["temp_slug"] = default_slug
                user_state[donor_id]["data"] = user_data

                kbd_post = InlineKeyboardMarkup(row_width=1)
                kbd_post.add(
                    InlineKeyboardButton(f"✅ Pakai {default_slug}.boontrack.com", callback_data="cp_confirm_default_slug"),
                    InlineKeyboardButton("✏️ Ketik Nama Custom Sendiri", callback_data="cp_change_slug_start")
                )

                don_thanks = (
                    f"🎉 <b>PEMBAYARAN CAREER PAGE TERKONFIRMASI!</b>\n\n"
                    f"Terima kasih atas dukunganmu sebesar <b>Rp{incoming_amount:,}</b>! 🙏\n\n"
                    f"Sekarang, silakan tentukan nama link subdomain untuk Career Page milikmu:\n\n"
                    f"<b>Rekomendasi Subdomain:</b>\n"
                    f"👉 <code>{default_slug}.boontrack.com</code>\n\n"
                    f"Apakah kamu mau memakai nama rekomendasi di atas, atau ingin mengetik nama custom sendiri?"
                )
                await bot.send_message(chat_id=donor_id, text=don_thanks, reply_markup=kbd_post, parse_mode="HTML")
                return web.json_response({"status": "success_donation", "donation_id": donation["donation_id"]}, status=200)

            print(f"[DANA WARNING] Nominal Rp{incoming_amount:,} tidak cocok dengan order/donasi pending manapun.", flush=True)
            return web.json_response({"status": "no_matching_transaction", "amount": incoming_amount}, status=200)

        print(f"[DANA PARSE ERROR] Gagal mengekstrak nominal dari teks: {message}", flush=True)
        return web.json_response({"status": "failed_parsing", "message": message}, status=200)

    except Exception as e:
        print(f"[Webhook Exception]: {e}", flush=True)
        return web.json_response({"status": "error", "detail": str(e)}, status=500)

async def health_check_handler(request):
    return web.json_response({"status": "healthy", "message": "Render is awake!"}, status=200)

async def funnel_report_handler(request):
    try:
        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=RealDictCursor)
        
        query = """
            SELECT 
                LOWER(c.source) as channel,
                COUNT(DISTINCT c.click_id) as total_klik,
                COUNT(DISTINCT CASE WHEN c.telegram_user_id IS NOT NULL THEN c.telegram_user_id END) as masuk_telegram,
                COUNT(DISTINCT CASE WHEN a.event = 'resume_generated' THEN a.user_id END) as selesai_cv,
                COUNT(DISTINCT CASE WHEN p.status = 'PAID' THEN p.telegram_id END) as purchase
            FROM click_logs c
            LEFT JOIN analytics a ON c.telegram_user_id = a.user_id
            LEFT JOIN product_orders p ON c.telegram_user_id = p.telegram_id
            GROUP BY LOWER(c.source);
        """
        cur.execute(query)
        rows = cur.fetchall()
        cur.close()
        conn.close()

        return web.json_response({"status": "success", "data": rows})
    except Exception as e:
        return web.json_response({"status": "error", "message": str(e)}, status=500)

from app.services.webchat_service import WebChatService
from app.services.lead_service import LeadService
from app.services.ai_gateway import AIGateway

# Inisialisasi Service Khusus B2B
_b2b_ai_gateway = AIGateway()
_b2b_lead_service = LeadService(ai_gateway=_b2b_ai_gateway)
_b2b_webchat_service = WebChatService(brain_engine=brain_engine, lead_service=_b2b_lead_service)

async def handle_b2b_webchat_http(request):
    try:
        data = await request.json()
        session_id = data.get("session_id", "default_session")
        message = data.get("message", "")

        if not message:
            return web.json_response({"error": "Message cannot be empty"}, status=400)

        result = await _b2b_webchat_service.process_business_chat(
            session_id=session_id,
            message=message
        )
        return web.json_response({
            "status": "success",
            "session_id": session_id,
            "reply": result["reply"],
            "is_lead_qualified": result["is_lead_qualified"]
        })
    except Exception as e:
        return web.json_response({"status": "error", "message": str(e)}, status=500)
    
# Tambahkan middleware CORS untuk aiohttp
@web.middleware
async def cors_middleware(request, handler):
    if request.method == "OPTIONS":
        return web.Response(
            status=200,
            headers={
                "Access-Control-Allow-Origin": "*",
                "Access-Control-Allow-Methods": "POST, GET, OPTIONS",
                "Access-Control-Allow-Headers": "Content-Type, Authorization, X-Requested-With",
            },
        )
    
    try:
        response = await handler(request)
    except web.HTTPException as ex:
        response = ex
    except Exception as e:
        response = web.json_response(
            {"status": "error", "message": str(e)},
            status=500
        )

    response.headers["Access-Control-Allow-Origin"] = "*"
    response.headers["Access-Control-Allow-Methods"] = "POST, GET, OPTIONS"
    response.headers["Access-Control-Allow-Headers"] = "Content-Type, Authorization, X-Requested-With"
    return response

async def start_web_server():
    app = web.Application(middlewares=[cors_middleware])
    app.router.add_get('/', health_check_handler)
    app.router.add_get('/health', health_check_handler)
    app.router.add_get('/source', tracker_handler)
    app.router.add_post('/webhook/dana', dana_webhook_handler)
    app.add_routes(commerce_routes)
    
    # Endpoint Karir (Support variasi dengan strip & tanpa strip)
    app.router.add_post('/api/webchat', handle_web_chat_http)
    app.router.add_post('/api/web-chat', handle_web_chat_http)
    
    # Endpoint B2B Business (Support variasi path)
    app.router.add_post('/api/webchat/business', handle_b2b_webchat_http)
    app.router.add_post('/api/b2b-webchat', handle_b2b_webchat_http)

    # Registrasi Seluruh Route Public Services (Webchat & WhatsApp)
    register_public_service_routes(app)

    # Registrasi Seluruh Route Multi-Tenant B2B & Device Auth Engine
    register_telegram_routes(app, async_session)
    register_whatsapp_routes(app, async_session)

    register_whatsapp_career_routes(app)
    register_payment_routes(app)

    # Device Lifecycle Endpoints (Android BoonTrack Reader)
    async def _wrap_pair(req):
        async with async_session() as session:
            return await pair_device_handler(req, session)

    async def _wrap_refresh(req):
        async with async_session() as session:
            return await refresh_token_handler(req, session)

    async def _wrap_revoke(req):
        async with async_session() as session:
            return await revoke_device_handler(req, session)

    app.router.add_post("/api/v1/devices/pair", _wrap_pair)
    app.router.add_post("/api/v1/devices/refresh", _wrap_refresh)
    app.router.add_post("/api/v1/devices/revoke", _wrap_revoke)

    port = int(os.getenv("PORT", 8080))
    runner = web.AppRunner(app)
    await runner.setup()
    
    site = web.TCPSite(runner, '0.0.0.0', port)
    await site.start()
    print(f"[BOOT] Web server listening on port {port}", flush=True)


if __name__ == '__main__':
    from aiogram.utils.exceptions import ConflictError
    
    # Daftarkan handler Career Page di sini
    register_career_page_handlers(dp)
    
    # ... baris executor.start_polling(dp, ...) bawaan kamu ...
    bot_token = os.getenv('TELEGRAM_BOT_TOKEN')

    print("========================================", flush=True)
    print("[BOOT] BoonTrack Telegram Bot STARTING", flush=True)
    print(f"[BOOT] PID          : {os.getpid()}", flush=True)
    print(f"[BOOT] HOSTNAME     : {os.getenv('HOSTNAME', 'unknown')}", flush=True)
    print(f"[BOOT] PORT         : {os.getenv('PORT', 'unknown')}", flush=True)
    print(f"[BOOT] TOKEN STATUS : {'TERBACA OK' if bot_token else 'KOSONG / UNDEFINED'}", flush=True)
    print("========================================", flush=True)

    loop = asyncio.new_event_loop()
    asyncio.set_event_loop(loop)

    print("[BOOT] Initializing database...", flush=True)
    loop.run_until_complete(init_db())

    # 1. Start Web Server di Main Loop
    print("[BOOT] Starting Web Server...", flush=True)
    loop.create_task(start_web_server())

    # 2. Start Telegram Polling langsung di loop yang sama
    async def start_telegram_polling():
        print("[TELEGRAM] Polling worker starting...", flush=True)
        try:
            await bot.delete_webhook(drop_pending_updates=True)
            await dp.start_polling(reset_webhook=True)
        except Exception as e:
            print(f"[TELEGRAM] ⚠️ Polling stopped ({e}). Web Server TETAP AKTIF.", flush=True)

    loop.create_task(start_telegram_polling())
    print("[BOOT] Telegram & Web Server running concurrently in main loop.", flush=True)

    # 3. Kunci Main Loop agar Web Server jalan terus
    try:
        loop.run_forever()
    except KeyboardInterrupt:
        pass
    finally:
        loop.stop()
        loop.close()