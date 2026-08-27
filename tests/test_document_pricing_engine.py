import io
import unittest
from unittest.mock import patch, MagicMock, AsyncMock
from docx import Document

from app.services.pricing_engine import (
    calculate_document_metrics,
    calculate_pricing,
    build_qris_invoice_payload,
    PRICE_POLISH_TIER_1,
    PRICE_POLISH_TIER_2,
    PRICE_POLISH_TIER_3,
    PRICE_POLISH_ADDON_RATE
)
from app.services.doc_builder import (
    render_ats_review_docx,
    render_cv_rewrite_docx,
    render_paraphrase_docx,
    build_document_result
)
from app.services.document_engine import (
    validate_document_file,
    intake_document_job,
    process_document_job_async,
    execute_ai_document_task,
    TASK_ATS_REVIEW,
    TASK_CV_REWRITE,
    TASK_PARAPHRASE,
    MAGIC_BYTES_PDF,
    MAGIC_BYTES_ZIP
)
from app.services.r2_storage_service import r2_storage_service


class TestDocumentPricingEngine(unittest.IsolatedAsyncioTestCase):

    # ==========================================
    # 1. Tests for Word Counter & Document Metrics
    # ==========================================
    def test_calculate_document_metrics_empty(self):
        metrics = calculate_document_metrics("")
        self.assertEqual(metrics["word_count"], 0)
        self.assertEqual(metrics["char_count"], 0)
        self.assertEqual(metrics["estimated_pages"], 0)

    def test_calculate_document_metrics_sample_text(self):
        sample_text = (
            "Halo nama saya Ahmad Dani. Saya adalah software engineer dengan pengalaman "
            "5 tahun membangun aplikasi berbasis microservices dan cloud native di Jakarta."
        )
        metrics = calculate_document_metrics(sample_text)
        self.assertGreater(metrics["word_count"], 15)
        self.assertEqual(metrics["estimated_pages"], 1)
        self.assertGreater(metrics["char_count"], len(sample_text) - 5)

    def test_calculate_document_metrics_large_pages(self):
        # 750 words -> should be ceil(750/250) = 3 pages
        long_text = " ".join(["kata"] * 750)
        metrics = calculate_document_metrics(long_text)
        self.assertEqual(metrics["word_count"], 750)
        self.assertEqual(metrics["estimated_pages"], 3)

    # ==========================================
    # 2. Tests for Dynamic Pricing Matrix Tiers
    # ==========================================
    def test_pricing_tier_1_under_500_words(self):
        # Tier 1: < 500 kata -> Rp5.000
        p_350 = calculate_pricing(TASK_PARAPHRASE, 350)
        self.assertEqual(p_350["pricing_tier"], "TIER_1")
        self.assertEqual(p_350["final_price"], 5000)

    def test_pricing_tier_2_between_500_and_2500_words(self):
        # Tier 2: 500 - 2500 kata -> Rp10.000
        p_1500 = calculate_pricing(TASK_PARAPHRASE, 1500)
        self.assertEqual(p_1500["pricing_tier"], "TIER_2")
        self.assertEqual(p_1500["final_price"], 10000)

    def test_pricing_tier_3_between_2500_and_6000_words(self):
        # Tier 3: 2500 - 6000 kata -> Rp20.000
        p_4000 = calculate_pricing(TASK_PARAPHRASE, 4000)
        self.assertEqual(p_4000["pricing_tier"], "TIER_3")
        self.assertEqual(p_4000["final_price"], 20000)

    def test_pricing_tier_4_above_6000_words(self):
        # Tier 4: > 6000 kata -> Rp40.000 (+Rp5.000 per 2000 kata jika > 12000 kata)
        p_8000 = calculate_pricing(TASK_PARAPHRASE, 8000)
        self.assertEqual(p_8000["pricing_tier"], "TIER_4")
        self.assertEqual(p_8000["final_price"], 40000)

        # 14000 words -> 40000 + (1 * 5000) = 45000
        p_14000 = calculate_pricing(TASK_PARAPHRASE, 14000)
        self.assertEqual(p_14000["pricing_tier"], "TIER_4")
        self.assertEqual(p_14000["final_price"], 45000)

    def test_build_qris_invoice_payload(self):
        inv = build_qris_invoice_payload("job-abc", TASK_PARAPHRASE, 1500, "628123456789")
        self.assertEqual(inv["job_id"], "job-abc")
        self.assertEqual(inv["amount"], 10000)
        self.assertEqual(inv["currency"], "IDR")
        self.assertEqual(inv["pricing_tier"], "TIER_2")
        self.assertEqual(inv["formatted_amount"], "Rp10.000")

    # ==========================================
    # 3. Tests for MIME & Magic Bytes Validation
    # ==========================================
    def test_validate_document_file_pdf(self):
        fake_pdf_bytes = MAGIC_BYTES_PDF + b"-1.7\n%fake pdf content"
        is_valid, mime, err = validate_document_file(fake_pdf_bytes, "cv_dani.pdf")
        self.assertTrue(is_valid)
        self.assertEqual(mime, "application/pdf")
        self.assertEqual(err, "")

    def test_validate_document_file_docx(self):
        fake_docx_bytes = MAGIC_BYTES_ZIP + b"\x00\x00fake zip docx data"
        is_valid, mime, err = validate_document_file(fake_docx_bytes, "resume.docx")
        self.assertTrue(is_valid)
        self.assertEqual(mime, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        self.assertEqual(err, "")

    def test_validate_document_file_invalid(self):
        invalid_bytes = b"GIF89a fake image"
        is_valid, mime, err = validate_document_file(invalid_bytes, "image.png")
        self.assertFalse(is_valid)
        self.assertIn("Format file tidak didukung", err)

    # ==========================================
    # 4. Tests for Document Builder (Docx Rendering)
    # ==========================================
    def test_render_ats_review_docx(self):
        data = {
            "overall_score": 88,
            "target_role": "Data Scientist",
            "summary": "CV kandidat sangat solid dengan analisis kuantitatif kuat.",
            "breakdown_scores": {"ats_compatibility": 95, "content_impact": 85},
            "strengths": ["Metrik bisnis terbukti jelas", "Struktur rapi"],
            "findings": [{"section": "Summary", "issue": "Sedikit kepanjangan", "recommendation": "Persingkat jadi 3 baris"}]
        }
        docx_bytes = render_ats_review_docx(data)
        self.assertIsInstance(docx_bytes, bytes)
        self.assertGreater(len(docx_bytes), 100)

        # Validasi file docx bisa dibuka kembali oleh python-docx
        doc = Document(io.BytesIO(docx_bytes))
        full_text = "\n".join([p.text for p in doc.paragraphs])
        self.assertIn("BOONTRACK ATS AUDIT", full_text)
        self.assertIn("Data Scientist", full_text)

    def test_render_cv_rewrite_docx(self):
        data = {
            "full_name": "Budi Santoso",
            "target_position": "Product Manager",
            "email": "budi@example.com",
            "phone": "+62812345678",
            "summary": "Product Manager berpengalaman memimpin 5 squad engineering.",
            "skills": {"technical": ["Jira", "SQL", "Figma"], "soft": ["Leadership", "Agile"]},
            "experience": [
                {
                    "role": "Lead PM",
                    "company": "Tech Corp",
                    "period": "2021 - 2024",
                    "bullets": ["Meningkatkan MAU sebesar 40%", "Meluncurkan produk fintech baru"]
                }
            ],
            "education": [{"degree": "S1 Manajemen", "institution": "Universitas Gadjah Mada", "year": "2019"}]
        }
        docx_bytes = render_cv_rewrite_docx(data)
        self.assertIsInstance(docx_bytes, bytes)
        self.assertGreater(len(docx_bytes), 100)

        doc = Document(io.BytesIO(docx_bytes))
        full_text = "\n".join([p.text for p in doc.paragraphs])
        self.assertIn("BUDI SANTOSO", full_text)
        self.assertIn("Product Manager", full_text)

    def test_render_paraphrase_docx(self):
        data = {
            "title": "Analisis Pengaruh AI dalam Rekrutmen",
            "tone": "Formal & Akademik",
            "original_word_count": 600,
            "paraphrased_word_count": 550,
            "key_takeaways": ["Penerapan AI meningkatkan akurasi screening kandidat."],
            "sections": [
                {"heading": "Pendahuluan", "content": "Penggunaan kecerdasan buatan dalam rekrutmen berkembang pesat."}
            ]
        }
        docx_bytes = render_paraphrase_docx(data)
        self.assertIsInstance(docx_bytes, bytes)
        doc = Document(io.BytesIO(docx_bytes))
        full_text = "\n".join([p.text for p in doc.paragraphs])
        self.assertIn("ANALISIS PENGARUH AI", full_text)

    # ==========================================
    # 5. Tests for Task Router & Intake Pipeline
    # ==========================================
    @patch("app.services.document_engine.get_supabase")
    async def test_intake_document_job_success(self, mock_get_supabase):
        mock_client = MagicMock()
        mock_table = MagicMock()
        mock_client.table.return_value = mock_table
        mock_table.insert.return_value = MagicMock(execute=MagicMock(return_value={"data": []}))
        mock_get_supabase.return_value = mock_client

        # Buat dokumen valid
        fake_doc = Document()
        fake_doc.add_paragraph("Ini adalah contoh CV sederhana untuk pengujian otomatis intake pipeline.")
        doc_io = io.BytesIO()
        fake_doc.save(doc_io)
        docx_bytes = doc_io.getvalue()

        res = await intake_document_job(
            tenant_id="boontrack-career",
            task_type=TASK_PARAPHRASE,
            filename="my_doc.docx",
            file_bytes=docx_bytes,
            user_id="user-123",
            user_phone="6281237450222"
        )

        self.assertEqual(res["status"], "WAITING_PAYMENT")
        self.assertEqual(res["task_type"], "POLISH_REPHRASE")
        self.assertIsNotNone(res["job_id"])
        self.assertGreater(res["word_count"], 0)
        self.assertEqual(res["pricing"]["pricing_tier"], "TIER_1")

    @patch("app.services.document_engine.update_job_status", new_callable=AsyncMock)
    @patch("app.services.document_engine.ai_gateway.generate", new_callable=AsyncMock)
    @patch("app.services.document_engine.send_whatsapp_document", new_callable=AsyncMock)
    @patch("app.services.document_engine.send_whatsapp_text", new_callable=AsyncMock)
    async def test_process_document_job_async_flow(self, mock_send_wa, mock_send_doc, mock_ai_gen, mock_update_status):
        mock_ai_gen.return_value = """```json
        {
            "overall_score": 90,
            "target_role": "Senior Engineer",
            "summary": "CV luar biasa.",
            "breakdown_scores": {"ats_compatibility": 95},
            "strengths": ["Berpengalaman luas"],
            "findings": []
        }
        ```"""

        await process_document_job_async(
            job_id="test-job-999",
            tenant_id="boontrack-career",
            task_type=TASK_ATS_REVIEW,
            filename="test_cv.docx",
            raw_text="Contoh naskah resume untuk testing.",
            user_phone="6281237450222"
        )

        # Verifikasi status updates
        self.assertTrue(mock_update_status.called)
        # Verifikasi WhatsApp notification dikirim
        mock_send_wa.assert_called_once()
        self.assertIn("Dokumen Anda telah selesai diproses", mock_send_wa.call_args[1]["text"])
        # Verifikasi WhatsApp document attachment dikirim
        mock_send_doc.assert_called_once()
        self.assertEqual(mock_send_doc.call_args[1]["filename"], "CV_Hasil_Polish.docx")


if __name__ == "__main__":
    unittest.main()
