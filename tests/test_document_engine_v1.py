"""tests/test_document_engine_v1.py
Comprehensive Test Suite for CTO Directive: Document Engine v1 Refactor & Strict Isolation Architecture.

Validates:
1. Deterministic routing for all 4 core products (CV_ATS, CV_REVIEW, POLISH_REPHRASE, CAREER_PRO_BUNDLE).
2. Strict rejection on unsupported task types and malformed documents.
3. Citation & statistical formula preservation during academic paraphrase.
4. Metric fabrication guard (asserting unprovided numbers are never hallucinated; placeholder brackets used).
5. Exact output filenames contract matching architecture specification.
6. JSON schema validation and clean python-docx compilation.
7. Anti-summarization length ratio assertion on academic paraphrase.
8. Tenant propagation and job status state machine.
"""

import io
import unittest
from unittest.mock import patch, AsyncMock, MagicMock
from docx import Document

from app.prompts import (
    get_prompt_for_task,
    get_fallback_for_task,
    normalize_prompt_task,
    TASK_CV_BUILD,
    TASK_CV_ATS,
    TASK_CV_REVIEW,
    TASK_POLISH_REPHRASE,
    TASK_CAREER_PRO_BUNDLE
)
from app.services.document_renderers import (
    render_document,
    render_cv_ats_docx,
    render_cv_review_docx,
    render_polish_rephrase_docx,
    render_career_pro_bundle_docx
)
from app.services.document_engine import (
    validate_document_file,
    intake_document_job,
    execute_ai_document_task,
    get_output_document_filename,
    process_document_job_async
)
from app.engines.rephrase_engine import academic_rephrase_engine


class TestDocumentEngineV1(unittest.IsolatedAsyncioTestCase):
    """Test suite validating Document Engine v1 architecture and strict isolation."""

    def test_exact_output_filenames_contract(self):
        """Matrix Verification: Exact output filenames matching CTO specification contract."""
        # 1. CV_BUILD / CV_ATS
        self.assertEqual(get_output_document_filename("CV_ATS"), "CV_ATS_Optimasi.docx")
        self.assertEqual(get_output_document_filename("CV_BUILD"), "CV_ATS_Optimasi.docx")
        self.assertEqual(get_output_document_filename("CV_POLISH_REWRITE"), "CV_ATS_Optimasi.docx")

        # 2. CV_REVIEW
        self.assertEqual(get_output_document_filename("CV_REVIEW"), "Laporan_Review_CV_HR.docx")
        self.assertEqual(get_output_document_filename("ATS_DIAGNOSTIC"), "Laporan_Review_CV_HR.docx")
        self.assertEqual(get_output_document_filename("ATS_REVIEW"), "Laporan_Review_CV_HR.docx")

        # 3. POLISH_REPHRASE
        self.assertEqual(get_output_document_filename("POLISH_REPHRASE"), "Naskah_Hasil_Parafrase.docx")
        self.assertEqual(get_output_document_filename("PARAPHRASE"), "Naskah_Hasil_Parafrase.docx")
        self.assertEqual(
            get_output_document_filename("POLISH_REPHRASE", "BAB_III_Metodologi.pdf"),
            "BAB_III_Metodologi_Hasil_Parafrase.docx"
        )

        # 4. CAREER_PRO_BUNDLE
        self.assertEqual(get_output_document_filename("CAREER_PRO_BUNDLE"), "Paket_Lengkap_Karir_Pro.docx")
        self.assertEqual(get_output_document_filename("BUNDLE_CAREER"), "Paket_Lengkap_Karir_Pro.docx")

    def test_rejection_on_unsupported_task_type(self):
        """Memvalidasi rejection deterministik jika task_type tidak didukung."""
        invalid_task = "UNSUPPORTED_HACK_TASK"

        # 1. Rejection in prompt dispatcher
        with self.assertRaises(ValueError) as ctx:
            get_prompt_for_task(invalid_task, "Sample text")
        self.assertIn("Unsupported document task_type", str(ctx.exception))

        # 2. Rejection in execute_ai_document_task
        with self.assertRaises(ValueError) as ctx2:
            import asyncio
            asyncio.run(execute_ai_document_task(invalid_task, "Sample text"))
        self.assertIn("Tipe task tidak didukung", str(ctx2.exception))

    @patch("app.services.document_engine.r2_storage_service.upload_file", new_callable=AsyncMock)
    @patch("app.services.document_engine.get_supabase")
    async def test_intake_rejection_on_invalid_task_or_malformed_document(self, mock_supa, mock_r2):
        """Memvalidasi intake_document_job menolak file korup dan task_type tak dikenal."""
        # A. Malformed Document (Empty / Non-PDF/DOCX)
        corrupt_bytes = b"bad"
        res_malformed = await intake_document_job(
            tenant_id="boontrack-career",
            task_type="CV_ATS",
            filename="bad.xyz",
            file_bytes=corrupt_bytes
        )
        self.assertEqual(res_malformed["status"], "REJECTED")
        self.assertFalse(res_malformed["is_valid"])

        # B. Unsupported task_type
        valid_magic_docx = b"PK\x03\x04dummy_docx_content"
        res_invalid_task = await intake_document_job(
            tenant_id="boontrack-career",
            task_type="ILLEGAL_TASK",
            filename="test.docx",
            file_bytes=valid_magic_docx
        )
        self.assertEqual(res_invalid_task["status"], "REJECTED")
        self.assertIn("tidak didukung", res_invalid_task["error"])

    def test_metric_fabrication_guard(self):
        """Hard Rule on Metrics: ZERO fabricated metrics on CVs."""
        # 1. Validasi instruksi prompt strategi CV_ATS
        prompt = get_prompt_for_task("CV_ATS", "Pengalaman kerja: staff logistik input barang gudang.")
        self.assertIn("ZERO FABRICATED METRICS", prompt)
        self.assertIn("DILARANG MEMALSUKAN ANGKA", prompt)
        self.assertIn("[Tambahkan", prompt)

        # 2. Validasi fallback data tidak mengarang angka / persen fiktif
        fallback_cv = get_fallback_for_task("CV_ATS")
        bullets = fallback_cv["experience"][0]["bullets"]
        for b in bullets:
            # Tidak boleh ada angka persentase fiktif
            self.assertNotIn("%", b)
            # Jika ada instruksi metrik, wajib menggunakan placeholder bracket
            if "metrik" in b.lower():
                self.assertIn("[Tambahkan", b)

    async def test_citation_and_formula_preservation(self):
        """Strict Preservation: Sitasi naratif, sitasi kurung, IEEE, dan formula matematika tidak boleh diubah."""
        raw_academic_sample = (
            "Penelitian ini merujuk pada teori kepemimpinan transformasional Sugiyono (2015). "
            "Sebagaimana diungkapkan dalam penelitian sebelumnya (Kuncoro, 2018: 45) serta literatur [1-3], "
            "model regresi linier berganda diformulasikan sebagai Y = a + b1X1 + b2X2. "
            "Uji signifikansi statistik menunjukkan hasil signifikan p < 0.05 dengan nilai R^2 = 0.85."
        )

        res = await academic_rephrase_engine.rephrase_document(raw_academic_sample, filename="Skripsi_Bab2.docx")
        full_output = res["full_text"]

        # Preservasi Sitasi
        self.assertIn("Sugiyono (2015)", full_output)
        self.assertIn("(Kuncoro, 2018: 45)", full_output)
        self.assertIn("[1-3]", full_output)

        # Preservasi Formula Matematika & Statistik
        self.assertIn("Y = a + b1X1 + b2X2", full_output)
        self.assertIn("p < 0.05", full_output)
        self.assertIn("R^2 = 0.85", full_output)

    async def test_polish_rephrase_anti_summarization_ratio(self):
        """Memvalidasi assertion rasio panjang teks output vs input (no aggressive summarization)."""
        input_text = (
            "Kajian ini meneliti efektivitas penerapan tata kelola manajemen rantai pasok modern pada industri manufaktur. "
            "Berbagai hambatan operasional dan fluktuasi pasokan bahan baku dianalisis secara komprehensif menggunakan pendekatan empiris. "
            "Data primer dikumpulkan melalui kuesioner terstruktur yang disebarkan kepada responden profesional di berbagai kota industri. "
            "Hasil pengujian hipotesis mengindikasikan korelasi positif yang signifikan antara integrasi sistem informasi dan performa logistik. "
            "Oleh karena itu, implementasi teknologi pelacakan terpadu menjadi esensial dalam rangka memitigasi risiko disrupsi operasional."
        )
        res = await academic_rephrase_engine.rephrase_document(input_text, filename="Naskah_Jurnal.docx")
        
        self.assertIn("length_ratio", res)
        self.assertTrue(res["anti_summarization_passed"])
        self.assertGreaterEqual(res["length_ratio"], 0.70)
        self.assertGreaterEqual(res["paraphrased_word_count"], int(res["original_word_count"] * 0.70))

    def test_json_schema_validation_and_clean_docx_compilation(self):
        """Memvalidasi rendering DOCX bersih (magic bytes PK\x03\x04) dan dapat di-parse ulang oleh python-docx."""
        # 1. CV ATS
        cv_data = {
            "full_name": "Ahmad Fauzi",
            "target_position": "Senior Backend Architect",
            "email": "ahmad@example.com",
            "phone": "+6281234567890",
            "location": "Bandung, Indonesia",
            "linkedin": "linkedin.com/in/ahmadfauzi",
            "summary": "Software Architect berpengalaman dalam distributed systems.",
            "skills": {"technical_skills": ["Python", "FastAPI", "PostgreSQL", "Docker"]},
            "experience": [{
                "role": "Lead Engineer",
                "company": "Tech Corp",
                "period": "2021 - Sekarang",
                "bullets": ["Merancang arsitektur mikroservis dengan uptime 99.9% [Tambahkan metrik SLA]"]
            }],
            "education": [{"degree": "S1 Teknik Informatika", "institution": "ITB", "year": "2020"}],
            "certifications": ["AWS Certified Solutions Architect"]
        }
        cv_bytes = render_document("CV_ATS", cv_data)
        self.assertTrue(cv_bytes.startswith(b"PK\x03\x04"))
        doc_cv = Document(io.BytesIO(cv_bytes))
        cv_text = "\n".join([p.text for p in doc_cv.paragraphs])
        self.assertIn("AHMAD FAUZI", cv_text)
        self.assertIn("Senior Backend Architect", cv_text)

        # 2. CV REVIEW
        review_data = {
            "ats_score": 85,
            "overall_score": 85,
            "target_role": "DevOps Engineer",
            "summary": "CV sangat baik dengan kesiapan ATS tinggi.",
            "breakdown_scores": {"ats_compatibility": 90, "content_impact": 80},
            "strengths": ["Pengalaman Kubernetes terbukti"],
            "red_flags": ["Format kontak belum menyertakan link GitHub"],
            "missing_keywords": ["Terraform", "CI/CD Pipeline"],
            "actionable_fixes": [{"section": "Experience", "issue": "Deskripsi pasif", "fix": "Ubah ke Action Verbs aktif"}]
        }
        review_bytes = render_document("CV_REVIEW", review_data)
        self.assertTrue(review_bytes.startswith(b"PK\x03\x04"))
        doc_rev = Document(io.BytesIO(review_bytes))
        rev_text = "\n".join([p.text for p in doc_rev.paragraphs])
        self.assertIn("LAPORAN EVALUASI & REVIEW CV HR", rev_text)
        self.assertIn("85/100", rev_text)

        # 3. POLISH REPHRASE
        rephrase_data = {
            "title": "Naskah Hasil Parafrase",
            "tone": "Akademik Formal (EYD V)",
            "original_word_count": 100,
            "paraphrased_word_count": 105,
            "full_text": "Naskah ilmiah hasil rekonstruksi komprehensif tanpa reduksi informasi."
        }
        rep_bytes = render_document("POLISH_REPHRASE", rephrase_data)
        self.assertTrue(rep_bytes.startswith(b"PK\x03\x04"))
        doc_rep = Document(io.BytesIO(rep_bytes))
        rep_text = "\n".join([p.text for p in doc_rep.paragraphs])
        self.assertIn("Naskah ilmiah hasil rekonstruksi", rep_text)

        # 4. CAREER PRO BUNDLE
        bundle_data = {
            "full_name": "Dewi Sartika",
            "target_position": "Product Manager",
            "summary": "PM berorientasi data dengan pengalaman growth.",
            "experience": [{"role": "PM", "company": "StartUp Indo", "period": "2022-Now", "bullets": ["Memimpin rilis fitur"]}],
            "hr_recommendations": {
                "profile_readiness": "Sangat Siap",
                "key_strengths": ["Product sense tajam"],
                "interview_tips": ["Gunakan metode STAR"]
            },
            "cover_letter": {
                "recipient": "Hiring Manager",
                "subject": "Lamaran PM",
                "opening": "Dengan hormat, saya mengajukan lamaran...",
                "body_paragraphs": ["Saya telah berpengalaman mengelola roadmap produk."],
                "closing": "Terima kasih atas perhatiannya."
            }
        }
        bundle_bytes = render_document("CAREER_PRO_BUNDLE", bundle_data)
        self.assertTrue(bundle_bytes.startswith(b"PK\x03\x04"))
        doc_bun = Document(io.BytesIO(bundle_bytes))
        bun_text = "\n".join([p.text for p in doc_bun.paragraphs])
        self.assertIn("BOONTRACK CAREER PRO BUNDLE", bun_text)
        self.assertIn("DEWI SARTIKA", bun_text)
        self.assertIn("SURAT LAMARAN KERJA (COVER LETTER)", bun_text)

    @patch("app.services.document_engine.r2_storage_service.upload_file", new_callable=AsyncMock)
    @patch("app.services.document_engine.update_job_status", new_callable=AsyncMock)
    @patch("app.services.document_engine.send_whatsapp_text", new_callable=AsyncMock)
    @patch("app.services.document_engine.send_whatsapp_document", new_callable=AsyncMock)
    async def test_tenant_propagation_and_job_state_machine(
        self,
        mock_send_doc,
        mock_send_txt,
        mock_update_status,
        mock_r2_upload
    ):
        """Memvalidasi tenant-agnostic design: tenant_id terpropagasi dengan benar ke storage key dan notifikasi."""
        partner_tenant = "partner-holding-agency"
        job_id = "job-tenant-prop-123"

        await process_document_job_async(
            job_id=job_id,
            tenant_id=partner_tenant,
            task_type="CV_ATS",
            filename="kandidat_cv.docx",
            raw_text="Contoh resume profesional.",
            user_phone="628999888777",
            amount=25000,
            invoice_id="INV-TENANT-123"
        )

        # 1. Memverifikasi status transition ke PROCESSING lalu COMPLETED
        self.assertTrue(mock_update_status.called)
        first_call = mock_update_status.call_args_list[0]
        self.assertEqual(first_call[1]["status"], "PROCESSING")
        
        last_call = mock_update_status.call_args_list[-1]
        self.assertEqual(last_call[1]["status"], "COMPLETED")
        self.assertEqual(last_call[1]["extra_fields"]["payment_status"], "PAID")

        # 2. Memverifikasi R2 storage key memuat tenant_id yang dipropagasikan
        mock_r2_upload.assert_called_once()
        upload_storage_key = mock_r2_upload.call_args[1]["storage_key"]
        self.assertIn(f"output/{partner_tenant}/", upload_storage_key)

        # 3. Memverifikasi WhatsApp delivery diteruskan dengan tenant_id partner
        mock_send_txt.assert_called_once()
        self.assertEqual(mock_send_txt.call_args[1]["tenant_id"], partner_tenant)

        mock_send_doc.assert_called_once()
        self.assertEqual(mock_send_doc.call_args[1]["tenant_id"], partner_tenant)
        self.assertEqual(mock_send_doc.call_args[1]["filename"], "CV_ATS_Optimasi.docx")


class TestPolishRephraseRendererPayload(unittest.TestCase):
    """Tests validating render_document('POLISH_REPHRASE', ...) payload contract & output size."""

    # --- 1000-word academic text fixture ---
    LONG_ACADEMIC_TEXT = (
        "Metodologi penelitian ini menggunakan pendekatan kuantitatif dengan desain survei analitik "
        "yang melibatkan instrumen kuesioner terstruktur berbasis skala Likert 1-5. "
        "Populasi penelitian mencakup seluruh mahasiswa aktif program studi Manajemen Bisnis "
        "Universitas X angkatan 2021-2023 dengan jumlah keseluruhan 1.240 responden.\n\n"
        "Teknik pengambilan sampel yang diterapkan adalah stratified random sampling guna "
        "memastikan representasi proporsional dari seluruh angkatan yang terlibat. "
        "Berdasarkan rumus Slovin dengan tingkat kesalahan toleransi 5%, ditetapkan sampel "
        "minimum sebesar 304 responden. Proses distribusi kuesioner dilaksanakan secara daring "
        "melalui platform Google Forms selama tiga minggu, menghasilkan 312 respons yang valid "
        "dan dapat digunakan dalam analisis lebih lanjut.\n\n"
        "Analisis data dilakukan menggunakan perangkat lunak SPSS versi 26 dan SmartPLS 3.0 "
        "untuk pengujian model struktural (Structural Equation Modeling / SEM). "
        "Uji validitas konstruk dilaksanakan melalui Confirmatory Factor Analysis (CFA) "
        "dengan threshold factor loading >= 0.70 sebagaimana direkomendasikan oleh Hair et al. (2019). "
        "Reliabilitas instrumen dikonfirmasi melalui nilai Cronbach Alpha yang berkisar antara "
        "0.82 hingga 0.91, jauh melampaui batas minimum yang dapat diterima sebesar 0.70.\n\n"
        "Variabel independen dalam penelitian ini adalah kualitas layanan digital (X1), "
        "kepercayaan pengguna platform (X2), dan kemudahan penggunaan antarmuka (X3). "
        "Variabel dependen yang diukur adalah kepuasan pelanggan (Y) dan loyalitas pengguna jangka panjang (Z). "
        "Model struktural yang diuji memuat empat hipotesis utama yang dirumuskan berdasarkan "
        "tinjauan literatur komprehensif terhadap 47 artikel ilmiah terindeks Scopus dan WoS.\n\n"
        "Hasil pengujian hipotesis menunjukkan bahwa kualitas layanan digital berpengaruh "
        "signifikan terhadap kepuasan pelanggan (β = 0.42, p < 0.001), kepercayaan pengguna "
        "berpengaruh positif terhadap loyalitas jangka panjang (β = 0.38, p < 0.05), "
        "dan kemudahan penggunaan antarmuka secara signifikan memediasi hubungan antara "
        "kepercayaan dan kepuasan (β = 0.29, p < 0.01). "
        "Secara keseluruhan, model yang diajukan mampu menjelaskan 67.3% varians dalam variabel "
        "kepuasan pelanggan, yang mengindikasikan daya prediktif model yang memadai.\n\n"
        "Temuan-temuan tersebut memberikan kontribusi teoretis yang bermakna bagi pengembangan "
        "literatur Service Quality dalam konteks platform digital ekonomi kreatif di Indonesia. "
        "Secara praktis, hasil penelitian ini merekomendasikan bahwa pengelola platform "
        "sebaiknya memprioritaskan investasi pada penyempurnaan antarmuka pengguna dan "
        "mekanisme pembangunan kepercayaan digital sebagai strategi retensi pelanggan yang efektif.\n\n"
    ) * 2  # ~1050+ words

    def test_render_polish_rephrase_with_sections_produces_large_docx(self):
        """render_document('POLISH_REPHRASE', data_with_sections) menghasilkan DOCX lebih besar
        dari dokumen kosong dan memuat konten teks naskah (verifikasi body text flow ke renderer)."""
        sections = [{"heading": "Bagian 1", "content": self.LONG_ACADEMIC_TEXT}]
        data = {
            "title": "Penyempurnaan Akademis: Bab Iii Metodologi",
            "tone": "Akademik Formal (EYD V)",
            "original_word_count": len(self.LONG_ACADEMIC_TEXT.split()),
            "paraphrased_word_count": len(self.LONG_ACADEMIC_TEXT.split()),
            "key_takeaways": [
                "Naskah telah disempurnakan sesuai kaidah EYD V.",
                "Seluruh sitasi akademik dan formula statistik dipertahankan 100%.",
            ],
            "sections": sections,
            "full_text": self.LONG_ACADEMIC_TEXT,
            "full_paraphrased_text": self.LONG_ACADEMIC_TEXT,
        }

        # Render with body text
        docx_bytes = render_document("POLISH_REPHRASE", data)
        word_count = len(self.LONG_ACADEMIC_TEXT.split())
        self.assertIsInstance(docx_bytes, bytes)
        self.assertGreater(len(docx_bytes), 36_000,
            f"DOCX output must be > 36kB (base docx overhead), got {len(docx_bytes)} bytes.")

        # Render empty doc to measure baseline (no body text)
        empty_data = {**data, "sections": [], "full_text": "", "full_paraphrased_text": "", "key_takeaways": []}
        empty_docx = render_document("POLISH_REPHRASE", empty_data)

        self.assertGreater(len(docx_bytes), len(empty_docx),
            f"DOCX with {word_count} words ({len(docx_bytes)} bytes) must be larger than empty doc ({len(empty_docx)} bytes). "
            "Body text is NOT flowing into renderer — key mismatch bug!")

    def test_render_polish_rephrase_full_text_fallback_produces_large_docx(self):
        """render_document('POLISH_REPHRASE', data) dengan sections=[] tapi full_text terpopulasi
        menghasilkan DOCX lebih besar dari dokumen dengan full_text kosong."""
        data_with_text = {
            "title": "Naskah Hasil Parafrase Akademis",
            "tone": "Akademik Formal (EYD V)",
            "original_word_count": len(self.LONG_ACADEMIC_TEXT.split()),
            "paraphrased_word_count": len(self.LONG_ACADEMIC_TEXT.split()),
            "key_takeaways": [],
            "sections": [],  # intentionally empty -> should fall back to full_text
            "full_text": self.LONG_ACADEMIC_TEXT,
        }
        data_empty = {**data_with_text, "full_text": ""}

        docx_full = render_document("POLISH_REPHRASE", data_with_text)
        docx_empty = render_document("POLISH_REPHRASE", data_empty)

        self.assertIsInstance(docx_full, bytes)
        self.assertGreater(len(docx_full), 36_000, f"DOCX must be > 36kB base, got {len(docx_full)} bytes.")
        self.assertGreater(len(docx_full), len(docx_empty),
            f"DOCX with full_text ({len(docx_full)} bytes) must be larger than empty ({len(docx_empty)} bytes). "
            "full_text fallback is NOT rendering body paragraphs!")

    def test_render_polish_rephrase_plain_string_input_produces_large_docx(self):
        """render_polish_rephrase_docx menerima plain string langsung dan menghasilkan
        DOCX lebih besar dari dokumen header-only (verifikasi _normalize_data bekerja)."""
        from app.services.document_renderers.polish_rephrase_renderer import render_polish_rephrase_docx
        docx_with_text = render_polish_rephrase_docx(self.LONG_ACADEMIC_TEXT)
        docx_empty = render_polish_rephrase_docx("")
        self.assertIsInstance(docx_with_text, bytes)
        self.assertGreater(len(docx_with_text), 36_000,
            f"String input must produce > 36kB DOCX base, got {len(docx_with_text)} bytes.")
        self.assertGreater(len(docx_with_text), len(docx_empty),
            f"String-input DOCX ({len(docx_with_text)} bytes) must be larger than empty DOCX ({len(docx_empty)} bytes). "
            "_normalize_data is not wrapping plain string into sections correctly!")

    def test_get_fallback_data_polish_rephrase_returns_correct_schema(self):
        """get_fallback_for_task('POLISH_REPHRASE', raw_text=...) mengembalikan schema sections/full_text yang benar untuk renderer."""
        from app.prompts import get_fallback_for_task, TASK_POLISH_REPHRASE
        fallback = get_fallback_for_task(TASK_POLISH_REPHRASE, raw_text=self.LONG_ACADEMIC_TEXT)

        # Must have correct keys for polish rephrase renderer
        self.assertIn("sections", fallback, "Fallback POLISH_REPHRASE harus memuat key 'sections'")
        self.assertIn("full_text", fallback, "Fallback POLISH_REPHRASE harus memuat key 'full_text'")
        self.assertIsInstance(fallback["sections"], list)
        self.assertGreater(len(fallback["sections"]), 0, "sections tidak boleh kosong")
        self.assertIn("content", fallback["sections"][0], "section dict harus memuat key 'content'")
        self.assertTrue(fallback["full_text"], "full_text tidak boleh string kosong")

        # Must NOT be CV schema (no full_name / target_position)
        self.assertNotIn("full_name", fallback,
            "CRITICAL: get_fallback_for_task(POLISH_REPHRASE) mengembalikan schema CV_ATS! Key mismatch bug!")
        self.assertNotIn("target_position", fallback,
            "CRITICAL: get_fallback_for_task(POLISH_REPHRASE) mengembalikan schema CV_ATS! Key mismatch bug!")

        # Rendering the fallback must produce a docx larger than an empty doc
        docx_bytes = render_document("POLISH_REPHRASE", fallback)
        docx_empty = render_document("POLISH_REPHRASE", {"sections": [], "full_text": ""})
        self.assertGreater(len(docx_bytes), len(docx_empty),
            f"Rendering fallback output ({len(docx_bytes)} bytes) is not larger than empty doc ({len(docx_empty)} bytes) — body text is empty/missing!")


if __name__ == "__main__":
    unittest.main()
