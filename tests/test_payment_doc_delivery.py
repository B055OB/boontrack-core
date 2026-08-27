import io
import os
import json
import unittest
from unittest.mock import AsyncMock, patch, MagicMock
from aiohttp import web
from aiohttp.test_utils import AioHTTPTestCase, unittest_run_loop
from docx import Document

from app.core.server import create_web_app
from app.payments.matcher import (
    extract_clean_dana_amount,
    find_matching_unpaid_job,
    match_and_fulfill_payment,
    handle_admin_verify_command,
    handle_admin_retry_doc_command
)
from app.services.whatsapp_service import send_whatsapp_document, send_whatsapp_text
from app.services.document_engine import (
    deliver_completed_document_job,
    process_document_job_async,
    intake_document_job
)
from app.services.reconciliation_service import PAYMENT_INTENTS
from app.services.cv_state_engine import GLOBAL_USER_STATES


class TestPaymentDocDelivery(AioHTTPTestCase):

    async def get_application(self):
        return create_web_app()

    def setUp(self):
        super().setUp()
        GLOBAL_USER_STATES.clear()
        PAYMENT_INTENTS.clear()

    # ==========================================
    # 1. DANA Nominal Extraction Tests
    # ==========================================
    def test_extract_clean_dana_amount_variations(self):
        """Validasi ekstraksi nominal bersih dari berbagai format teks notifikasi DANA Bisnis."""
        # Format standard DANA notification
        self.assertEqual(extract_clean_dana_amount("DANA Masuk Rp5.083 dari Siti"), 5083)
        self.assertEqual(extract_clean_dana_amount("DANA: Berhasil menerima pembayaran QRIS sebesar Rp 5.083"), 5083)
        self.assertEqual(extract_clean_dana_amount("Mutasi Masuk DANA Bisnis Rp5.083,00"), 5083)
        self.assertEqual(extract_clean_dana_amount("Transfer masuk Rp5083"), 5083)
        self.assertEqual(extract_clean_dana_amount("Pembayaran QRIS Rp. 5.083 berhasil"), 5083)
        self.assertEqual(extract_clean_dana_amount("IDR 5,083.00 diterima"), 5083)
        
        # Dict payload
        self.assertEqual(extract_clean_dana_amount({"amount": 5083}), 5083)
        self.assertEqual(extract_clean_dana_amount({"nominal": "5.083"}), 5083)
        self.assertEqual(extract_clean_dana_amount({"raw_text": "DANA QRIS Masuk Rp5.083"}), 5083)

        # Non-matching
        self.assertEqual(extract_clean_dana_amount(""), 0)
        self.assertEqual(extract_clean_dana_amount("Pesan halo biasa tanpa nominal"), 0)

    # ==========================================
    # 2. Payment Matching & Document Delivery Flow
    # ==========================================
    @patch("app.payments.matcher.get_supabase")
    @patch("app.services.document_engine.get_supabase")
    @patch("app.services.document_engine.r2_storage_service.download_file", new_callable=AsyncMock)
    @patch("app.services.document_engine.send_whatsapp_document", new_callable=AsyncMock)
    @patch("app.services.document_engine.send_whatsapp_text", new_callable=AsyncMock)
    async def test_dana_callback_5083_matches_unpaid_job_and_delivers_docx(
        self,
        mock_send_text,
        mock_send_doc,
        mock_r2_download,
        mock_doc_engine_supabase,
        mock_matcher_supabase
    ):
        """Simulasi notifikasi pembayaran DANA Rp5.083 -> status job berubah ke PAID & file DOCX terkirim."""
        job_id = "job-uuid-5083-abc"
        user_phone = "6281299988877"

        # Mock Supabase document_jobs record
        mock_job_record = {
            "id": job_id,
            "tenant_id": "boontrack-career",
            "user_phone": user_phone,
            "price_amount": 5083,
            "payment_status": "UNPAID",
            "status": "COMPLETED",
            "task_type": "POLISH_REPHRASE",
            "result_storage_key": f"output/boontrack-career/{job_id}_result.docx"
        }

        # Mock query return with fluent builder
        mock_client = MagicMock()
        mock_table = MagicMock()
        mock_client.table.return_value = mock_table

        query_builder = MagicMock()
        query_builder.eq.return_value = query_builder
        query_builder.order.return_value = query_builder
        query_builder.limit.return_value = query_builder
        query_builder.execute.return_value = MagicMock(data=[mock_job_record])

        mock_table.select.return_value = query_builder
        mock_table.update.return_value.eq.return_value.execute.return_value = MagicMock(data=[mock_job_record])

        mock_matcher_supabase.return_value = mock_client
        mock_doc_engine_supabase.return_value = mock_client

        # Mock binary docx
        fake_doc = Document()
        fake_doc.add_paragraph("Hasil Polish CV ATS")
        doc_buf = io.BytesIO()
        fake_doc.save(doc_buf)
        mock_r2_download.return_value = doc_buf.getvalue()

        # Mock send_whatsapp_document success
        mock_send_doc.return_value = {"messages": [{"id": "wamid.123"}]}
        mock_send_text.return_value = {"messages": [{"id": "wamid.456"}]}

        # Eksekusi matching pembayaran Rp5.083
        result = await match_and_fulfill_payment(
            amount=5083,
            raw_text="DANA: Masuk Rp5.083 dari User Budi",
            tenant_id="boontrack-career",
            source="dana_reader"
        )

        self.assertEqual(result["status"], "SUCCESS")
        self.assertEqual(result["action"], "JOB_PAID_AND_DELIVERED")
        self.assertEqual(result["job_id"], job_id)
        self.assertEqual(result["user_phone"], user_phone)
        self.assertTrue(result["delivered"])

        # Verifikasi status di memory
        self.assertTrue(GLOBAL_USER_STATES[user_phone]["is_premium_paid"])
        self.assertEqual(GLOBAL_USER_STATES[user_phone]["tier"], "premium_unlocked")

        # Verifikasi attachment DOCX terkirim dengan parameter yang benar
        mock_send_doc.assert_called_once()
        call_kwargs = mock_send_doc.call_args[1]
        self.assertEqual(call_kwargs["to_phone"], user_phone)
        self.assertEqual(call_kwargs["filename"], "Naskah_Hasil_Parafrase.docx")
        self.assertEqual(call_kwargs["mime_type"], "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        self.assertIn("Naskah_Hasil_Parafrase.docx", call_kwargs["caption"])
        self.assertIn("kebijakan integritas", call_kwargs["caption"])

        # Verifikasi teks konfirmasi terpadu (Pesan 1) juga terkirim
        mock_send_text.assert_called_once()
        sent_text = mock_send_text.call_args[1]["text"]
        self.assertIn("PEMBAYARAN TERVERIFIKASI & LUNAS", sent_text)
        self.assertIn("Rp5,083", sent_text)
        self.assertIn("Dokumen Anda telah selesai diproses", sent_text)

    # ==========================================
    # 3. Webhook Listener Route Integration Test
    # ==========================================
    @unittest_run_loop
    @patch("app.routes.payment_webhook.match_and_fulfill_payment", new_callable=AsyncMock)
    async def test_webhook_listener_dana_5083(self, mock_fulfill):
        """Memvalidasi route POST /api/webhook/dana mengekstrak Rp5.083 dan memanggil matcher."""
        mock_fulfill.return_value = {
            "status": "SUCCESS",
            "action": "JOB_PAID_AND_DELIVERED",
            "job_id": "job-5083",
            "amount": 5083
        }

        payload = {
            "message": "DANA Bisnis: Masuk pembayaran QRIS Rp5.083",
            "package_name": "id.dana"
        }

        resp = await self.client.post("/api/webhook/dana", json=payload)
        self.assertEqual(resp.status, 200)
        data = await resp.json()
        self.assertEqual(data["status"], "SUCCESS")
        self.assertEqual(data["amount"], 5083)

        mock_fulfill.assert_called_once()
        self.assertEqual(mock_fulfill.call_args[1]["amount"], 5083)

    # ==========================================
    # 4. WhatsApp Service Attachment & Error Logging
    # ==========================================
    @patch("app.services.whatsapp_service.get_wa_credentials")
    @patch("httpx.AsyncClient.post")
    async def test_send_whatsapp_document_error_logging_on_failure(self, mock_http_post, mock_creds):
        """Memvalidasi penanganan dan logging jika WhatsApp Cloud API mengembalikan status non-200."""
        mock_creds.return_value = ("fake_token", "fake_phone_id", "v21.0")

        # Mock API returning 400 Bad Request
        mock_resp = MagicMock()
        mock_resp.status_code = 400
        mock_resp.text = '{"error":{"message":"Invalid parameter: media_id"}}'
        mock_http_post.return_value = mock_resp

        # Harus mengembalikan None dan mencatat error tanpa raise exception tak tertangani
        res = await send_whatsapp_document(
            to_phone="628123456789",
            file_path_or_bytes=b"dummy docx bytes",
            filename="CV_Hasil_Polish.docx",
            caption="Test Caption"
        )
        self.assertIsNone(res)

    # ==========================================
    # 5. Admin Fallback Commands (/verify & /retry_doc)
    # ==========================================
    @patch("app.payments.matcher.match_and_fulfill_payment", new_callable=AsyncMock)
    async def test_admin_verify_command_success(self, mock_fulfill):
        """Memvalidasi eksekusi perintah admin /verify 5083."""
        mock_fulfill.return_value = {
            "status": "SUCCESS",
            "action": "JOB_PAID_AND_DELIVERED",
            "job_id": "job-5083-xyz",
            "user_phone": "62812345678"
        }

        reply = await handle_admin_verify_command("/verify 5083")
        self.assertIn("Verifikasi Manual Berhasil", reply)
        self.assertIn("5,083", reply)
        self.assertIn("job-5083-xyz", reply)

    @patch("app.payments.matcher.get_supabase")
    @patch("app.services.document_engine.get_supabase")
    @patch("app.services.document_engine.r2_storage_service.download_file", new_callable=AsyncMock)
    @patch("app.services.document_engine.send_whatsapp_document", new_callable=AsyncMock)
    @patch("app.services.document_engine.send_whatsapp_text", new_callable=AsyncMock)
    async def test_admin_verify_command_triggers_full_fulfillment_and_docx_delivery(
        self,
        mock_send_text,
        mock_send_doc,
        mock_r2_download,
        mock_doc_engine_supabase,
        mock_matcher_supabase
    ):
        """Memvalidasi bahwa admin mengetik /verify 5083 langsung menandai job PAID dan mengirim DOCX ke WhatsApp."""
        job_id = "job-admin-verify-5083"
        user_phone = "628123456789"
        
        mock_job_record = {
            "id": job_id,
            "tenant_id": "boontrack-career",
            "user_phone": user_phone,
            "price_amount": 5083,
            "unique_code": 83,
            "payment_status": "UNPAID",
            "status": "COMPLETED",
            "task_type": "POLISH_REPHRASE",
            "result_storage_key": f"output/boontrack-career/{job_id}_result.docx"
        }

        mock_client = MagicMock()
        mock_table = MagicMock()
        mock_client.table.return_value = mock_table

        query_builder = MagicMock()
        query_builder.eq.return_value = query_builder
        query_builder.order.return_value = query_builder
        query_builder.limit.return_value = query_builder
        query_builder.execute.return_value = MagicMock(data=[mock_job_record])

        mock_table.select.return_value = query_builder
        mock_table.update.return_value = query_builder

        mock_matcher_supabase.return_value = mock_client
        mock_doc_engine_supabase.return_value = mock_client

        # Mock download DOCX bytes
        fake_doc = Document()
        fake_doc.add_paragraph("Hasil Polish & Rephrase Naskah.")
        buf = io.BytesIO()
        fake_doc.save(buf)
        mock_r2_download.return_value = buf.getvalue()
        mock_send_doc.return_value = {"messages": [{"id": "wamid.admin_verify"}]}

        # Eksekusi command admin /verify 5083
        reply = await handle_admin_verify_command("/verify 5083")
        
        # Validasi respon admin
        self.assertIn("Verifikasi Manual Berhasil", reply)
        self.assertIn("5,083", reply)
        self.assertIn(job_id, reply)

        # Validasi file terkirim ke WhatsApp user
        mock_send_doc.assert_called_once()
        self.assertEqual(mock_send_doc.call_args[1]["to_phone"], user_phone)
        self.assertEqual(mock_send_doc.call_args[1]["filename"], "Naskah_Hasil_Parafrase.docx")

    @patch("app.services.document_engine.deliver_completed_document_job", new_callable=AsyncMock)
    async def test_admin_retry_doc_command_success(self, mock_deliver):
        """Memvalidasi eksekusi perintah admin /retry_doc <job_id>."""
        mock_deliver.return_value = True

        reply = await handle_admin_retry_doc_command("/retry_doc job-5083-xyz")
        self.assertIn("Pengiriman Ulang Berhasil", reply)
        self.assertIn("job-5083-xyz", reply)

    # ==========================================
    # 6. Strict State Machine & Payment Lock Tests
    # ==========================================
    @patch("app.services.document_engine.get_supabase")
    @patch("app.services.document_engine.r2_storage_service.upload_file", new_callable=AsyncMock)
    @patch("app.services.document_engine.send_whatsapp_document", new_callable=AsyncMock)
    @patch("app.services.document_engine.send_whatsapp_text", new_callable=AsyncMock)
    async def test_intake_sets_status_waiting_payment_and_no_doc_sent_before_payment(
        self,
        mock_send_text,
        mock_send_doc,
        mock_r2_upload,
        mock_supabase
    ):
        """Memvalidasi bahwa saat dokumen berbayar di-intake, statusnya WAITING_PAYMENT dan TIDAK mengirim attachment .docx."""
        mock_client = MagicMock()
        mock_table = MagicMock()
        mock_client.table.return_value = mock_table
        mock_supabase.return_value = mock_client

        # Buat dokumen dummy berbayar (misal 600 kata -> Tier 2 Rp10.000 / Rp10.482)
        fake_text = "Metodologi penelitian ini menggunakan pendekatan kuantitatif. " * 50
        fake_doc = Document()
        fake_doc.add_paragraph(fake_text)
        buf = io.BytesIO()
        fake_doc.save(buf)
        file_bytes = buf.getvalue()

        # Eksekusi intake job berbayar
        intake_res = await intake_document_job(
            tenant_id="boontrack-career",
            task_type="POLISH_REPHRASE",
            filename="Naskah_Skripsi.docx",
            file_bytes=file_bytes,
            user_id="628111222333",
            user_phone="628111222333",
            exact_price_amount=10482
        )

        # 1. Validasi Status Response adalah WAITING_PAYMENT
        self.assertEqual(intake_res["status"], "WAITING_PAYMENT")
        self.assertEqual(intake_res["price_amount"], 10482)

        # 2. Validasi status record yang diinsert ke Supabase
        mock_table.insert.assert_called_once()
        inserted_record = mock_table.insert.call_args[0][0]
        self.assertEqual(inserted_record["status"], "WAITING_PAYMENT")
        self.assertEqual(inserted_record["payment_status"], "UNPAID")
        self.assertEqual(inserted_record["price_amount"], 10482)

        # 3. CRITICAL: JANGAN PERNAH panggil send_whatsapp_document sebelum status PAID!
        mock_send_doc.assert_not_called()
        mock_send_text.assert_not_called()

    @patch("app.services.document_engine.get_supabase")
    @patch("app.services.document_engine.send_whatsapp_document", new_callable=AsyncMock)
    @patch("app.services.document_engine.send_whatsapp_text", new_callable=AsyncMock)
    async def test_unpaid_job_delivery_is_strictly_blocked(
        self,
        mock_send_text,
        mock_send_doc,
        mock_supabase
    ):
        """Memvalidasi bahwa pemanggilan deliver_completed_document_job pada job UNPAID langsung ditolak (Strict Payment Lock)."""
        job_id = "job-unpaid-test-999"
        mock_client = MagicMock()
        mock_table = MagicMock()
        mock_client.table.return_value = mock_table

        mock_job_record = {
            "id": job_id,
            "tenant_id": "boontrack-career",
            "user_phone": "62811223344",
            "price_amount": 10482,
            "payment_status": "UNPAID",
            "status": "WAITING_PAYMENT",
            "task_type": "POLISH_REPHRASE"
        }

        query_builder = MagicMock()
        query_builder.eq.return_value = query_builder
        query_builder.execute.return_value = MagicMock(data=[mock_job_record])
        mock_table.select.return_value = query_builder
        mock_supabase.return_value = mock_client

        # Panggil deliver_completed_document_job
        res = await deliver_completed_document_job(job_id=job_id, tenant_id="boontrack-career")
        
        # Harus mengembalikan False dan TIDAK mengirim dokumen
        self.assertFalse(res)
        mock_send_doc.assert_not_called()
        mock_send_text.assert_not_called()

    # ==========================================
    # 7. Academic Rephrase Engine Tests
    # ==========================================
    async def test_academic_rephrase_engine_precleaning_and_citation_protection(self):
        """Memvalidasi pembersihan artefak PDF, spasi terputus, dan proteksi sitasi akademik."""
        from app.engines.rephrase_engine import academic_rephrase_engine

        raw_dirty_text = (
            "--- Page 1 ---\n"
            "Peneli tian ini meman faatkan va riabel terikat guna menguji hipo tesis.\n"
            "Hal ini sejalan dengan teori kognitif (Sugiyono, 2015) serta model [1-3].\n"
            "Di mana model regresi yang digunakan adalah Y = a + b1X1 + e dengan signifikansi p < 0.05.\n\n"
            "12\n\n"
            "Halaman 1 dari 10"
        )

        # 1. Test Pre-cleaning
        cleaned = academic_rephrase_engine.clean_academic_text(raw_dirty_text)
        self.assertNotIn("--- Page 1 ---", cleaned)
        self.assertNotIn("Halaman 1 dari 10", cleaned)
        self.assertIn("variabel", cleaned)
        self.assertIn("penelitian", cleaned)
        self.assertIn("hipotesis", cleaned)

        # 2. Test Masking & Unmasking
        masked, mask_map = academic_rephrase_engine.mask_academic_entities(cleaned)
        self.assertIn("__CIT_", masked)
        self.assertIn("__FORM_", masked)

        unmasked = academic_rephrase_engine.unmask_academic_entities(masked, mask_map)
        self.assertIn("(Sugiyono, 2015)", unmasked)
        self.assertIn("[1-3]", unmasked)
        self.assertIn("Y = a + b1X1 + e", unmasked)
        self.assertIn("p < 0.05", unmasked)

        # 3. Test Full Rephrase
        result = await academic_rephrase_engine.rephrase_document(raw_dirty_text, filename="Jurnal_Ilmiah.docx")
        self.assertEqual(result["tone"], "Akademik Formal (EYD V)")
        self.assertIn("(Sugiyono, 2015)", result["full_text"])
        self.assertIn("Y = a + b1X1 + e", result["full_text"])
        self.assertGreater(len(result["sections"]), 0)

    async def test_academic_rephrase_engine_chunking_and_seamless_stitching(self):
        """Memvalidasi bahwa dokumen panjang di-chunk per 600-800 kata dan dijahit secara utuh."""
        from app.engines.rephrase_engine import academic_rephrase_engine

        # Generate naskah 1.500 kata
        paragraph = (
            "Penerapan tata kelola sistem informasi yang efektif dan terstruktur sangat penting "
            "bagi kelancaran operasional organisasi modern. Dengan ini manajemen dapat mengukur kinerja secara akurat. "
        ) * 15 # ~300 kata per paragraf
        long_text = "\n\n".join([paragraph for _ in range(5)]) # ~1.500 kata

        chunks = academic_rephrase_engine.chunk_document_smart(long_text, max_words_per_chunk=650)
        self.assertGreaterEqual(len(chunks), 2)

        result = await academic_rephrase_engine.rephrase_document(long_text, filename="Tesis_Lengkap.docx")
        self.assertEqual(len(result["sections"]), len(chunks))
        self.assertGreaterEqual(result["paraphrased_word_count"], 1000)
        self.assertIn("Naskah telah disempurnakan", result["key_takeaways"][0])

    # ==========================================
    # 8. WhatsApp Media Upload & QRIS Image Tests
    # ==========================================
    @patch("app.services.whatsapp_service.get_wa_credentials")
    @patch("httpx.AsyncClient.post")
    @patch("app.services.whatsapp_service.log_to_supabase_messages", new_callable=AsyncMock)
    async def test_upload_media_success_and_send_whatsapp_image_with_media_id(
        self,
        mock_log,
        mock_post,
        mock_creds
    ):
        """Memvalidasi upload_media multipart dan pengiriman send_whatsapp_image dengan payload media_id."""
        from app.services.whatsapp_service import upload_media, send_whatsapp_image

        mock_creds.return_value = ("fake-token-xyz", "1029384756", "v20.0")

        # Mock respons: upload_media -> 200 dengan id, send message -> 200 dengan wamid
        mock_res_upload = MagicMock()
        mock_res_upload.status_code = 200
        mock_res_upload.json.return_value = {"id": "meta-media-id-998877"}

        mock_res_send = MagicMock()
        mock_res_send.status_code = 200
        mock_res_send.json.return_value = {"messages": [{"id": "wamid.img.12345"}]}

        mock_post.side_effect = [mock_res_upload, mock_res_send]

        fake_qr_bytes = b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR"
        caption = "Total Pembayaran: Rp25.482 (Termasuk Kode Unik)"

        res = await send_whatsapp_image(
            to_phone="628123456789",
            image_path_or_bytes=fake_qr_bytes,
            caption=caption,
            tenant_id="boontrack-career"
        )

        self.assertIsNotNone(res)
        self.assertEqual(mock_post.call_count, 2)

        # Cek call pertama: multipart upload_media
        upload_call = mock_post.call_args_list[0]
        self.assertIn("https://graph.facebook.com/v20.0/1029384756/media", upload_call[0][0])
        self.assertEqual(upload_call[1]["data"]["messaging_product"], "whatsapp")
        self.assertEqual(upload_call[1]["data"]["type"], "image/png")
        self.assertIn("file", upload_call[1]["files"])

        # Cek call kedua: kirim pesan image dengan media_id
        send_call = mock_post.call_args_list[1]
        self.assertIn("https://graph.facebook.com/v20.0/1029384756/messages", send_call[0][0])
        payload = send_call[1]["json"]
        self.assertEqual(payload["messaging_product"], "whatsapp")
        self.assertEqual(payload["type"], "image")
        self.assertEqual(payload["image"]["id"], "meta-media-id-998877")
        self.assertEqual(payload["image"]["caption"], caption)

    @patch("app.services.whatsapp_service.get_wa_credentials")
    @patch("httpx.AsyncClient.post")
    @patch("app.services.whatsapp_service.send_whatsapp_text", new_callable=AsyncMock)
    async def test_send_whatsapp_image_upload_failure_fallback_to_text(
        self,
        mock_send_text,
        mock_post,
        mock_creds
    ):
        """Memvalidasi jika upload_media gagal, sistem fallback mengirim teks instruksi."""
        from app.services.whatsapp_service import send_whatsapp_image

        mock_creds.return_value = ("fake-token-xyz", "1029384756", "v20.0")

        # Mock respons upload media gagal (HTTP 500)
        mock_res_upload = MagicMock()
        mock_res_upload.status_code = 500
        mock_res_upload.text = "Internal Server Error"
        mock_post.return_value = mock_res_upload

        mock_send_text.return_value = {"messages": [{"id": "wamid.txt.fallback"}]}

        fake_qr_bytes = b"bad_image_bytes"
        caption = "Total Pembayaran: Rp25.482"

        res = await send_whatsapp_image(
            to_phone="628123456789",
            image_path_or_bytes=fake_qr_bytes,
            caption=caption,
            tenant_id="boontrack-career"
        )

        self.assertIsNotNone(res)
        mock_send_text.assert_called_once_with("628123456789", caption, tenant_id="boontrack-career")


    # ==========================================
    # 8. document_jobs INSERT Guarantee Tests
    # ==========================================
    @patch("app.services.payment_service.get_supabase")
    def test_create_dynamic_order_inserts_to_document_jobs(self, mock_get_supabase):
        """Memvalidasi bahwa create_dynamic_order SELALU INSERT record ke document_jobs
        dengan price_amount = total_amount dan payment_status = 'UNPAID'."""
        from app.services.payment_service import payment_service as ps

        # Setup mock Supabase client
        mock_client = MagicMock()
        mock_table = MagicMock()
        mock_client.table.return_value = mock_table
        mock_table.upsert.return_value.execute.return_value = MagicMock(data=[])
        mock_table.insert.return_value.execute.return_value = MagicMock(data=[])
        mock_get_supabase.return_value = mock_client

        order = ps.create_dynamic_order(
            user_id="6281237450222",
            base_amount=25000,
            tenant_id="boontrack-career",
            meta={"product": "career_pro_bundle"}
        )

        # Pastikan order dibuat
        self.assertIn("total_amount", order)
        total_amount = order["total_amount"]
        self.assertGreater(total_amount, 25000)  # base + unique_code

        # Verifikasi document_jobs insert dipanggil
        table_calls = [call[0][0] for call in mock_client.table.call_args_list]
        self.assertIn("document_jobs", table_calls, "CRITICAL: document_jobs insert tidak dipanggil!")

        # Verifikasi field kritis dalam record document_jobs
        insert_record = mock_table.insert.call_args[0][0]
        self.assertEqual(insert_record["price_amount"], total_amount)
        self.assertEqual(insert_record["payment_status"], "UNPAID")
        self.assertEqual(insert_record["status"], "WAITING_PAYMENT")
        self.assertEqual(insert_record["user_id"], "6281237450222")

    @patch("app.services.payment_service.get_supabase")
    def test_create_dynamic_order_single_rewrite_document_jobs_record(self, mock_get_supabase):
        """Validasi Single CV Rewrite (10k) juga membuat record di document_jobs."""
        from app.services.payment_service import payment_service as ps

        mock_client = MagicMock()
        mock_table = MagicMock()
        mock_client.table.return_value = mock_table
        mock_table.upsert.return_value.execute.return_value = MagicMock(data=[])
        mock_table.insert.return_value.execute.return_value = MagicMock(data=[])
        mock_get_supabase.return_value = mock_client

        order = ps.create_dynamic_order(
            user_id="628888777666",
            base_amount=10000,
            tenant_id="boontrack-career",
            meta={"product": "single_cv_rewrite"}
        )

        total_amount = order["total_amount"]
        self.assertGreater(total_amount, 10000)

        table_calls = [call[0][0] for call in mock_client.table.call_args_list]
        self.assertIn("document_jobs", table_calls, "CRITICAL: document_jobs tidak di-insert untuk single CV rewrite!")

        # Verifikasi price_amount sesuai total_amount
        insert_record = mock_table.insert.call_args[0][0]
        self.assertEqual(insert_record["price_amount"], total_amount)
        self.assertEqual(insert_record["payment_status"], "UNPAID")

    @patch("app.services.document_engine.get_supabase")
    @patch("app.services.document_engine.r2_storage_service.upload_file", new_callable=AsyncMock)
    async def test_intake_document_job_invalid_file_still_inserts_document_jobs(
        self, mock_r2_upload, mock_get_supabase
    ):
        """Memvalidasi bahwa intake_document_job dengan file invalid (non-DOCX/PDF)
        TETAP INSERT record ke document_jobs jika exact_price_amount disediakan."""
        mock_client = MagicMock()
        mock_table = MagicMock()
        mock_client.table.return_value = mock_table
        mock_table.insert.return_value.execute.return_value = MagicMock(data=[])
        mock_get_supabase.return_value = mock_client

        # Kirim plain UTF-8 text bytes (bukan DOCX/PDF) — ini yang terjadi di rewrite handler
        plain_text_bytes = "Draft CV Budi Santoso. Software Engineer 3 tahun.".encode("utf-8")

        result = await intake_document_job(
            tenant_id="boontrack-career",
            task_type="CV_POLISH_REWRITE",
            filename="CV_Draft.docx",
            file_bytes=plain_text_bytes,
            user_id="6281237450222",
            user_phone="6281237450222",
            exact_price_amount=10432
        )

        # File memang invalid (REJECTED), tapi INSERT ke document_jobs HARUS tetap terjadi
        self.assertEqual(result["status"], "REJECTED")

        # Verifikasi bahwa document_jobs.insert dipanggil meski file invalid
        table_calls = [call[0][0] for call in mock_client.table.call_args_list]
        self.assertIn(
            "document_jobs", table_calls,
            "CRITICAL: document_jobs insert tidak dipanggil meski file invalid dan exact_price_amount=10432!"
        )

        # Verifikasi record yang di-insert memiliki price_amount yang tepat
        insert_record = mock_table.insert.call_args[0][0]
        self.assertEqual(insert_record["price_amount"], 10432)
        self.assertEqual(insert_record["payment_status"], "UNPAID")
        self.assertEqual(insert_record["status"], "WAITING_PAYMENT")

    def test_dynamic_output_document_filename_by_task_type(self):
        """Memvalidasi aturan penamaan file dinamis berdasarkan task_type."""
        from app.services.document_engine import get_output_document_filename

        # 1. POLISH_REPHRASE (Naskah skripsi / dokumen akademik)
        self.assertEqual(
            get_output_document_filename("POLISH_REPHRASE"),
            "Naskah_Hasil_Parafrase.docx"
        )
        self.assertEqual(
            get_output_document_filename("POLISH_REPHRASE", "BAB III WORD.pdf"),
            "BAB_III_WORD_Hasil_Parafrase.docx"
        )
        self.assertEqual(
            get_output_document_filename("PARAPHRASE", "Tesis_Bab_1.docx"),
            "Tesis_Bab_1_Hasil_Parafrase.docx"
        )

        # 2. CV_REVIEW / CV_ATS / CV_POLISH_REWRITE
        self.assertEqual(
            get_output_document_filename("CV_REVIEW"),
            "CV_Hasil_Optimasi_ATS.docx"
        )
        self.assertEqual(
            get_output_document_filename("CV_ATS"),
            "CV_Hasil_Optimasi_ATS.docx"
        )
        self.assertEqual(
            get_output_document_filename("CV_POLISH_REWRITE"),
            "CV_Hasil_Optimasi_ATS.docx"
        )

        # 3. CAREER_PRO_BUNDLE
        self.assertEqual(
            get_output_document_filename("CAREER_PRO_BUNDLE"),
            "Paket_Lengkap_Karir_ATS.docx"
        )


if __name__ == "__main__":
    unittest.main()
