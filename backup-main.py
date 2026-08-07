import os
import asyncio
import json
import re
import requests
from datetime import datetime
from dotenv import load_dotenv
import psycopg2
from psycopg2.extras import RealDictCursor
from aiogram import Bot, Dispatcher, executor, types
from aiogram.types import InlineKeyboardMarkup, InlineKeyboardButton, InputFile
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import google.generativeai as genai
from aiohttp import web

load_dotenv()

POSTGRES_HOST = os.getenv("POSTGRES_HOST", "postgres")
POSTGRES_PORT = os.getenv("POSTGRES_PORT", "5432")
POSTGRES_DB = os.getenv("POSTGRES_DB")
POSTGRES_USER = os.getenv("POSTGRES_USER")
POSTGRES_PASSWORD = os.getenv("POSTGRES_PASSWORD")

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "")
GROQ_API_KEY = os.getenv("GROQ_API_KEY", "")
OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY", "")

BOT_TOKEN = os.getenv("TELEGRAM_BOT_TOKEN")

if GEMINI_API_KEY:
    genai.configure(api_key=GEMINI_API_KEY)

bot = Bot(token=BOT_TOKEN)
dp = Dispatcher(bot)

user_state = {}

CV_QUESTIONS = {
    1: "👤 Siapa nama lengkapmu?",
    2: "📧 Email aktif yang bisa dihubungi recruiter?",
    3: "📱 Nomor WhatsApp / HP aktif? <i>(contoh: 081234567890)</i>",
    4: "📍 Di kota mana kamu berdomisili saat ini?",
    5: "🔗 Link akun LinkedIn kamu? <i>(Ketik '-' jika tidak ada)</i>",
    6: "🎯 Posisi/pekerjaan apa yang ingin kamu lamar?",
    7: (
        "💼 <b>Pengalaman kerja terakhirmu?</b>\n\n"
        "Tuliskan nama posisi, tempat kerja, dan tahunnya.\n"
        "Jika ada lebih dari 1 pekerjaan, pisahkan dengan <b>pindah baris (Enter)</b>, tanda <b>garis tegak ('|')</b>, atau <b>koma (',')</b>.\n\n"
        "<i>Contoh:</i>\n"
        "Kasir — Toko Makmur (2020 - 2022)\n"
        "Staff Admin — PT ABC (2022 - 2024)\n\n"
        "<i>(Ketik '-' jika fresh graduate)</i>"
    ),
    8: "🏆 Ceritakan tugas atau pencapaian utamamu di pekerjaan tersebut. <i>(Tulis santai saja, nanti saya bantu rapikan menjadi poin-poin profesional)</i>",
    9: "🎓 Pendidikan terakhirmu? <i>(contoh: S1 Manajemen, Universitas Terbuka, 2023)</i>",
    10: "🛠️ Apa saja skill atau keahlian utamamu? <i>(contoh: Microsoft Excel, Pelayanan Pelanggan, Kasir)</i>"
}

TOTAL_STEPS = 10

def get_progress_bar(step):
    return f"📍 <b>Langkah {step} dari {TOTAL_STEPS}</b>\n━━━━━━━━━━"

def get_db_connection():
    return psycopg2.connect(
        host=POSTGRES_HOST,
        port=POSTGRES_PORT,
        dbname=POSTGRES_DB,
        user=POSTGRES_USER,
        password=POSTGRES_PASSWORD
    )

def _init_db_sync():
    conn = get_db_connection()
    cur = conn.cursor()
    
    cur.execute("""
        CREATE TABLE IF NOT EXISTS users (
            telegram_id BIGINT PRIMARY KEY,
            username VARCHAR(255),
            first_name VARCHAR(255),
            last_name VARCHAR(255),
            language_code VARCHAR(10),
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        );
    """)
    
    cur.execute("""
        CREATE TABLE IF NOT EXISTS analytics (
            id SERIAL PRIMARY KEY,
            user_id BIGINT,
            event VARCHAR(100),
            meta JSONB,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        );
    """)

    cur.execute("""
        CREATE TABLE IF NOT EXISTS cv_documents (
            id SERIAL PRIMARY KEY,
            user_id BIGINT,
            version INT DEFAULT 1,
            position VARCHAR(255),
            data JSONB,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        );
    """)

    cur.execute("""
        CREATE TABLE IF NOT EXISTS user_progress (
            user_id BIGINT PRIMARY KEY,
            last_step INT,
            data JSONB,
            updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        );
    """)
    
    conn.commit()
    cur.close()
    conn.close()

async def init_db():
    await asyncio.to_thread(_init_db_sync)

def _track_event_sync(user_id, event, meta=None):
    try:
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("INSERT INTO analytics (user_id, event, meta) VALUES (%s, %s, %s)",
                    (user_id, event, json.dumps(meta or {})))
        conn.commit()
        cur.close()
        conn.close()
    except Exception as e:
        print(f"Analytics DB Error: {e}")

async def track_event(user_id, event, meta=None):
    await asyncio.to_thread(_track_event_sync, user_id, event, meta)

def _save_user_sync(user: types.User):
    try:
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("""
            INSERT INTO users (telegram_id, username, first_name, last_name, language_code)
            VALUES (%s, %s, %s, %s, %s)
            ON CONFLICT (telegram_id) DO UPDATE SET
                username = EXCLUDED.username,
                first_name = EXCLUDED.first_name,
                last_name = EXCLUDED.last_name;
        """, (user.id, user.username, user.first_name, user.last_name, user.language_code))
        conn.commit()
        cur.close()
        conn.close()
    except Exception as e:
        print(f"Save User DB Error: {e}")

async def save_user(user: types.User):
    await asyncio.to_thread(_save_user_sync, user)

def _save_dropoff_sync(user_id, step, data):
    try:
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("""
            INSERT INTO user_progress (user_id, last_step, data, updated_at)
            VALUES (%s, %s, %s, CURRENT_TIMESTAMP)
            ON CONFLICT (user_id) DO UPDATE SET
                last_step = EXCLUDED.last_step,
                data = EXCLUDED.data,
                updated_at = CURRENT_TIMESTAMP;
        """, (user_id, step, json.dumps(data)))
        conn.commit()
        cur.close()
        conn.close()
    except Exception as e:
        print(f"Dropoff DB Error: {e}")

async def save_dropoff(user_id, step, data):
    await asyncio.to_thread(_save_dropoff_sync, user_id, step, data)

def _get_user_history_sync(user_id):
    try:
        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=RealDictCursor)
        
        cur.execute("SELECT last_step, data FROM user_progress WHERE user_id = %s", (user_id,))
        progress = cur.fetchone()
        
        cur.execute("SELECT version, position, data, created_at FROM cv_documents WHERE user_id = %s ORDER BY id DESC LIMIT 1", (user_id,))
        last_cv = cur.fetchone()
        
        cur.close()
        conn.close()
        return progress, last_cv
    except Exception as e:
        print(f"Get User History Error: {e}")
        return None, None

async def get_user_history(user_id):
    return await asyncio.to_thread(_get_user_history_sync, user_id)

def _save_cv_version_sync(user_id, position, data):
    try:
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("SELECT COUNT(*) FROM cv_documents WHERE user_id = %s", (user_id,))
        count = cur.fetchone()[0]
        version = count + 1
        
        cur.execute("""
            INSERT INTO cv_documents (user_id, version, position, data)
            VALUES (%s, %s, %s, %s)
        """, (user_id, version, position, json.dumps(data)))
        conn.commit()
        cur.close()
        conn.close()
    except Exception as e:
        print(f"Save CV Version Error: {e}")

async def save_cv_version(user_id, position, data):
    await asyncio.to_thread(_save_cv_version_sync, user_id, position, data)

def clean_val(val):
    if not val:
        return ""
    v = val.strip().lower()
    if v in ["-", "skip", "tidak ada", "ga ada", "ngga ada", "belum ada", "hangus", "hilang", "lupa", "kosong"]:
        return ""
    return val.strip()

def ai_generate_summary(position):
    if not position:
        return "Profesional berdedikasi tinggi yang terbiasa bekerja secara terstruktur, adaptif, serta berkomitmen memberikan kontribusi operasional terbaik bagi perusahaan."
    return f"Profesional berpengalaman di bidang {position} dengan rekam jejak yang terbukti dalam mengeksekusi target operasional dan manajemen kerja secara efisien."

def ai_rewrite_achievement(text):
    prompt_text = (
        "Ubah deskripsi tugas/pengalaman kerja berikut menjadi 2-4 poin bullet point (menggunakan simbol •) "
        "berbahasa Indonesia profesional standar HR yang ATS-friendly. Buat lugas, aksi-orientasi, dan profesional:\n\n"
        f"Input Pengalaman: {text}"
    )

    if GEMINI_API_KEY:
        try:
            model = genai.GenerativeModel('gemini-2.5-flash')
            response = model.generate_content(prompt_text)
            if response and response.text:
                return response.text.strip()
        except Exception as e:
            print(f"[Fallback Alert] Gemini API Error/Limit: {e}. Beralih ke Groq...")

    if GROQ_API_KEY:
        try:
            headers = {
                "Authorization": f"Bearer {GROQ_API_KEY}",
                "Content-Type": "application/json"
            }
            payload = {
                "model": "llama3-8b-8192",
                "messages": [{"role": "user", "content": prompt_text}],
                "temperature": 0.5
            }
            res = requests.post("https://api.groq.com/openai/v1/chat/completions", json=payload, headers=headers, timeout=3)
            if res.status_code == 200:
                return res.json()['choices'][0]['message']['content'].strip()
        except Exception as e:
            print(f"[Fallback Alert] Groq API Error: {e}. Beralih ke OpenRouter...")

    if OPENROUTER_API_KEY:
        try:
            headers = {
                "Authorization": f"Bearer {OPENROUTER_API_KEY}",
                "Content-Type": "application/json"
            }
            payload = {
                "model": "meta-llama/llama-3-8b-instruct:free",
                "messages": [{"role": "user", "content": prompt_text}]
            }
            res = requests.post("https://openrouter.ai/api/v1/chat/completions", json=payload, headers=headers, timeout=3)
            if res.status_code == 200:
                return res.json()['choices'][0]['message']['content'].strip()
        except Exception as e:
            print(f"[Fallback Alert] OpenRouter Error: {e}. Gunakan aturan lokal.")

    clean = text.lower()
    if "galon" in clean or "antar" in clean or "kurir" in clean or "ojek" in clean or "driver" in clean:
        return "• Mengantarkan produk atau barang kepada pelanggan secara tepat waktu dan aman.\n• Menjaga komunikasi yang baik dengan pelanggan guna memastikan kepuasan pelayanan.\n• Mencatat dan melaporkan setiap transaksi pengiriman harian secara akurat."
    elif "jaga" in clean or "warung" in clean or "toko" in clean or "kasir" in clean or "semen" in clean:
        return "• Melayani pelanggan dengan ramah dan profesional guna meningkatkan kepuasan transaksi.\n• Mengelola transaksi serta pencatatan inventaris harian secara akurat.\n• Memastikan kebersihan dan keteraturan area kerja."
    else:
        lines = [line.strip().lstrip("-*• ") for line in text.split("\n") if line.strip()]
        formatted = "\n".join([f"• {line}" for line in lines])
        if len(lines) == 1:
            formatted += "\n• Berkomitmen melaksanakan seluruh tanggung jawab kerja dengan kedisiplinan dan ketelitian tinggi.\n• Berkontribusi aktif dalam mendukung tercapainya efisiensi target operasional tim."
        return formatted

def create_cv_docx(user_id, data):
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    name = clean_val(data.get("step_1", "NAMA LENGKAP"))
    email = clean_val(data.get("step_2", ""))
    phone = clean_val(data.get("step_3", ""))
    domicile = clean_val(data.get("step_4", ""))
    linkedin = clean_val(data.get("step_5", ""))

    p_name = doc.add_paragraph()
    p_name.paragraph_format.space_after = Pt(2)
    r_name = p_name.add_run(name.upper())
    r_name.font.name = 'Calibri'
    r_name.font.size = Pt(16)
    r_name.font.bold = True
    r_name.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    contact_parts = [p for p in [email, phone, domicile, linkedin] if p]
    if contact_parts:
        p_contact = doc.add_paragraph()
        p_contact.paragraph_format.space_after = Pt(12)
        r_contact = p_contact.add_run(" | ".join(contact_parts))
        r_contact.font.name = 'Calibri'
        r_contact.font.size = Pt(10)
        r_contact.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    def add_section_header(title):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(10)
        h.paragraph_format.space_after = Pt(4)
        run = h.add_run(title.upper())
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
        
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '1F4E78')
        pBdr.append(bottom)
        h._p.get_or_add_pPr().append(pBdr)

    add_section_header("PROFESSIONAL SUMMARY")
    summary_text = ai_generate_summary(clean_val(data.get("step_6", "")))
    p_sum = doc.add_paragraph(summary_text)
    p_sum.paragraph_format.space_after = Pt(8)
    for r in p_sum.runs:
        r.font.name = 'Calibri'
        r.font.size = Pt(10.5)

    exp = clean_val(data.get("step_7", ""))
    ach_raw = clean_val(data.get("step_8", ""))

    if exp:
        add_section_header("PROFESSIONAL EXPERIENCE")
        raw_jobs = [j.strip() for j in re.split(r'[\n|]', exp) if j.strip()]
        
        for job_title in raw_jobs:
            if not job_title:
                continue
            
            p_job = doc.add_paragraph()
            p_job.paragraph_format.space_before = Pt(6)
            p_job.paragraph_format.space_after = Pt(2)
            r_job = p_job.add_run(job_title)
            r_job.font.name = 'Calibri'
            r_job.font.size = Pt(10.5)
            r_job.font.bold = True

            if ach_raw:
                ach_formatted = ai_rewrite_achievement(ach_raw)
                for bullet in ach_formatted.split("\n"):
                    b_text = bullet.strip().lstrip("-*• ")
                    if b_text:
                        p_b = doc.add_paragraph(style='List Bullet')
                        p_b.paragraph_format.space_after = Pt(2)
                        r_b = p_b.add_run(b_text)
                        r_b.font.name = 'Calibri'
                        r_b.font.size = Pt(10)

    edu = clean_val(data.get("step_9", ""))
    if edu:
        add_section_header("EDUCATION")
        p_edu = doc.add_paragraph(edu)
        p_edu.paragraph_format.space_after = Pt(8)
        for r in p_edu.runs:
            r.font.name = 'Calibri'
            r.font.size = Pt(10.5)

    skill = clean_val(data.get("step_10", ""))
    if skill:
        add_section_header("SKILLS")
        for line in skill.split("\n"):
            line_str = line.strip()
            if line_str:
                p_skill = doc.add_paragraph(line_str)
                p_skill.paragraph_format.space_after = Pt(3)
                for r in p_skill.runs:
                    r.font.name = 'Calibri'
                    r.font.size = Pt(10)

    file_path = f"/tmp/CV_{user_id}.docx"
    doc.save(file_path)
    return file_path

@dp.message_handler(commands=['start'])
async def send_welcome(message: types.Message):
    user_id = message.from_user.id
    await save_user(message.from_user)
    
    text_parts = message.text.split()
    args = text_parts[1] if len(text_parts) > 1 else "direct"

    meta_data = {}
    if args.startswith("ref_"):
        meta_data = {
            "utm_source": "referral",
            "referrer_id": args.replace("ref_", "")
        }
    else:
        meta_data = {
            "utm_source": args
        }

    await track_event(user_id, "start", meta=meta_data)
    
    progress, last_cv = await get_user_history(user_id)
    first_name = message.from_user.first_name or "Teman"

    if progress and progress.get("last_step", 0) > 1 and progress.get("last_step", 0) < TOTAL_STEPS:
        last_step = progress["last_step"]
        saved_data = progress.get("data", {})
        user_state[user_id] = {"step": last_step, "data": saved_data}
        
        kbd = InlineKeyboardMarkup(row_width=2)
        kbd.add(
            InlineKeyboardButton("▶️ Lanjutkan CV", callback_data="resume_flow"),
            InlineKeyboardButton("🔄 Mulai Baru", callback_data="restart_flow")
        )
        
        await message.reply(
            f"Halo lagi, {first_name}! 👋\n\n"
            f"Kemarin kita sempat membuat CV sampai di <b>Langkah {last_step} dari {TOTAL_STEPS}</b>.\n\n"
            "Mau kita tuntaskan sekarang agar CV kamu siap dipakai melamar kerja?\n\n"
            "💡 <i>Tips: Ketik /cancel kapan saja jika ingin membatalkan atau mengulang dari awal.</i>",
            reply_markup=kbd,
            parse_mode="HTML"
        )
        return

    if last_cv:
        last_pos = last_cv.get("position", "pekerjaan kamu")
        kbd = InlineKeyboardMarkup(row_width=2)
        kbd.add(
            InlineKeyboardButton("📄 Buat CV Posisi Lain", callback_data="restart_flow"),
            InlineKeyboardButton("💡 FAQ & Tips ATS", callback_data="show_faq")
        )
        
        await message.reply(
            f"Halo kembali, {first_name}! 👋\n\n"
            f"Terakhir kali kita membuat CV untuk posisi <b>{last_pos}</b>.\n"
            "Ada yang bisa saya bantu hari ini?\n\n"
            "💡 <i>Tips: Ketik /cancel kapan saja untuk membatalkan proses.</i>",
            reply_markup=kbd,
            parse_mode="HTML"
        )
        return

    user_state[user_id] = {"step": 1, "data": {}}
    await save_dropoff(user_id, 1, {})
    
    greeting = (
        "<b>👋 Halo! Saya BoonTrack Assistant.</b>\n\n"
        "Saya akan membantumu membuat CV yang bersih, profesional, dan mudah dibaca oleh HR serta sistem rekrutmen perusahaan modern (seperti JobStreet, LinkedIn, Glints, dll).\n\n"
        "<b>📌 Catatan Penting:</b>\n"
        "CV ini sengaja dibuat <b>tanpa foto & desain berlebihan</b> agar fokus utama HR langsung ke pengalaman kerjamu, dan peluang lolos seleksi awal jauh lebih besar.\n\n"
        "Cukup jawab beberapa pertanyaan singkat (~5 menit) dan hasilnya siap di-download dalam format Word (.docx).\n\n"
        "💡 <i>Tips: Ketik /cancel kapan saja jika ingin membatalkan atau mengulang dari awal.</i>\n\n"
        "Kalau sudah siap, kita mulai ya 😃\n\n"
        f"{get_progress_bar(1)}\n"
        f"{CV_QUESTIONS[1]}"
    )
    await message.reply(greeting, parse_mode="HTML")

@dp.callback_query_handler(lambda c: c.data in ["resume_flow", "restart_flow", "show_faq"])
async def handle_callback_navigation(callback_query: types.CallbackQuery):
    user_id = callback_query.from_user.id
    code = callback_query.data
    
    await bot.edit_message_reply_markup(user_id, callback_query.message.message_id, reply_markup=None)
    
    if code == "resume_flow":
        state = user_state.get(user_id, {"step": 1, "data": {}})
        step = state["step"]
        await bot.send_message(
            user_id,
            f"Sip, mari kita lanjutkan! 👍\n\n{get_progress_bar(step)}\n{CV_QUESTIONS[step]}",
            parse_mode="HTML"
        )
    elif code == "restart_flow":
        user_state[user_id] = {"step": 1, "data": {}}
        await save_dropoff(user_id, 1, {})
        await bot.send_message(
            user_id,
            f"Sip, kita mulai dari awal ya! 😊\n\n{get_progress_bar(1)}\n{CV_QUESTIONS[1]}",
            parse_mode="HTML"
        )
    elif code == "show_faq":
        faq_text = (
            "❓ <b>Kenapa CV BoonTrack Tanpa Foto?</b>\n"
            "Banyak perusahaan saat ini lebih berfokus pada skill & pengalaman. Format bersih tanpa foto juga memastikan CV mudah dibaca sistem ATS tanpa error.\n\n"
            "❓ <b>Apakah File .docx Bisa Di-edit?</b>\n"
            "Bisa banget! Kamu bebas mengedit kembali tulisan atau menambah foto secara manual di Microsoft Word jika melamar ke perusahaan yang mewajibkannya."
        )
        await bot.send_message(user_id, faq_text, parse_mode="HTML")

@dp.message_handler(commands=['cancel'])
async def cancel_handler(message: types.Message):
    user_id = message.from_user.id
    user_state[user_id] = {"step": 0, "data": {}}
    await save_dropoff(user_id, 0, {})
    await message.reply("❌ <b>Proses pembuatan CV dibatalkan.</b>\n\nKetik /start kapan saja jika ingin memulai kembali dari awal!", parse_mode="HTML")

@dp.message_handler()
async def handle_message(message: types.Message):
    user_id = message.from_user.id
    text = message.text.strip()

    if user_id not in user_state or user_state[user_id].get("step", 0) == 0:
        progress, _ = await get_user_history(user_id)
        if progress and progress.get("last_step", 0) > 0:
            user_state[user_id] = {"step": progress["last_step"], "data": progress.get("data", {})}
        else:
            await message.reply("Ketik <b>/start</b> untuk mulai membuat CV ya!", parse_mode="HTML")
            return

    state = user_state[user_id]
    step = state["step"]

    if step == 1:
        if len(text) < 2 or text.lower() in ["saya", "gak tahu", "gak tau", "nggak tahu", "bot", "tes", "test"]:
            await message.reply("Maaf, sepertinya itu bukan nama lengkapmu 😊\n\nBolehkah minta tolong tuliskan nama lengkap sesuai KTP/Ijazah?")
            return

    elif step == 2:
        if text != "-" and not re.match(r"[^@]+@[^@]+\.[^@]+", text):
            await message.reply(
                "Sepertinya format email belum sesuai (contoh: <i>budi@gmail.com</i>) 😊\n\n"
                "CV modern sebaiknya memiliki email agar recruiter mudah menghubungi.\n"
                "Kalau belum punya, ketik <b>-</b> untuk dikosongkan dulu ya!",
                parse_mode="HTML"
            )
            return

    elif step == 3:
        lower_t = text.lower()
        if any(w in lower_t for w in ["hangus", "hilang", "lupa", "mati", "tidak ada", "ga ada"]):
            await message.reply(
                "Tidak apa-apa 😊\n\n"
                "Kalau nomor HP kamu sedang tidak aktif, kamu bisa ketik <b>-</b> untuk mengosongkannya terlebih dahulu di CV."
            )
            return

        digits = re.sub(r"\D", "", text)
        if text != "-" and (len(digits) < 10 or not (digits.startswith("08") or digits.startswith("628"))):
            await message.reply(
                "Sepertinya nomor HP/WhatsApp yang kamu masukkan belum lengkap 😊\n\n"
                "Nomor seluler yang valid biasanya terdiri dari <b>10–13 digit</b> (contoh: <i>081234567890</i>).\n\n"
                "Coba periksa dan masukkan ulang ya! Atau ketik <b>-</b> jika belum ingin mencantumkannya di CV.",
                parse_mode="HTML"
            )
            return

    elif step == 8 and state.get("sub_step") != "daily_tasks":
        lower_text = text.lower()
        if len(text.split()) < 3 or any(w in lower_text for w in ["belum", "tidak ada", "ngga ada", "ga ada", "bingung", "cuma", "pikir"]):
            state["sub_step"] = "daily_tasks"
            await message.reply(
                "Tidak apa-apa 😊 Tuliskan saja dengan santai apa tugas atau kegiatanmu sehari-hari di pekerjaan tersebut (misal: <i>antar galon ke pelanggan</i>, <i>jaga toko</i>, <i>input data</i>).\n\n"
                "Nanti saya akan bantu menerjemahkannya menjadi kalimat profesional di CV kamu!"
            )
            return
            
    if state.get("sub_step") == "daily_tasks":
        state["sub_step"] = None

    state["data"][f"step_{step}"] = text
    await save_dropoff(user_id, step, state["data"])
    
    ack = ""
    if step == 1:
        ack = f"Sip 👍 Senang bertemu denganmu, <b>{text}</b>!\n\n"
    elif step == 6:
        ack = f"Sip, posisi <b>{text}</b>. Saya catat ya!\n\n"
    elif step == 7:
        ack = "Terima kasih! Informasi pengalaman kerjamu sudah dicatat 👍\n\n"
    elif step == 8:
        ack = "Mantap! Poin pengalamanmu sudah diperkaya dengan bahasa profesional oleh AI 😊\n\n"

    if step < TOTAL_STEPS:
        next_step = step + 1
        state["step"] = next_step
        
        cheer = ""
        if next_step == 4:
            cheer = "Sedikit lagi kita menyelesaikan data diri... 😊\n\n"
        elif next_step == 9:
            cheer = "Sedikit lagi CV kamu selesai!\n\n"

        prompt = f"{ack}{get_progress_bar(next_step)}\n{cheer}{CV_QUESTIONS[next_step]}"
        await message.reply(prompt, parse_mode="HTML")
    else:
        await track_event(user_id, "finish_cv")
        
        msg_load = await message.reply("🛠️ <b>Memproses AI & merapikan bahasa pengalaman kerjamu...</b>\n<code>[██████░░░░]</code>", parse_mode="HTML")
        await asyncio.sleep(0.5)
        
        await msg_load.edit_text("✍️ <b>Menyesuaikan format profesional & tata letak Word...</b>\n<code>[████████░░]</code>", parse_mode="HTML")
        await asyncio.sleep(0.5)

        await msg_load.edit_text("📄 <b>Hampir selesai...</b>\n<code>[██████████]</code>", parse_mode="HTML")
        
        file_path = None
        try:
            file_path = await asyncio.to_thread(create_cv_docx, user_id, state["data"])
            await save_cv_version(user_id, state["data"].get("step_6", "General"), state["data"])
            
            await bot.send_document(
                chat_id=user_id, 
                document=InputFile(file_path), 
                caption=(
                    "🎉 <b>CV profesionalmu sudah selesai!</b>\n\n"
                    "Selamat. Satu langkah penting menuju pekerjaan impianmu sudah berhasil kamu selesaikan hari ini.\n\n"
                    "📄 <b>Kenapa CV ini sederhana & tanpa foto?</b>\n"
                    "Banyak perusahaan saat ini menggunakan sistem otomatis (ATS) dan lebih memfokuskan penilaian pada pengalaman kerja serta keahlian. Format bersih ini dirancang agar mudah dibaca sistem maupun HRD."
                ),
                parse_mode="HTML"
            )
            await track_event(user_id, "download_cv")
        finally:
            state["step"] = 0
            await save_dropoff(user_id, 0, {})
            if file_path and os.path.exists(file_path):
                os.remove(file_path)

        await asyncio.sleep(1.5)
        touch_msg = (
            "💡 <b>Catatan Karir:</b>\n"
            "Mencari pekerjaan memang membutuhkan ketekunan. Semoga CV yang baru kamu buat hari ini menjadi perantara terbaik menuju karir impianmu. 🚀\n\n"
            "Kalau suatu hari nanti kamu menerima kabar baik atau panggilan interview, kembali ke sini ya dan ceritakan kabar baiknya ke saya! 😊"
        )
        await bot.send_message(user_id, touch_msg, parse_mode="HTML")

        await asyncio.sleep(1)
        keyboard = InlineKeyboardMarkup(row_width=2)
        keyboard.add(
            InlineKeyboardButton("⭐⭐⭐⭐⭐ Mantap", callback_data="feedback_5star"),
            InlineKeyboardButton("👍 Membantu", callback_data="feedback_membantu"),
            InlineKeyboardButton("😐 Lumayan", callback_data="feedback_lumayan"),
            InlineKeyboardButton("👎 Belum", callback_data="feedback_belum")
        )
        await bot.send_message(user_id, "Saya penasaran 😊\nMenurutmu, hasil CV ini bagaimana?", reply_markup=keyboard)

@dp.callback_query_handler(lambda c: c.data.startswith('feedback_'))
async def process_feedback(callback_query: types.CallbackQuery):
    user_id = callback_query.from_user.id
    code = callback_query.data
    await track_event(user_id, "feedback", {"choice": code})
    
    await bot.edit_message_reply_markup(user_id, callback_query.message.message_id, reply_markup=None)
    
    if code in ["feedback_5star", "feedback_membantu"]:
        bot_info = await bot.get_me()
        bot_link = f"https://t.me/{bot_info.username}"
        
        qris_text = (
            "Saya senang sekali mendengarnya! ❤️\n\n"
            "Hari ini kamu mendapatkan CV profesional tanpa biaya sama sekali. Kalau suatu hari nanti CV ini membantumu mendapatkan interview, aku akan sangat bahagia!\n\n"
            "Kalau kamu ingin ikut menjaga BoonTrack tetap hidup supaya bisa membantu pencari kerja lainnya, kamu boleh mentraktir developer secangkir kopi ☕ (Rp10.000–15.000).\n\n"
            "💳 <b>NMID QRIS:</b> ID1026564075103 (BoonTrack)\n\n"
            "📌 <b>Bantu Temanmu:</b>\n"
            f"Kalau punya teman yang sedang cari kerja, bagikan bot ini ya: {bot_link}\n\n"
            "Terima kasih banyak dan semoga sukses ya! 🚀"
        )
        
        possible_paths = [
            "/app/qris.jpg",
            "qris.jpg",
            "/root/boontrack-core/qris.jpg"
        ]
        
        found_qris_path = None
        for path in possible_paths:
            if os.path.exists(path):
                found_qris_path = path
                break

        if found_qris_path:
            try:
                await bot.send_photo(chat_id=user_id, photo=InputFile(found_qris_path), caption=qris_text, parse_mode="HTML")
            except Exception as err:
                print(f"Error mengirim foto QRIS: {err}")
                await bot.send_message(user_id, qris_text, parse_mode="HTML")
        else:
            await bot.send_message(user_id, qris_text, parse_mode="HTML")
            
        await track_event(user_id, "donation_clicked")
        
    else:
        await bot.send_message(user_id, "Terima kasih banyak atas masukannya! Kami akan terus membenahi alurnya agar semakin membantu ke depannya 😊 Ketik /start jika ingin mencoba lagi.")

async def health_check_handler(request):
    return web.json_response({"status": "healthy", "message": "Render is awake!"}, status=200)

async def start_web_server():
    app = web.Application()
    app.router.add_get('/', health_check_handler)
    app.router.add_get('/health', health_check_handler)
    
    port = int(os.getenv("PORT", 10000))
    runner = web.AppRunner(app)
    await runner.setup()
    
    try:
        site = web.TCPSite(runner, '0.0.0.0', port)
        await site.start()
        print(f"Health Check Web Server berjalan di port {port}...")
    except OSError:
        fallback_port = port + 1
        site = web.TCPSite(runner, '0.0.0.0', fallback_port)
        await site.start()
        print(f"Port {port} terpakai, Health Check Server dialihkan ke port {fallback_port}...")

if __name__ == '__main__':
    print("Inisialisasi Database & Schema (Async)...")
    loop = asyncio.new_event_loop()
    asyncio.set_event_loop(loop)
    loop.run_until_complete(init_db())
    
    loop.create_task(start_web_server())
    
    print("Bot Telegram Boontrack Berjalan...")
    executor = executor.Executor(dp, skip_updates=True, loop=loop)
    executor.start_polling()